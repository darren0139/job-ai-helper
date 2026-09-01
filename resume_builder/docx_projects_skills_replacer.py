"""
resume_builder/docx_projects_skills_replacer.py

Updated DOCX resume copy replacer.

Workflow:
    Upload original_resume.docx
    -> optionally save a local copy
    -> generate tailored projects + tailored skills
    -> copy original DOCX
    -> replace only SKILLS and PROJECTS sections
    -> download tailored resume copy

This version improves formatting preservation:
1. Copies paragraph style from the original Projects section.
2. Copies bullet style/numbering from the original bullet paragraphs.
3. Uses a right-aligned tab stop for project dates.
4. Preserves spacing between projects where possible.
5. Inserts real bullet paragraphs instead of plain text paragraphs.

Important:
    This does not overwrite the original uploaded file.
    Work Experience is not changed.
"""



from __future__ import annotations

import base64
import hashlib
import os
import re
import shutil
import subprocess
import time
import uuid
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.document import Document as DocumentObject
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from utils.date_sorting import period_sort_value
from resume_builder.project_header_format import (
    build_project_title,
    format_project_metadata,
    normalise_project_header_layout,
    normalise_project_metadata_style,
    split_legacy_project_title,
)
from resume_builder.skills_section_compactor import (
    compact_skills_one_step,
    count_skill_items,
    count_skill_reduction_candidates,
    restore_skill_change,
    skill_restoration_quality_gain,
)
from resume_builder.evidence_aware_fitting import (
    PHASE6C_FITTING_VERSION,
    build_evidence_aware_project_reductions,
    restore_removed_bullet_metadata,
    sync_project_bullet_metadata,
)
from resume_builder.fitting_render_optimizer import (
    PHASE6C1_OPTIMIZATION_VERSION,
    build_render_state_fingerprint,
    group_candidates_by_protection_tier,
    rendered_candidate_is_effective,
    source_docx_signature,
)
from resume_builder.fitting_provenance import (
    PHASE6C_SEARCH_ALGORITHM_VERSION,
    build_fitting_input_snapshot,
)
from tailoring.tailoring_generation_fingerprint import (
    build_fitting_lock_policy,
)

SAVED_RESUME_DIR = Path("data/saved_resumes")
TAILORED_RESUME_DIR = Path("outputs/tailored_resumes")
PREVIEW_DIR = Path("outputs/resume_previews")

# Phase 6C renders actual DOCX/PDF output as the fit authority. These limits
# bound external work without making any fitting or evidence decision from an
# estimate alone.
FIT_RENDER_BUDGET = 96
FIT_COARSE_INITIAL_REMOVALS = 4
FIT_LOCAL_REFINEMENT_LIMIT = 8
FIT_RESTORATION_RENDER_BUDGET = 8
DEFAULT_MINIMUM_TOTAL_SKILLS = 8
LIBREOFFICE_SINGLE_CONVERSION_TIMEOUT_SECONDS = 120
LIBREOFFICE_BATCH_BASE_TIMEOUT_SECONDS = 60
LIBREOFFICE_BATCH_TIMEOUT_PER_CANDIDATE_SECONDS = 15
LIBREOFFICE_BATCH_MAX_TIMEOUT_SECONDS = 300

KNOWN_SECTION_HEADINGS = {
    "EDUCATION",
    "WORK EXPERIENCE",
    "EXPERIENCE",
    "PROFESSIONAL EXPERIENCE",
    "PROJECTS",
    "SKILLS",
    "TECHNICAL SKILLS",
    "CERTIFICATIONS",
    "ACHIEVEMENTS",
    "AWARDS",
    "COURSEWORK",
    "SUMMARY",
    "PROFILE",
}

from tailoring.deterministic_project_rules import (
    apply_deterministic_evidence_floors,
)

# ---------------------------------------------------------------------------
# File save helpers
# ---------------------------------------------------------------------------

def _safe_filename(filename: str) -> str:
    """Create a safe filename for local storage."""
    name = Path(filename).name
    name = re.sub(r"[^A-Za-z0-9_. -]", "_", name)
    return name.strip() or "uploaded_resume.docx"


def save_uploaded_docx_for_editing(
    uploaded_file: Any,
    *,
    application_id: int | None = None,
    replace_existing: bool = True,
) -> Path:
    """
    Save the uploaded DOCX so the app can later generate an edited copy.

    This should only be called if the user ticks the save checkbox.

    Keeps only the latest saved DOCX for the same application session
    when replace_existing=True.
    """
    if uploaded_file is None:
        raise ValueError("No uploaded file provided.")

    if not uploaded_file.name.lower().endswith(".docx"):
        raise ValueError("Only DOCX files can be saved for editable resume-copy generation.")

    SAVED_RESUME_DIR.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    app_part = f"app_{application_id}_" if application_id is not None else ""
    filename = f"{app_part}{timestamp}_{_safe_filename(uploaded_file.name)}"
    saved_path = SAVED_RESUME_DIR / filename

    saved_path.write_bytes(uploaded_file.getbuffer())

    if replace_existing and application_id is not None:
        for old_path in SAVED_RESUME_DIR.glob(f"app_{application_id}_*.docx"):
            if old_path != saved_path:
                try:
                    old_path.unlink()
                except OSError:
                    pass

    return saved_path


# ---------------------------------------------------------------------------
# Section detection helpers
# ---------------------------------------------------------------------------

def _paragraph_text(paragraph: Paragraph) -> str:
    """Return normalized paragraph text."""
    return " ".join(paragraph.text.replace("\xa0", " ").split())


def _is_heading_named(paragraph: Paragraph, names: set[str]) -> bool:
    """Return True if paragraph is one of the target headings."""
    text = _paragraph_text(paragraph).strip().rstrip(":")
    return text.upper() in {name.upper() for name in names}


def _is_probable_section_heading(paragraph: Paragraph) -> bool:
    """
    Detect the next resume section heading.

    This is intentionally conservative. It checks:
    - Word heading styles
    - known resume heading names
    - short all-caps headings
    """
    text = _paragraph_text(paragraph).strip().rstrip(":")
    if not text:
        return False

    style_name = ""
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        style_name = ""

    if "Heading" in style_name:
        return True

    upper_text = text.upper()

    if upper_text in KNOWN_SECTION_HEADINGS:
        return True

    # Example: "EDUCATION", "PROJECTS", "TECHNICAL SKILLS"
    # Avoid treating long bullet sentences as headings.
    if (
        text == upper_text
        and 1 <= len(text.split()) <= 4
        and len(text) <= 40
        and not text.startswith("•")
        and not text.startswith("-")
    ):
        return True

    return False


def _find_section_range(
    document: DocumentObject,
    heading_names: set[str],
) -> tuple[int, int | None]:
    """
    Find a section by heading name.

    Returns:
        start_index: index of heading paragraph
        end_index: index of next section heading, or None if this is last section
    """
    paragraphs = document.paragraphs

    start_index = None
    for index, paragraph in enumerate(paragraphs):
        if _is_heading_named(paragraph, heading_names):
            start_index = index
            break

    if start_index is None:
        raise ValueError(f"Could not find section heading: {', '.join(sorted(heading_names))}")

    end_index = None
    for index in range(start_index + 1, len(paragraphs)):
        if _is_probable_section_heading(paragraphs[index]):
            end_index = index
            break

    return start_index, end_index


def _is_bullet_paragraph(paragraph: Paragraph) -> bool:
    """Return True if paragraph appears to be a bullet/list item."""
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style else ""

    has_numbering = (
        paragraph._p.pPr is not None
        and paragraph._p.pPr.numPr is not None
    )

    return has_numbering or "Bullet" in style_name or text.startswith("•")


def _find_templates_in_section(
    document: DocumentObject,
    heading_names: set[str],
) -> tuple[Paragraph | None, Paragraph | None]:
    """
    Find likely title/normal paragraph template and bullet template in a section.

    Returns:
        normal_template: first non-empty non-bullet paragraph
        bullet_template: first bullet/list paragraph
    """
    start_index, end_index = _find_section_range(document, heading_names)
    paragraphs = document.paragraphs
    section_end = end_index if end_index is not None else len(paragraphs)

    normal_template = None
    bullet_template = None

    for paragraph in paragraphs[start_index + 1 : section_end]:
        text = paragraph.text.strip()
        if not text:
            continue

        if bullet_template is None and _is_bullet_paragraph(paragraph):
            bullet_template = paragraph
            continue

        if normal_template is None and not _is_bullet_paragraph(paragraph):
            normal_template = paragraph

        if normal_template is not None and bullet_template is not None:
            break

    return normal_template, bullet_template


def _normalise_source_project_key(value: Any) -> str:
    text = str(value or "").casefold().strip()
    text = re.sub(r"\(.*?\)", "", text)
    text = re.sub(r"[-–—].*$", "", text)
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return " ".join(text.split())


def _split_source_project_title(value: Any) -> dict[str, Any]:
    """Parse source-DOCX project title semantics without inventing metadata."""
    raw = str(value or "").replace("\u00a0", " ").strip()
    title_part = raw.split("\t", 1)[0].strip()
    legacy = split_legacy_project_title(title_part)
    semantic = " ".join(str(legacy.get("title") or "").split()).strip()
    title = semantic
    subtitle = ""
    for separator in (" — ", " – "):
        if separator in semantic:
            left, right = semantic.split(separator, 1)
            if left.strip() and right.strip():
                title = left.strip()
                subtitle = right.strip()
                break
    return {
        "title": title,
        "subtitle": subtitle,
        "resume_header_tools": list(legacy.get("resume_header_tools") or []),
        "resume_header_context": list(legacy.get("resume_header_context") or []),
    }


def _extract_source_project_display_metadata(
    document: DocumentObject,
) -> dict[str, dict[str, Any]]:
    """Read display-only project metadata from the exact source DOCX."""
    try:
        start_index, end_index = _find_section_range(document, {"PROJECTS"})
    except ValueError:
        return {}

    paragraphs = document.paragraphs
    section_end = end_index if end_index is not None else len(paragraphs)
    result: dict[str, dict[str, Any]] = {}
    current_key = ""

    for paragraph in paragraphs[start_index + 1 : section_end]:
        raw_text = str(paragraph.text or "").strip()
        if not raw_text:
            continue
        if _is_bullet_paragraph(paragraph):
            continue

        if "\t" in paragraph.text:
            parsed = _split_source_project_title(paragraph.text)
            key = _normalise_source_project_key(parsed.get("title"))
            if not key:
                current_key = ""
                continue
            current_key = key
            result[current_key] = {
                "title": str(parsed.get("title") or "").strip(),
                "subtitle": str(parsed.get("subtitle") or "").strip(),
                "resume_header_tools": list(
                    parsed.get("resume_header_tools") or []
                ),
                "resume_header_context": list(
                    parsed.get("resume_header_context") or []
                ),
                "metadata_seen": False,
            }
            continue

        if not current_key or current_key not in result:
            continue
        current = result[current_key]
        if current.get("metadata_seen"):
            continue

        groups = [
            " ".join(part.replace("\u00a0", " ").split()).strip()
            for part in raw_text.split("|")
        ]
        groups = [group for group in groups if group]
        if not groups:
            continue

        tools = [
            " ".join(part.split()).strip()
            for part in groups[0].split(",")
            if " ".join(part.split()).strip()
        ]
        if tools:
            current["resume_header_tools"] = tools
        if len(groups) > 1:
            current["resume_header_context"] = groups[1:]
        current["metadata_seen"] = True

    for value in result.values():
        value.pop("metadata_seen", None)
    return result


def apply_source_project_display_fallbacks(
    saved_resume_docx_path: str | Path,
    tailored_projects: dict[str, Any],
) -> dict[str, Any]:
    """Fill missing display-only fields from the exact source resume.

    Explicit structured fields already present in the tailored project always
    win. The source fallback preserves Base Resume presentation without
    mutating canonical Evidence Library facts.
    """
    enriched = deepcopy(tailored_projects)
    try:
        source_document = Document(str(saved_resume_docx_path))
    except Exception:
        return enriched

    source_by_key = _extract_source_project_display_metadata(source_document)
    for project in enriched.get("recommended_projects", []) or []:
        if not isinstance(project, dict):
            continue
        key = _normalise_source_project_key(
            project.get("title")
            or project.get("display_title")
        )
        source = source_by_key.get(key)
        if not source:
            continue

        if not str(project.get("subtitle") or "").strip():
            project["subtitle"] = str(source.get("subtitle") or "").strip()
        if not project.get("resume_header_tools"):
            project["resume_header_tools"] = list(
                source.get("resume_header_tools") or []
            )
        if not project.get("resume_header_context"):
            project["resume_header_context"] = list(
                source.get("resume_header_context") or []
            )
        project["display_title"] = build_project_title(project)

    return enriched


def _find_project_templates_in_section(
    document: DocumentObject,
) -> tuple[Paragraph | None, Paragraph | None, Paragraph | None]:
    """Find source title, stacked-metadata, and bullet paragraph templates."""
    start_index, end_index = _find_section_range(document, {"PROJECTS"})
    paragraphs = document.paragraphs
    section_end = end_index if end_index is not None else len(paragraphs)

    title_template: Paragraph | None = None
    metadata_template: Paragraph | None = None
    bullet_template: Paragraph | None = None

    for paragraph in paragraphs[start_index + 1 : section_end]:
        raw_text = str(paragraph.text or "").strip()
        if not raw_text:
            continue
        if _is_bullet_paragraph(paragraph):
            if bullet_template is None:
                bullet_template = paragraph
            if title_template is not None and metadata_template is not None:
                break
            continue

        if title_template is None:
            title_template = paragraph
            continue

        if metadata_template is None and "\t" not in paragraph.text:
            metadata_template = paragraph

        if (
            title_template is not None
            and metadata_template is not None
            and bullet_template is not None
        ):
            break

    return title_template, metadata_template, bullet_template


# ---------------------------------------------------------------------------
# Paragraph insertion and formatting helpers
# ---------------------------------------------------------------------------

def _delete_paragraph(paragraph: Paragraph) -> None:
    """Remove a paragraph from the DOCX XML tree."""
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def _insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    """Insert a paragraph after another paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)

    if style:
        try:
            new_paragraph.style = style
        except Exception:
            pass

    if text:
        new_paragraph.add_run(text)

    return new_paragraph


def _copy_paragraph_format(source: Paragraph | None, target: Paragraph) -> None:
    """Copy basic paragraph styling/formatting from source to target."""
    if source is None:
        return

    try:
        target.style = source.style
    except Exception:
        pass

    source_format = source.paragraph_format
    target_format = target.paragraph_format

    for attr in (
        "left_indent",
        "right_indent",
        "first_line_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "alignment",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
    ):
        try:
            setattr(target_format, attr, getattr(source_format, attr))
        except Exception:
            pass


def _copy_run_format(source_run: Any, target_run: Any) -> None:
    """Copy basic run/font formatting from source run to target run."""
    if source_run is None:
        return

    try:
        target_run.style = source_run.style
    except Exception:
        pass

    for attr in ("bold", "italic", "underline"):
        try:
            setattr(target_run, attr, getattr(source_run, attr))
        except Exception:
            pass

    try:
        target_run.font.name = source_run.font.name
    except Exception:
        pass

    try:
        target_run.font.size = source_run.font.size
    except Exception:
        pass

    try:
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
    except Exception:
        pass


def _get_first_run_template(paragraph: Paragraph | None) -> Any:
    """Return first run from a paragraph, or None."""
    if paragraph is not None and paragraph.runs:
        return paragraph.runs[0]
    return None


def _copy_numbering(source: Paragraph | None, target: Paragraph) -> None:
    """
    Copy bullet/numbering XML from source paragraph to target paragraph.

    This helps preserve bullet indentation and bullet style.
    """
    if source is None:
        return

    source_ppr = source._p.pPr
    if source_ppr is None or source_ppr.numPr is None:
        return

    target_ppr = target._p.get_or_add_pPr()

    # Remove existing numbering if present.
    if target_ppr.numPr is not None:
        try:
            target_ppr.remove(target_ppr.numPr)
        except Exception:
            pass

    target_ppr.append(deepcopy(source_ppr.numPr))


def _clear_section_content(
    document: DocumentObject,
    heading_names: set[str],
) -> Paragraph:
    """
    Delete all paragraphs inside a section, keeping the section heading.

    Returns:
        The heading paragraph to use as insertion anchor.
    """
    start_index, end_index = _find_section_range(document, heading_names)
    paragraphs = document.paragraphs

    delete_start = start_index + 1
    delete_end = end_index if end_index is not None else len(paragraphs)

    for index in range(delete_end - 1, delete_start - 1, -1):
        _delete_paragraph(paragraphs[index])

    return document.paragraphs[start_index]


def _add_blank_line_after(anchor: Paragraph) -> Paragraph:
    """
    Insert a blank paragraph after anchor.
    This behaves like pressing Enter once in Word.
    """
    spacer = _insert_paragraph_after(anchor)

    try:
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = 1
    except Exception:
        pass

    return spacer


_MARGIN_PROFILE_TARGET_INCHES: dict[str, float | None] = {
    "source": None,
    "compact_075": 0.75,
    "compact_065": 0.65,
    "compact_050": 0.50,
}


def _normalise_margin_profile(value: str) -> str:
    profile = str(value or "source").strip().lower()
    return profile if profile in _MARGIN_PROFILE_TARGET_INCHES else "source"


def _apply_margin_profile(document: DocumentObject, profile: str) -> bool:
    # Shrink section margins conservatively; never expand them.
    profile = _normalise_margin_profile(profile)
    target_inches = _MARGIN_PROFILE_TARGET_INCHES[profile]
    if target_inches is None:
        return False
    target = Inches(float(target_inches))
    changed = False
    for section in document.sections:
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            current = getattr(section, attr, None)
            if current is not None and int(current) > int(target):
                setattr(section, attr, target)
                changed = True
    return changed


def _available_margin_compaction_profiles(saved_resume_docx_path: str | Path) -> list[str]:
    # Return only profiles that would actually shrink at least one margin.
    try:
        document = Document(str(saved_resume_docx_path))
    except Exception:
        return []
    profiles: list[str] = []
    for profile in ("compact_075", "compact_065", "compact_050"):
        target_inches = _MARGIN_PROFILE_TARGET_INCHES[profile]
        if target_inches is None:
            continue
        target = Inches(float(target_inches))
        if any(
            int(getattr(section, attr)) > int(target)
            for section in document.sections
            for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin")
            if getattr(section, attr, None) is not None
        ):
            profiles.append(profile)
    return profiles


# ---------------------------------------------------------------------------
# One page helpers
# ---------------------------------------------------------------------------
def count_pdf_pages(pdf_path: str | Path) -> int:
    """Count pages in a generated PDF."""
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    return len(reader.pages)



def measure_pdf_page_fill(
    pdf_path: str | Path,
) -> dict[str, Any]:
    """
    Measure vertical text occupancy for one-page and multi-page PDFs.

    The one-page fields preserve the existing Phase 5 diagnostics. For
    multi-page candidates, ``last_page_fill_ratio`` and ``overflow_ratio``
    allow the fitter to detect real layout progress before the document has
    reached one page.
    """
    try:
        import fitz  # PyMuPDF
    except Exception:
        return {
            "page_fill_ratio": None,
            "estimated_unused_page_ratio": None,
            "last_page_fill_ratio": None,
            "overflow_ratio": None,
            "occupied_page_units": None,
            "measurement_method": "unavailable",
        }

    try:
        document = fitz.open(str(pdf_path))
        page_count = len(document)
        page_fill_ratios: list[float] = []

        for page in document:
            page_height = float(page.rect.height)
            text_blocks = [
                block
                for block in page.get_text("blocks")
                if len(block) >= 5 and str(block[4]).strip()
            ]

            if not text_blocks or page_height <= 0:
                page_fill_ratios.append(0.0)
                continue

            lowest_text_y = max(float(block[3]) for block in text_blocks)
            page_fill_ratios.append(
                max(0.0, min(1.0, lowest_text_y / page_height))
            )

        document.close()

        if not page_fill_ratios:
            return {
                "page_fill_ratio": 0.0,
                "estimated_unused_page_ratio": 1.0,
                "last_page_fill_ratio": 0.0,
                "overflow_ratio": 0.0,
                "occupied_page_units": 0.0,
                "measurement_method": "pymupdf_text_blocks",
            }

        last_page_fill_ratio = page_fill_ratios[-1]
        single_page_fill_ratio = (
            page_fill_ratios[0] if page_count == 1 else None
        )
        overflow_ratio = (
            0.0
            if page_count <= 1
            else max(0.0, float(page_count - 2) + last_page_fill_ratio)
        )

        return {
            "page_fill_ratio": (
                round(single_page_fill_ratio, 3)
                if single_page_fill_ratio is not None
                else None
            ),
            "estimated_unused_page_ratio": (
                round(1.0 - single_page_fill_ratio, 3)
                if single_page_fill_ratio is not None
                else None
            ),
            "last_page_fill_ratio": round(last_page_fill_ratio, 3),
            "overflow_ratio": round(overflow_ratio, 3),
            "occupied_page_units": round(sum(page_fill_ratios), 3),
            "measurement_method": "pymupdf_text_blocks",
        }

    except Exception:
        return {
            "page_fill_ratio": None,
            "estimated_unused_page_ratio": None,
            "last_page_fill_ratio": None,
            "overflow_ratio": None,
            "occupied_page_units": None,
            "measurement_method": "failed",
        }



def _delete_generated_output(
    docx_path: str | Path | None,
    pdf_path: str | Path | None,
) -> None:
    """Delete an unselected temporary DOCX/PDF candidate."""
    for raw_path in (
        docx_path,
        pdf_path,
    ):
        if not raw_path:
            continue

        try:
            Path(raw_path).unlink()
        except FileNotFoundError:
            pass
        except OSError:
            pass

# def _project_priority_score(project: dict[str, Any]) -> int:
#     """Rank projects for space reduction."""
#     priority = str(project.get("priority", "")).lower()

#     if priority == "high":
#         base = 300
#     elif priority == "medium":
#         base = 200
#     elif priority == "low":
#         base = 100
#     else:
#         base = 150

#     matched_count = len(project.get("matched_jd_requirements", []) or [])

#     return base + matched_count

def _project_priority_score(
    project: dict[str, Any],
) -> int:
    """Rank projects for space reduction."""
    priority = str(
        project.get("priority", "")
    ).lower()

    if priority == "high":
        base = 300
    elif priority == "medium":
        base = 200
    elif priority == "low":
        base = 100
    else:
        base = 150

    fit_score = int(
        project.get("project_fit_score", 0)
        or 0
    )

    direct_matches = len(
        project.get("matched_jd_requirements", [])
        or []
    )

    transferable_matches = len(
        project.get(
            "transferable_jd_requirements",
            [],
        )
        or []
    )

    return (
        base
        + fit_score
        + direct_matches * 10
        + transferable_matches * 5
    )

# _WEAK_TRAILING_WORDS = {
#     "and",
#     "or",
#     "but",
#     "for",
#     "to",
#     "of",
#     "in",
#     "on",
#     "with",
#     "by",
#     "as",
#     "the",
#     "a",
#     "an",
#     "that",
#     "which",
#     "including",
# }


# def _trim_words(text: str, max_words: int) -> str:
#     """
#     Trim a bullet to a rough word limit without ending on a weak
#     connector such as 'and', 'to', or 'with'.
#     """
#     cleaned_text = " ".join(str(text).split())
#     words = cleaned_text.split()

#     if len(words) <= max_words:
#         return cleaned_text

#     trimmed_words = words[:max_words]

#     # Avoid results such as:
#     # "... degree fit, and."
#     while trimmed_words:
#         final_word = re.sub(
#             r"[^a-z0-9]+",
#             "",
#             trimmed_words[-1].lower(),
#         )

#         if final_word not in _WEAK_TRAILING_WORDS:
#             break

#         trimmed_words.pop()

#     if not trimmed_words:
#         return cleaned_text

#     trimmed_text = " ".join(trimmed_words).rstrip(" ,;:.-")
#     return trimmed_text + "."

# def _trim_words(text: str, max_words: int) -> str:
#     """Trim a bullet to a rough word limit."""
#     words = str(text).split()

#     if len(words) <= max_words:
#         return str(text).strip()

#     trimmed = " ".join(words[:max_words]).rstrip(".,;:")
#     return trimmed + "."


# def compact_tailored_projects_for_space(
#     tailored_projects: dict[str, Any],
#     *,
#     attempt: int,
# ) -> dict[str, Any]:
#     """
#     Reduce Projects section size only after the generated DOCX exceeds one page.

#     attempt 1:
#         Keep projects and bullet counts mostly intact, only trim long bullets.
#     attempt 2:
#         Keep 3 projects. High priority gets up to 2 bullets, others up to 1-2.
#     attempt 3:
#         Keep 3 projects. High priority gets up to 2 bullets;
#         medium and low priority get 1 bullet.

#     attempt 4+:
#         Keep 2 strongest projects with 1 concise bullet each.
#     """
#     compacted = deepcopy(tailored_projects)
#     projects = compacted.get("recommended_projects", [])

#     projects = sorted(
#         projects,
#         key=_project_priority_score,
#         reverse=True,
#     )

#     if attempt == 1:
#         max_projects = 3
#         max_words = 24
#         bullet_limits_by_priority = {
#             "high": 3,
#             "medium": 2,
#             "low": 1,
#         }
#         note = "Trimmed long project bullets but kept useful evidence where possible."

#     elif attempt == 2:
#         max_projects = 3
#         max_words = 24
#         bullet_limits_by_priority = {
#             "high": 2,
#             "medium": 2,
#             "low": 1,
#         }
#         note = "Reduced lower-priority project detail because the resume exceeded one page."

#     elif attempt == 3:
#         max_projects = 3
#         max_words = 20
#         bullet_limits_by_priority = {
#             "high": 2,
#             "medium": 1,
#             "low": 1,
#         }
#         note = (
#         "Kept up to two bullets for high-priority projects and one bullet "
#         "for medium- and low-priority projects."
#         )

#     else:
#         max_projects = 2
#         max_words = 18
#         bullet_limits_by_priority = {
#             "high": 1,
#             "medium": 1,
#             "low": 1,
#         }
#         note = "Kept only the two strongest projects because the resume still exceeded one page."

#     kept_projects = []

#     for project in projects[:max_projects]:
#         priority = str(project.get("priority", "")).lower()
#         bullet_limit = bullet_limits_by_priority.get(priority, 1)

#         original_bullets = project.get("draft_bullets", []) or []


#         compacted_bullets = [
#             _trim_words(bullet, max_words)
#             for bullet in original_bullets[:bullet_limit]
#             if str(bullet).strip()
#         ]

#         project["draft_bullets"] = compacted_bullets

#         if len(compacted_bullets) == 1:
#             project["space_action"] = "single_bullet"
#         elif compacted_bullets != original_bullets:
#             project["space_action"] = "shorten"
#         else:
#             project["space_action"] = "keep_full"

#         # bullets = project.get("draft_bullets", []) or []
#         # project["draft_bullets"] = [
#         #     _trim_words(bullet, max_words)
#         #     for bullet in bullets[:bullet_limit]
#         #     if str(bullet).strip()
#         # ]

#         # if len(project["draft_bullets"]) == 1:
#         #     project["space_action"] = "single_bullet"
#         # elif attempt >= 1:
#         #     project["space_action"] = "shorten"

#         kept_projects.append(project)



#     removed_projects = projects[max_projects:]

#     kept_projects = sorted(
#     kept_projects,
#     key=lambda project: period_sort_value(project.get("period", "")),
#     reverse=True,
#     )

#     compacted["recommended_projects"] = kept_projects
#     compacted["projects_to_remove_or_deprioritize"] = compacted.get(
#         "projects_to_remove_or_deprioritize",
#         [],
#     )

#     for project in removed_projects:
#         compacted["projects_to_remove_or_deprioritize"].append(
#             {
#                 "title": project.get("display_title") or project.get("title", "Untitled Project"),
#                 "reason": "Removed during compacting because the resume exceeded one page.",
#             }
#         )

#     compacted.setdefault("notes_for_user", []).append(
#         f"Applied compact mode attempt {attempt}: {note}"
#     )

#     return compacted

def _bullet_word_count(values: list[str]) -> int:
    """Return the total word count for a bullet list."""
    return sum(
        len(str(value).split())
        for value in values
    )


_COMPACT_ACTION_VERBS = {
    "applied",
    "automated",
    "built",
    "collaborated",
    "configured",
    "containerised",
    "containerized",
    "contributed",
    "created",
    "deployed",
    "designed",
    "developed",
    "engineered",
    "implemented",
    "integrated",
    "led",
    "managed",
    "optimised",
    "optimized",
    "scripted",
    "secured",
    "tested",
    "validated",
}


def _compact_bullets_are_quality_preserving(
    full_bullets: list[str],
    compact_bullets: list[str],
) -> bool:
    """
    Return True only for conservative, one-to-one compact rewrites.

    Complete-bullet removal is handled later by the fitting loop. The compact
    pass should save lines without silently deleting evidence or producing
    weak sentence fragments.
    """
    if not full_bullets or not compact_bullets:
        return False

    if len(compact_bullets) != len(full_bullets):
        return False

    full_word_count = _bullet_word_count(
        full_bullets
    )
    compact_word_count = _bullet_word_count(
        compact_bullets
    )

    if full_word_count <= 0:
        return False

    if compact_word_count >= full_word_count:
        return False

    if compact_word_count / full_word_count < 0.65:
        return False

    for bullet in compact_bullets:
        cleaned = str(bullet or "").strip()
        words = cleaned.split()

        if not 10 <= len(words) <= 24:
            return False

        first_word = re.sub(
            r"[^a-z]+",
            "",
            words[0].lower(),
        )

        if first_word not in _COMPACT_ACTION_VERBS:
            return False

        if cleaned[-1:] not in {".", ";", ":"}:
            return False

    return True


def _protected_compact_bullet_indexes(
    project: dict[str, Any],
    *,
    bullet_count: int,
) -> list[int]:
    """Return Phase 6B.1-protected bullet indexes without mutating state."""
    protected: set[int] = set()

    for raw_index in project.get("protected_bullet_indexes", []) or []:
        try:
            index = int(raw_index)
        except (TypeError, ValueError):
            continue
        if 0 <= index < bullet_count:
            protected.add(index)

    for row in project.get("bullet_evidence_priorities", []) or []:
        if not isinstance(row, dict):
            continue
        try:
            index = int(row.get("bullet_index"))
        except (TypeError, ValueError):
            continue
        if not 0 <= index < bullet_count:
            continue
        if (
            row.get("protect_during_fitting")
            or row.get("protected_requirement_ids")
            or int(row.get("unique_required_core_count", 0) or 0) > 0
        ):
            protected.add(index)

    return sorted(protected)


def _build_evidence_preserving_compact_bullets(
    project: dict[str, Any],
    *,
    full_bullets: list[str],
    compact_bullets: list[str],
) -> tuple[list[str], list[int], list[int]] | None:
    """Compact only unprotected indexes; protected text stays unchanged."""
    if not _compact_bullets_are_quality_preserving(
        full_bullets,
        compact_bullets,
    ):
        return None

    protected_indexes = _protected_compact_bullet_indexes(
        project,
        bullet_count=len(full_bullets),
    )
    protected_set = set(protected_indexes)

    mixed_bullets = [
        full_bullets[index]
        if index in protected_set
        else compact_bullets[index]
        for index in range(len(full_bullets))
    ]
    compacted_indexes = [
        index
        for index in range(len(full_bullets))
        if (
            index not in protected_set
            and compact_bullets[index] != full_bullets[index]
        )
    ]

    if not compacted_indexes:
        return None
    if _bullet_word_count(mixed_bullets) >= _bullet_word_count(full_bullets):
        return None

    return mixed_bullets, protected_indexes, compacted_indexes


def _count_quality_compact_candidates(
    tailored_projects: dict[str, Any],
) -> int:
    """Count projects with an unused quality-preserving compact alternative."""
    count = 0

    for project in (
        tailored_projects.get(
            "recommended_projects",
            [],
        )
        or []
    ):
        full_bullets = [
            str(value).strip()
            for value in (
                project.get(
                    "draft_bullets",
                    [],
                )
                or []
            )
            if str(value).strip()
        ]

        compact_bullets = [
            str(value).strip()
            for value in (
                project.get(
                    "compact_bullets",
                    [],
                )
                or []
            )
            if str(value).strip()
        ]

        if _build_evidence_preserving_compact_bullets(
            project,
            full_bullets=full_bullets,
            compact_bullets=compact_bullets,
        ) is not None:
            count += 1

    return count


def apply_compact_bullets_once(
    tailored_projects: dict[str, Any],
) -> tuple[dict[str, Any], bool, dict[str, Any]]:
    """
    Apply one quality-preserving compact project rewrite.

    Only one project is compacted per render attempt. This prevents every
    project from being shortened when a smaller change may already make the
    resume fit. Lower-relevance projects are compacted first.
    """
    compacted = deepcopy(
        tailored_projects
    )

    projects = (
        compacted.get(
            "recommended_projects",
            [],
        )
        or []
    )

    eligible_projects: list[
        tuple[
            dict[str, Any],
            list[str],
            list[str],
            list[str],
            list[int],
            list[int],
        ]
    ] = []

    for project in projects:
        if (
            str(
                project.get(
                    "space_action",
                    "",
                )
            )
            == "compact_rewrite"
        ):
            continue

        full_bullets = [
            str(value).strip()
            for value in (
                project.get(
                    "draft_bullets",
                    [],
                )
                or []
            )
            if str(value).strip()
        ]

        compact_bullets = [
            str(value).strip()
            for value in (
                project.get(
                    "compact_bullets",
                    [],
                )
                or []
            )
            if str(value).strip()
        ]

        mixed = _build_evidence_preserving_compact_bullets(
            project,
            full_bullets=full_bullets,
            compact_bullets=compact_bullets,
        )
        if mixed is not None:
            mixed_bullets, protected_indexes, compacted_indexes = mixed
            eligible_projects.append(
                (
                    project,
                    full_bullets,
                    compact_bullets,
                    mixed_bullets,
                    protected_indexes,
                    compacted_indexes,
                )
            )

    if not eligible_projects:
        return (
            compacted,
            False,
            {
                "change_type": (
                    "compact_rewrite_unavailable"
                ),
                "reason": (
                    "No unused quality-preserving compact "
                    "project bullets were available."
                ),
            },
        )

    (
        target_project,
        full_bullets,
        source_compact_bullets,
        mixed_bullets,
        protected_indexes,
        compacted_indexes,
    ) = min(
        eligible_projects,
        key=lambda item: (
            _project_priority_score(
                item[0]
            ),
            -(
                _bullet_word_count(
                    item[1]
                )
                - _bullet_word_count(
                    item[3]
                )
            ),
        ),
    )

    previous_space_action = str(
        target_project.get(
            "space_action",
            "keep_full",
        )
    )

    target_project[
        "draft_bullets"
    ] = list(mixed_bullets)

    target_project[
        "space_action"
    ] = "compact_rewrite"

    # Protected bullets retain their exact full text. Unprotected bullets may
    # use their one-to-one compact alternative; metadata remains aligned by index.
    sync_project_bullet_metadata(
        target_project,
        bullet_texts=mixed_bullets,
    )

    project_title = (
        target_project.get(
            "display_title"
        )
        or target_project.get(
            "title"
        )
        or "Untitled Project"
    )

    change = {
        "change_type": "compact_rewrite",
        "project": project_title,
        "previous_space_action": (
            previous_space_action
        ),
        "original_draft_bullets": list(
            full_bullets
        ),
        "source_compact_bullets": list(
            source_compact_bullets
        ),
        "compact_draft_bullets": list(
            mixed_bullets
        ),
        "protected_bullet_indexes_preserved": list(
            protected_indexes
        ),
        "compacted_bullet_indexes": list(
            compacted_indexes
        ),
        "protected_bullet_text_preserved": True,
        "project_priority_score": (
            _project_priority_score(
                target_project
            )
        ),
        "full_bullet_count": len(
            full_bullets
        ),
        "compact_bullet_count": len(
            compact_bullets
        ),
        "full_word_count": (
            _bullet_word_count(
                full_bullets
            )
        ),
        "compact_word_count": (
            _bullet_word_count(
                mixed_bullets
            )
        ),
    }

    compacted[
        "recommended_projects"
    ] = sorted(
        projects,
        key=lambda project: (
            period_sort_value(
                project.get(
                    "period",
                    "",
                )
            )
        ),
        reverse=True,
    )

    compacted.setdefault(
        "notes_for_user",
        [],
    ).append(
        "Applied a quality-preserving compact rewrite "
        f"to {project_title} before deleting any "
        "complete project bullet."
    )

    return (
        compacted,
        True,
        change,
    )

def compact_tailored_projects_one_step(
    tailored_projects: dict[str, Any],
    *,
    minimum_bullets_per_project: int = 1,
    minimum_projects_to_keep: int = 3,
    prefer_balanced_bullets: bool = False,
) -> tuple[dict[str, Any], bool, dict[str, Any]]:
    """
    Return the safest deterministic Phase 6C reduction.

    The full one-page fitting loop calls
    ``build_evidence_aware_project_reductions`` directly so it can render and
    compare every eligible bullet candidate. This wrapper remains compatible
    with callers and tests that expect one reduction.
    """
    candidates = build_evidence_aware_project_reductions(
        tailored_projects,
        minimum_bullets_per_project=minimum_bullets_per_project,
        minimum_projects_to_keep=minimum_projects_to_keep,
        prefer_balanced_bullets=prefer_balanced_bullets,
    )

    if not candidates:
        return (
            deepcopy(tailored_projects),
            False,
            {
                "fitting_version": PHASE6C_FITTING_VERSION,
                "change_type": "none",
                "reason": "No evidence-aware project reduction is available.",
            },
        )

    compacted, change = candidates[0]
    return compacted, True, change

def _project_change_title(
    project: dict[str, Any],
) -> str:
    """Return the stable display title used by fitting changes."""
    return str(
        project.get("display_title")
        or project.get("title")
        or "Untitled Project"
    )


def _find_project_for_change(
    tailored_projects: dict[str, Any],
    project_title: str,
) -> dict[str, Any] | None:
    """Find a retained project by the title recorded in a change."""
    for project in (
        tailored_projects.get(
            "recommended_projects",
            [],
        )
        or []
    ):
        if (
            _project_change_title(
                project
            )
            == project_title
        ):
            return project

    return None


def _restore_fitting_change(
    tailored_projects: dict[str, Any],
    change: dict[str, Any],
) -> tuple[
    dict[str, Any],
    bool,
    dict[str, Any],
]:
    """
    Reverse one earlier fitting reduction without changing other projects.
    """
    restored = deepcopy(
        tailored_projects
    )

    change_type = str(
        change.get(
            "change_type",
            "",
        )
    )

    project_title = str(
        change.get(
            "project",
            "",
        )
    )

    if change_type == "compact_rewrite":
        project = _find_project_for_change(
            restored,
            project_title,
        )

        original_bullets = [
            str(value).strip()
            for value in (
                change.get(
                    "original_draft_bullets",
                    [],
                )
                or []
            )
            if str(value).strip()
        ]

        if (
            project is None
            or not original_bullets
        ):
            return (
                restored,
                False,
                {
                    "change_type": (
                        "restore_unavailable"
                    ),
                    "reason": (
                        "The compact rewrite could not "
                        "be restored."
                    ),
                },
            )

        project[
            "draft_bullets"
        ] = original_bullets

        # Compacting updates metadata bullet_text values to the mixed rendered
        # wording. Restore those rows to the full text as well.
        sync_project_bullet_metadata(
            project,
            bullet_texts=original_bullets,
        )

        project[
            "space_action"
        ] = str(
            change.get(
                "previous_space_action",
                "keep_full",
            )
        )

        restore_info = {
            "change_type": (
                "restore_compact_rewrite"
            ),
            "project": project_title,
            "restored_change_type": (
                change_type
            ),
            "restored_word_count": (
                _bullet_word_count(
                    original_bullets
                )
            ),
        }

    elif change_type == "remove_bullet":
        project = _find_project_for_change(
            restored,
            project_title,
        )

        removed_bullet = str(
            change.get(
                "removed_bullet",
                "",
            )
        ).strip()

        if project is None or not removed_bullet:
            return (
                restored,
                False,
                {
                    "change_type": (
                        "restore_unavailable"
                    ),
                    "reason": (
                        "The removed bullet could not "
                        "be restored."
                    ),
                },
            )

        bullets = [
            str(value).strip()
            for value in (
                project.get(
                    "draft_bullets",
                    [],
                )
                or []
            )
            if str(value).strip()
        ]

        insert_index = int(
            change.get(
                "removed_bullet_index",
                len(bullets),
            )
            or 0
        )

        insert_index = max(
            0,
            min(
                insert_index,
                len(bullets),
            ),
        )

        bullets.insert(
            insert_index,
            removed_bullet,
        )

        project[
            "draft_bullets"
        ] = bullets

        restore_removed_bullet_metadata(
            project,
            bullet_index=insert_index,
            bullet_text=removed_bullet,
            removed_metadata=change.get("removed_bullet_metadata"),
        )

        project[
            "space_action"
        ] = str(
            change.get(
                "previous_space_action",
                (
                    "single_bullet"
                    if len(bullets) == 1
                    else "keep_full"
                ),
            )
        )

        restore_info = {
            "change_type": (
                "restore_removed_bullet"
            ),
            "project": project_title,
            "restored_change_type": (
                change_type
            ),
            "restored_bullet": (
                removed_bullet
            ),
        }

    elif change_type == "remove_project":
        removed_project = change.get(
            "removed_project_data"
        )

        if not isinstance(
            removed_project,
            dict,
        ):
            return (
                restored,
                False,
                {
                    "change_type": (
                        "restore_unavailable"
                    ),
                    "reason": (
                        "The removed project data is "
                        "not available."
                    ),
                },
            )

        projects = (
            restored.get(
                "recommended_projects",
                [],
            )
            or []
        )

        if any(
            _project_change_title(project)
            == project_title
            for project in projects
        ):
            return (
                restored,
                False,
                {
                    "change_type": (
                        "restore_unavailable"
                    ),
                    "reason": (
                        "The project is already present."
                    ),
                },
            )

        projects.append(
            deepcopy(
                removed_project
            )
        )

        restored[
            "recommended_projects"
        ] = sorted(
            projects,
            key=lambda project: (
                period_sort_value(
                    project.get(
                        "period",
                        "",
                    )
                )
            ),
            reverse=True,
        )

        restore_info = {
            "change_type": (
                "restore_removed_project"
            ),
            "project": project_title,
            "restored_change_type": (
                change_type
            ),
        }

    else:
        return (
            restored,
            False,
            {
                "change_type": (
                    "restore_unavailable"
                ),
                "reason": (
                    "Unsupported fitting change."
                ),
            },
        )

    restored.setdefault(
        "notes_for_user",
        [],
    ).append(
        "Restored stronger project content after "
        "confirming the resume still fits on one page."
    )

    return (
        restored,
        True,
        restore_info,
    )


def _restoration_quality_gain(
    change: dict[str, Any],
) -> int:
    """Score evidence recovered by reversing a fitting change."""
    project_priority = int(change.get("project_priority_score", 0) or 0)
    change_type = str(change.get("change_type", ""))

    if change_type == "compact_rewrite":
        recovered_words = max(
            0,
            int(change.get("full_word_count", 0) or 0)
            - int(change.get("compact_word_count", 0) or 0),
        )
        return project_priority * 10 + recovered_words

    if change_type == "remove_bullet":
        recovered_words = len(
            str(change.get("removed_bullet", "")).split()
        )
        legacy_gain = project_priority * 10 + recovered_words * 3
        phase6c_gain = (
            int(change.get("evidence_loss_score", 0) or 0) * 10
            + recovered_words
        )
        return max(legacy_gain, phase6c_gain)

    if change_type == "remove_project":
        removed_project = change.get("removed_project_data") or {}
        recovered_words = (
            _bullet_word_count(
                removed_project.get("draft_bullets", []) or []
            )
            if isinstance(removed_project, dict)
            else 0
        )
        legacy_gain = 10000 + project_priority * 10 + recovered_words
        phase6c_gain = (
            int(change.get("evidence_loss_score", 0) or 0) * 10
            + recovered_words
        )
        return max(legacy_gain, phase6c_gain)

    return 0

def _restorable_change_indices(
    active_changes: list[
        dict[str, Any]
    ],
) -> list[int]:
    """
    Return changes that can be reversed without conflicting with a later
    change to the same project.
    """
    restorable: list[int] = []

    for index, change in enumerate(
        active_changes
    ):
        project_title = str(
            change.get(
                "project",
                "",
            )
        )

        later_same_project = any(
            str(
                later_change.get(
                    "project",
                    "",
                )
            )
            == project_title
            for later_change in active_changes[
                index + 1:
            ]
        )

        if not later_same_project:
            restorable.append(index)

    return restorable


# ---------------------------------------------------------------------------
# Skills replacement
# ---------------------------------------------------------------------------

def _add_skill_line_after(
    anchor: Paragraph,
    *,
    category: str,
    items: list[str],
    template: Paragraph | None = None,
) -> Paragraph:
    """
    Add one skill line after anchor.

    If the original Skills section used bullets, this preserves that bullet style.
    """
    new_paragraph = _insert_paragraph_after(anchor)

    if template is not None:
        _copy_paragraph_format(template, new_paragraph)
        _copy_numbering(template, new_paragraph)
    else:
        try:
            new_paragraph.style = "List Bullet"
        except Exception:
            pass

    source_run = _get_first_run_template(template)

    if category:
        category_run = new_paragraph.add_run(f"{category}: ")
        _copy_run_format(source_run, category_run)
        category_run.bold = True

        items_run = new_paragraph.add_run(", ".join(items))
        _copy_run_format(source_run, items_run)
        items_run.bold = False
    else:
        line_run = new_paragraph.add_run(", ".join(items))
        _copy_run_format(source_run, line_run)

    return new_paragraph


def replace_skills_section(
    document: DocumentObject,
    tailored_skills: dict[str, Any],
) -> None:
    """
    Replace SKILLS / TECHNICAL SKILLS section content while preserving formatting.
    """
    _, skill_bullet_template = _find_templates_in_section(document, {"SKILLS", "TECHNICAL SKILLS"})

    anchor = _clear_section_content(document, {"SKILLS", "TECHNICAL SKILLS"})
    skill_lines = tailored_skills.get("skill_lines", [])

    if not skill_lines:
        _insert_paragraph_after(anchor, "No tailored skills were generated.")
        return

    for row in skill_lines:
        category = str(row.get("category", "")).strip()
        items = [str(item).strip() for item in row.get("items", []) if str(item).strip()]

        if items:
            anchor = _add_skill_line_after(
                anchor,
                category=category,
                items=items,
                template=skill_bullet_template,
            )


# ---------------------------------------------------------------------------
# Projects replacement
# ---------------------------------------------------------------------------

# def _format_project_heading(project: dict[str, Any]) -> str:
#     """
#     Build a project heading from optional project fields.

#     Supported fields:
#         title: "QueryAI"
#         role: "AI Programmer"
#         display_details: "React, Team of 4"

#     Output examples:
#         QueryAI (React, Team of 4)
#         Job AI Helper – AI Programmer (Python, Streamlit, Solo)

#     If your title already contains bracket details, you can leave role/details empty.
#     """
#     title = str(project.get("title", "Untitled Project")).strip() or "Untitled Project"
#     role = str(project.get("role", "")).strip()
#     details = str(project.get("display_details", "")).strip()

#     heading = title

#     if role and role.lower() not in heading.lower():
#         heading += f" – {role}"

#     if details and details.lower() not in heading.lower():
#         heading += f" ({details})"

#     return heading

def _format_project_heading(project: dict[str, Any]) -> str:
    """Return semantic project title/subtitle only; metadata renders separately."""
    return build_project_title(project)

def _add_project_title_after(
    anchor: Paragraph,
    *,
    title: str,
    period: str = "",
    inline_metadata: str = "",
    metadata_style: str = "pipes",
    template: Paragraph | None = None,
    right_tab_position: Any = None,
    add_space_before: bool = False,
    space_before_pt: int = 10,
) -> Paragraph:
    """
    Add project title line after anchor, preserving original formatting.

    Uses a right-aligned tab stop for the period/date.
    """
    new_paragraph = _insert_paragraph_after(anchor)

    if template is not None:
        _copy_paragraph_format(template, new_paragraph)
    
    if add_space_before:
        new_paragraph.paragraph_format.space_before = Pt(space_before_pt)

    # Keep title and following bullet together where Word supports it.
    try:
        new_paragraph.paragraph_format.keep_with_next = True
    except Exception:
        pass

    # Right tab stop for date.
    if right_tab_position is not None:
        try:
            new_paragraph.paragraph_format.tab_stops.add_tab_stop(
                right_tab_position,
                WD_TAB_ALIGNMENT.RIGHT,
            )
        except Exception:
            pass

    source_run = _get_first_run_template(template)

    title_run = new_paragraph.add_run(title)
    _copy_run_format(source_run, title_run)
    title_run.bold = True

    if inline_metadata:
        separator = (
            " "
            if normalise_project_metadata_style(metadata_style) == "parentheses"
            else " | "
        )
        metadata_run = new_paragraph.add_run(separator + inline_metadata)
        _copy_run_format(source_run, metadata_run)
        metadata_run.bold = False
        metadata_run.italic = True

    if period:
        date_run = new_paragraph.add_run(f"\t{period}")
        _copy_run_format(source_run, date_run)
        date_run.bold = False

    return new_paragraph

# def _add_project_title_after(
#     anchor: Paragraph,
#     *,
#     title: str,
#     period: str = "",
#     template: Paragraph | None = None,
#     right_tab_position: Any = None,
#     add_space_before: bool = False,
# ) -> Paragraph:
#     """
#     Add project title line after anchor, preserving original formatting.

#     Uses a right-aligned tab stop for the period/date.
#     """
#     new_paragraph = _insert_paragraph_after(anchor)

#     if template is not None:
#         _copy_paragraph_format(template, new_paragraph)
    
#     if add_space_before:
#         new_paragraph.paragraph_format.space_before = Pt(10)

#     # Keep title and following bullet together where Word supports it.
#     try:
#         new_paragraph.paragraph_format.keep_with_next = True
#     except Exception:
#         pass

#     # Right tab stop for date.
#     if right_tab_position is not None:
#         try:
#             new_paragraph.paragraph_format.tab_stops.add_tab_stop(
#                 right_tab_position,
#                 WD_TAB_ALIGNMENT.RIGHT,
#             )
#         except Exception:
#             pass

#     source_run = _get_first_run_template(template)

#     title_run = new_paragraph.add_run(title)
#     _copy_run_format(source_run, title_run)
#     title_run.bold = True

#     if period:
#         date_run = new_paragraph.add_run(f"\t{period}")
#         _copy_run_format(source_run, date_run)
#         date_run.bold = False

#     return new_paragraph

# def _add_project_title_after(
#     anchor: Paragraph,
#     *,
#     title: str,
#     period: str = "",
#     template: Paragraph | None = None,
#     right_tab_position: Any = None,
# ) -> Paragraph:
#     """
#     Add project title line after anchor, preserving original formatting.

#     Uses a right-aligned tab stop for the period/date.
#     """
#     new_paragraph = _insert_paragraph_after(anchor)

#     if template is not None:
#         _copy_paragraph_format(template, new_paragraph)

#     # Keep title and following bullet together where Word supports it.
#     try:
#         new_paragraph.paragraph_format.keep_with_next = True
#     except Exception:
#         pass

#     # Right tab stop for date.
#     if right_tab_position is not None:
#         try:
#             new_paragraph.paragraph_format.tab_stops.add_tab_stop(
#                 right_tab_position,
#                 WD_TAB_ALIGNMENT.RIGHT,
#             )
#         except Exception:
#             pass

#     source_run = _get_first_run_template(template)

#     title_run = new_paragraph.add_run(title)
#     _copy_run_format(source_run, title_run)
#     title_run.bold = True

#     if period:
#         date_run = new_paragraph.add_run(f"\t{period}")
#         _copy_run_format(source_run, date_run)
#         date_run.bold = False

#     return new_paragraph



def _add_project_metadata_after(
    anchor: Paragraph,
    *,
    metadata: str,
    template: Paragraph | None = None,
) -> Paragraph:
    """Add stacked metadata using the source metadata-line presentation."""
    new_paragraph = _insert_paragraph_after(anchor)
    if template is not None:
        _copy_paragraph_format(template, new_paragraph)
    else:
        try:
            new_paragraph.paragraph_format.space_before = Pt(0)
            new_paragraph.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
    try:
        new_paragraph.paragraph_format.keep_with_next = True
    except Exception:
        pass
    source_run = _get_first_run_template(template)
    run = new_paragraph.add_run(metadata)
    _copy_run_format(source_run, run)
    if source_run is None:
        run.bold = False
    return new_paragraph

def _paragraph_has_native_numbering(
    paragraph: Paragraph | None,
) -> bool:
    # True when the paragraph carries explicit Word list numbering.
    if paragraph is None:
        return False
    ppr = paragraph._p.pPr
    return bool(ppr is not None and ppr.numPr is not None)


def _literal_bullet_prefix(
    paragraph: Paragraph | None,
) -> str:
    # Preserve literal Unicode bullets only for non-numbered templates.
    # Native Word numbered/list paragraphs deliberately return no text prefix
    # so they never receive a second bullet.
    if paragraph is None or _paragraph_has_native_numbering(paragraph):
        return ""

    text = str(paragraph.text or "")
    match = re.match(
        r"^\s*([•●▪◦‣∙])(?:[ \t\u00a0]+)?",
        text,
    )
    if match is None:
        return ""
    return f"{match.group(1)} "


def _add_project_bullet_after(
    anchor: Paragraph,
    *,
    bullet: str,
    template: Paragraph | None = None,
) -> Paragraph:
    # Add a project bullet while preserving the template's bullet mechanism.
    new_paragraph = _insert_paragraph_after(anchor)

    if template is not None:
        _copy_paragraph_format(template, new_paragraph)
        _copy_numbering(template, new_paragraph)
    else:
        try:
            new_paragraph.style = "List Bullet"
        except Exception:
            pass

    source_run = _get_first_run_template(template)

    bullet_text = str(bullet).strip()
    literal_prefix = _literal_bullet_prefix(template)

    # Preserve a literal-bullet template such as Resume (Tech) 3, but do not
    # duplicate a bullet if generated text already contains one.
    if (
        literal_prefix
        and bullet_text
        and not re.match(r"^[•●▪◦‣∙](?:\s|$)", bullet_text)
    ):
        bullet_text = literal_prefix + bullet_text

    run = new_paragraph.add_run(bullet_text)
    _copy_run_format(source_run, run)

    return new_paragraph


def replace_projects_section(
    document: DocumentObject,
    tailored_projects: dict[str, Any],
    *,
    max_projects: int = 3,
    max_bullets_per_project: int = 2,
    spacing_mode: str = "paragraph_spacing",
    project_spacing_pt: int = 10,
    after_projects_spacing_pt: int = 10,
    blank_lines_between_projects: int = 1,
    blank_lines_after_projects: int = 1,
    add_spacing_before_first_project: bool = False,
    project_header_layout: str = "auto",
    project_metadata_style: str = "pipes",
) -> None:
    """
    Replace PROJECTS section content while preserving original formatting.
    """
    (
        project_title_template,
        project_metadata_template,
        project_bullet_template,
    ) = _find_project_templates_in_section(document)

    section = document.sections[0]
    right_tab_position = section.page_width - section.left_margin - section.right_margin

    anchor = _clear_section_content(document, {"PROJECTS"})
    projects = tailored_projects.get("recommended_projects", [])[:max_projects]
    resolved_layout = normalise_project_header_layout(project_header_layout)
    if resolved_layout == "auto":
        resolved_layout = "stacked"
    resolved_metadata_style = normalise_project_metadata_style(project_metadata_style)

    if not projects:
        _insert_paragraph_after(anchor, "No tailored projects were generated.")
        return
    
    for project_index, project in enumerate(projects):
        title = _format_project_heading(project)
        metadata = format_project_metadata(project, style=resolved_metadata_style)
        period = str(project.get("period", "")).strip()
        bullets = project.get("draft_bullets", [])[:max_bullets_per_project]

        should_add_spacing_before = project_index > 0 or add_spacing_before_first_project

        if spacing_mode == "blank_line" and should_add_spacing_before:
            for _ in range(blank_lines_between_projects):
                anchor = _add_blank_line_after(anchor)

        anchor = _add_project_title_after(
            anchor,
            title=title,
            period=period,
            inline_metadata=(metadata if resolved_layout == "inline" else ""),
            metadata_style=resolved_metadata_style,
            template=project_title_template,
            right_tab_position=right_tab_position,
            add_space_before=(
                spacing_mode == "paragraph_spacing"
                and should_add_spacing_before
            ),
            space_before_pt=project_spacing_pt,
        )
        if metadata and resolved_layout == "stacked":
            anchor = _add_project_metadata_after(
                anchor,
                metadata=metadata,
                template=(project_metadata_template or project_title_template),
            )

        for bullet in bullets:
            if str(bullet).strip():
                anchor = _add_project_bullet_after(
                    anchor,
                    bullet=str(bullet).strip(),
                    template=project_bullet_template,
                )

    if spacing_mode == "blank_line":
        for _ in range(blank_lines_after_projects):
            anchor = _add_blank_line_after(anchor)
    else:
        try:
            anchor.paragraph_format.space_after = Pt(after_projects_spacing_pt)
        except Exception:
            pass

    # for project_index, project in  enumerate(projects):
    #     title = _format_project_heading(project)
    #     period = str(project.get("period", "")).strip()
    #     bullets = project.get("draft_bullets", [])[:max_bullets_per_project]

    #     anchor = _add_project_title_after(
    #         anchor,
    #         title=title,
    #         period=period,
    #         template=project_title_template,
    #         right_tab_position=right_tab_position,
    #         add_space_before=project_index > 0,
    #     )

    #     for bullet in bullets:
    #         if str(bullet).strip():
    #             anchor = _add_project_bullet_after(
    #                 anchor,
    #                 bullet=str(bullet).strip(),
    #                 template=project_bullet_template,
    #             )
    
    #     # Add spacing between the final project entry and the next section heading, e.g. SKILLS.
    # try:
    #     anchor.paragraph_format.space_after = Pt(10)
    # except Exception:
    #     pass

# ---------------------------------------------------------------------------
# Saved Docx Loader
# ---------------------------------------------------------------------------
def get_latest_saved_docx_for_application(application_id: int | None) -> Path | None:
    """
    Return the latest saved DOCX for an application session.

    This works because saved files are named with app_{application_id}_.
    """
    if application_id is None:
        return None

    if not SAVED_RESUME_DIR.exists():
        return None

    matches = sorted(
        SAVED_RESUME_DIR.glob(f"app_{application_id}_*.docx"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )

    if not matches:
        return None

    return matches[0]

# ---------------------------------------------------------------------------
# Main generation function
# ---------------------------------------------------------------------------

def generate_tailored_resume_copy(
    *,
    saved_resume_docx_path: str | Path,
    tailored_projects: dict[str, Any] | None = None,
    tailored_skills: dict[str, Any] | None = None,
    application_id: int | None = None,
    max_projects: int = 3,
    max_bullets_per_project: int = 2,
    spacing_mode: str = "paragraph_spacing",
    project_spacing_pt: int = 10,
    after_projects_spacing_pt: int = 10,
    blank_lines_between_projects: int = 1,
    blank_lines_after_projects: int = 1,
    add_spacing_before_first_project: bool = False,
    margin_profile: str = "source",
    project_header_layout: str = "auto",
    project_metadata_style: str = "pipes",
) -> Path:
    """
    Generate a new tailored resume DOCX copy.

    Only Skills and Projects are changed.
    Work Experience is not changed.
    """
    saved_resume_docx_path = Path(saved_resume_docx_path)

    if not saved_resume_docx_path.exists():
        raise FileNotFoundError(f"Saved resume DOCX not found: {saved_resume_docx_path}")

    TAILORED_RESUME_DIR.mkdir(parents=True, exist_ok=True)

    # cleanup_old_tailored_outputs_for_application(application_id)

    # timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")

    app_part = f"app_{application_id}_" if application_id is not None else ""
    output_path = TAILORED_RESUME_DIR / f"{app_part}tailored_resume_{timestamp}.docx"

    # Never edit the saved original directly. Always copy first.
    shutil.copy2(saved_resume_docx_path, output_path)

    document = Document(output_path)
    _apply_margin_profile(document, margin_profile)

    # Change only the sections that were generated.
    if not tailored_projects and not tailored_skills:
        raise ValueError("Generate a tailored Projects section or Skills section first.")

    if tailored_skills:
        replace_skills_section(document, tailored_skills)

    if tailored_projects:
        rendered_projects = apply_source_project_display_fallbacks(
            saved_resume_docx_path,
            tailored_projects,
        )
        replace_projects_section(
            document,
            rendered_projects,
            max_projects=max_projects,
            max_bullets_per_project=max_bullets_per_project,
            spacing_mode=spacing_mode,
            project_spacing_pt=project_spacing_pt,
            after_projects_spacing_pt=after_projects_spacing_pt,
            blank_lines_between_projects=blank_lines_between_projects,
            blank_lines_after_projects=blank_lines_after_projects,
            add_spacing_before_first_project=add_spacing_before_first_project,
            project_header_layout=project_header_layout,
            project_metadata_style=project_metadata_style,
        )

    # # Change only these sections.
    # replace_skills_section(document, tailored_skills)
    # replace_projects_section(
    #     document,
    #     tailored_projects,
    #     max_projects=max_projects,
    #     max_bullets_per_project=max_bullets_per_project,
    # )

    document.save(output_path)

    return output_path


# def _project_priority_score(project: dict[str, Any]) -> int:
#     """Rank projects so weaker projects are reduced first."""
#     priority = str(project.get("priority", "")).lower()

#     if priority == "high":
#         base = 300
#     elif priority == "medium":
#         base = 200
#     elif priority == "low":
#         base = 100
#     else:
#         base = 150

#     matched_count = len(project.get("matched_jd_requirements", []) or [])

#     return base + matched_count


# def _trim_words(text: str, max_words: int) -> str:
#     """Trim a bullet to a rough word limit."""
#     words = str(text).split()

#     if len(words) <= max_words:
#         return str(text).strip()

#     trimmed = " ".join(words[:max_words]).rstrip(".,;:")
#     return trimmed + "."

# old version
# def compact_tailored_projects_for_space(
#     tailored_projects: dict[str, Any],
#     *,
#     attempt: int,
# ) -> dict[str, Any]:
#     """
#     Reduce Projects section size when the generated resume exceeds one page.

#     attempt 1: keep 3 projects, high priority gets up to 2 bullets, others 1.
#     attempt 2: keep 3 projects, everyone gets 1 shorter bullet.
#     attempt 3: keep 2 strongest projects, everyone gets 1 short bullet.
#     """
#     compacted = deepcopy(tailored_projects)
#     projects = compacted.get("recommended_projects", [])

#     projects = sorted(projects, key=_project_priority_score, reverse=True)

#     if attempt == 1:
#         max_projects = 3
#         high_bullets = 2
#         other_bullets = 1
#         max_words = 24
#     elif attempt == 2:
#         max_projects = 3
#         high_bullets = 1
#         other_bullets = 1
#         max_words = 20
#     else:
#         max_projects = 2
#         high_bullets = 1
#         other_bullets = 1
#         max_words = 18

#     kept_projects = []

#     for project in projects[:max_projects]:
#         priority = str(project.get("priority", "")).lower()
#         bullet_limit = high_bullets if priority == "high" else other_bullets

#         bullets = project.get("draft_bullets", []) or []
#         project["draft_bullets"] = [
#             _trim_words(bullet, max_words)
#             for bullet in bullets[:bullet_limit]
#             if str(bullet).strip()
#         ]

#         if len(project["draft_bullets"]) <= 1:
#             project["space_action"] = "single_bullet"
#         else:
#             project["space_action"] = "shorten"

#         kept_projects.append(project)

#     removed_projects = projects[max_projects:]
#     compacted["recommended_projects"] = kept_projects

#     compacted.setdefault("projects_to_remove_or_deprioritize", [])

#     for project in removed_projects:
#         compacted["projects_to_remove_or_deprioritize"].append(
#             {
#                 "title": project.get("display_title") or project.get("title", "Untitled Project"),
#                 "reason": "Removed during compacting because the generated resume exceeded one page.",
#             }
#         )

#     compacted.setdefault("notes_for_user", []).append(
#         f"Applied compact mode attempt {attempt} to reduce page overflow."
#     )

#     return compacted

# def compact_tailored_projects_for_space(
#     tailored_projects: dict[str, Any],
#     *,
#     attempt: int,
# ) -> dict[str, Any]:
#     """
#     Reduce the Projects section only after the full generated resume
#     exceeds one page.

#     attempt 1:
#         Keep 3 projects.
#         High priority: up to 3 bullets.
#         Medium priority: up to 2 bullets.
#         Low priority: up to 1 bullet.

#     attempt 2:
#         Keep 3 projects.
#         High and medium priority: up to 2 bullets.
#         Low priority: 1 bullet.

#     attempt 3:
#         Keep 3 projects.
#         High priority: up to 2 bullets.
#         Medium and low priority: 1 bullet.

#     attempt 4+:
#         Keep the 2 strongest projects with 1 bullet each.
#     """
#     compacted = deepcopy(tailored_projects)
#     projects = compacted.get("recommended_projects", [])

#     # Sort by relevance only to decide which projects survive compaction.
#     projects = sorted(
#         projects,
#         key=_project_priority_score,
#         reverse=True,
#     )

#     if attempt == 1:
#         max_projects = 3
#         max_words = 24
#         bullet_limits_by_priority = {
#             "high": 3,
#             "medium": 2,
#             "low": 1,
#         }
#         note = (
#             "Kept most selected content while trimming long bullets. "
#             "Medium-priority projects retain up to two bullets."
#         )

#     elif attempt == 2:
#         max_projects = 3
#         max_words = 22
#         bullet_limits_by_priority = {
#             "high": 2,
#             "medium": 2,
#             "low": 1,
#         }
#         note = (
#             "Reduced project detail moderately because the resume "
#             "still exceeded one page."
#         )

#     elif attempt == 3:
#         max_projects = 3
#         max_words = 20
#         bullet_limits_by_priority = {
#             "high": 2,
#             "medium": 1,
#             "low": 1,
#         }
#         note = (
#             "Reduced lower-priority project detail further because "
#             "the resume still exceeded one page."
#         )

#     else:
#         max_projects = 2
#         max_words = 18
#         bullet_limits_by_priority = {
#             "high": 1,
#             "medium": 1,
#             "low": 1,
#         }
#         note = (
#             "Kept only the two strongest projects because the resume "
#             "still exceeded one page."
#         )

#     kept_projects = []

#     for project in projects[:max_projects]:
#         priority = str(project.get("priority", "")).lower()
#         bullet_limit = bullet_limits_by_priority.get(priority, 1)

#         original_bullets = project.get("draft_bullets", []) or []

#         project["draft_bullets"] = [
#             _trim_words(bullet, max_words)
#             for bullet in original_bullets[:bullet_limit]
#             if str(bullet).strip()
#         ]

#         if len(project["draft_bullets"]) == 1:
#             project["space_action"] = "single_bullet"
#         elif len(project["draft_bullets"]) < len(original_bullets):
#             project["space_action"] = "shorten"
#         else:
#             project["space_action"] = "keep_full"

#         kept_projects.append(project)

#     removed_projects = projects[max_projects:]

#     # Sort the final displayed projects by latest period first.
#     kept_projects = sorted(
#         kept_projects,
#         key=lambda project: period_sort_value(project.get("period", "")),
#         reverse=True,
#     )

#     compacted["recommended_projects"] = kept_projects
#     compacted.setdefault("projects_to_remove_or_deprioritize", [])

#     for project in removed_projects:
#         compacted["projects_to_remove_or_deprioritize"].append(
#             {
#                 "title": (
#                     project.get("display_title")
#                     or project.get("title", "Untitled Project")
#                 ),
#                 "reason": (
#                     "Removed during compacting because the generated "
#                     "resume exceeded one page."
#                 ),
#             }
#         )

#     compacted.setdefault("notes_for_user", []).append(
#         f"Applied compact mode attempt {attempt}: {note}"
#     )

#     return compacted


def _project_reduction_quality_loss(
    change: dict[str, Any],
) -> int:
    """Estimate evidence loss for one project fitting change."""
    phase6c_loss = change.get("evidence_loss_score")
    if phase6c_loss is not None:
        return max(0, int(phase6c_loss or 0))

    priority = int(change.get("project_priority_score", 0) or 0)
    change_type = str(change.get("change_type", ""))

    if change_type == "compact_rewrite":
        removed_words = max(
            0,
            int(change.get("full_word_count", 0) or 0)
            - int(change.get("compact_word_count", 0) or 0),
        )
        return 200 + priority // 2 + removed_words

    if change_type == "remove_bullet":
        removed_words = len(
            str(change.get("removed_bullet", "")).split()
        )
        return 600 + priority + removed_words * 5

    if change_type == "remove_project":
        return 10000 + priority * 10

    return 100000

def _skill_reduction_quality_loss(change: dict[str, Any]) -> int:
    """Estimate evidence loss for removing one Skills item."""
    return 100 + int(change.get("skill_priority_score", 0) or 0)


def _whole_resume_change_key(change: dict[str, Any]) -> str:
    section = str(change.get("section", "projects"))

    if section == "skills":
        return (
            "skills:"
            + str(change.get("category", ""))
        )

    return (
        "projects:"
        + str(change.get("project", ""))
    )


def _whole_resume_restorable_change_indices(
    active_changes: list[dict[str, Any]],
) -> list[int]:
    """Return fitting changes that can be reversed without later conflicts."""
    restorable: list[int] = []

    for index, change in enumerate(active_changes):
        resource_key = _whole_resume_change_key(change)
        later_same_resource = any(
            _whole_resume_change_key(later_change) == resource_key
            for later_change in active_changes[index + 1:]
        )

        if not later_same_resource:
            restorable.append(index)

    return restorable


def _restore_whole_resume_change(
    tailored_projects: dict[str, Any] | None,
    tailored_skills: dict[str, Any] | None,
    change: dict[str, Any],
) -> tuple[
    dict[str, Any] | None,
    dict[str, Any] | None,
    bool,
    dict[str, Any],
]:
    section = str(change.get("section", "projects"))

    if section == "skills":
        if not isinstance(tailored_skills, dict):
            return tailored_projects, tailored_skills, False, {
                "section": "skills",
                "change_type": "restore_unavailable",
                "reason": "The tailored Skills result is unavailable.",
            }

        restored_skills, restored, restore_info = restore_skill_change(
            tailored_skills,
            change,
        )
        return tailored_projects, restored_skills, restored, restore_info

    if not isinstance(tailored_projects, dict):
        return tailored_projects, tailored_skills, False, {
            "section": "projects",
            "change_type": "restore_unavailable",
            "reason": "The tailored Projects result is unavailable.",
        }

    restored_projects, restored, restore_info = _restore_fitting_change(
        tailored_projects,
        change,
    )
    restore_info["section"] = "projects"
    return restored_projects, tailored_skills, restored, restore_info


def _whole_resume_restoration_quality_gain(change: dict[str, Any]) -> int:
    if str(change.get("section", "projects")) == "skills":
        return skill_restoration_quality_gain(change)
    return _restoration_quality_gain(change)




_LAYOUT_EFFECT_THRESHOLD = 0.002
_PAGE_DENSITY_MAX_FILL: dict[str, float | None] = {
    "none": None,
    "balanced": 0.92,
    "maximize": 0.97,
}


def _normalise_page_density_mode(value: str) -> str:
    mode = str(value or "balanced").strip().lower()
    return mode if mode in _PAGE_DENSITY_MAX_FILL else "balanced"


def _rendered_overflow_value(rendered: dict[str, Any]) -> float:
    """Return comparable overflow units above the one-page target."""
    page_count = rendered.get("page_count")
    if page_count is None:
        return float("inf")

    if int(page_count) <= 1:
        return 0.0

    metrics = rendered.get("fill_metrics", {}) or {}
    overflow_ratio = metrics.get("overflow_ratio")
    if overflow_ratio is not None:
        return max(0.0, float(overflow_ratio))

    return float(max(1, int(page_count) - 1))


def _change_identity(change: dict[str, Any]) -> str:
    """Create a deterministic identity for reduction/restoration probes."""
    return "|".join(
        [
            str(change.get("section", "projects")),
            str(change.get("change_type", "")),
            str(change.get("project", "")),
            str(change.get("category", "")),
            str(change.get("removed_skill", "")),
            str(change.get("removed_bullet_index", "")),
        ]
    )


def _choose_layout_aware_reduction(
    candidates: list[dict[str, Any]],
) -> dict[str, Any]:
    """
    Choose a rendered reduction while protecting requirement evidence.

    Phase 6C first limits consideration to candidates with the lowest
    protection tier that produce a measurable layout improvement. It then uses
    the existing one-page/efficiency comparison inside that safer tier.
    """
    if not candidates:
        raise ValueError("No rendered fitting candidates were supplied.")

    effective = [
        candidate
        for candidate in candidates
        if candidate.get("reaches_one_page")
        or float(candidate.get("space_saved_ratio", 0.0) or 0.0)
        >= _LAYOUT_EFFECT_THRESHOLD
    ]
    pool = effective or list(candidates)

    def protection_tier(candidate: dict[str, Any]) -> int:
        change = candidate.get("change", {}) or {}
        return max(0, int(change.get("protection_tier", 0) or 0))

    safest_tier = min(protection_tier(candidate) for candidate in pool)
    pool = [
        candidate
        for candidate in pool
        if protection_tier(candidate) == safest_tier
    ]

    def key(
        candidate: dict[str, Any],
    ) -> tuple[float, float, int, int]:
        reaches_one_page = bool(candidate.get("reaches_one_page"))
        space_saved = float(
            candidate.get("space_saved_ratio", 0.0) or 0.0
        )
        quality_loss = float(candidate.get("quality_loss", 0) or 0)
        candidate_order = int(
            candidate.get("candidate_order", 99) or 99
        )
        change = candidate.get("change", {}) or {}
        evidence_priority = max(
            0,
            int(change.get("evidence_priority", 0) or 0),
        )

        if reaches_one_page:
            return (
                0.0,
                quality_loss,
                -evidence_priority,
                candidate_order,
            )

        if space_saved >= _LAYOUT_EFFECT_THRESHOLD:
            return (
                1.0,
                quality_loss / space_saved,
                -evidence_priority,
                candidate_order,
            )

        return (
            2.0,
            quality_loss,
            -evidence_priority,
            candidate_order,
        )

    return min(pool, key=key)


def resolve_fitting_bullet_allocation_mode(
    tailored_projects: dict[str, Any] | None,
    *,
    fallback_mode: str,
) -> str:
    """Return persisted allocation provenance when the payload carries it."""
    allocation = (
        (((tailored_projects or {}).get("deterministic_rule_debug") or {}).get(
            "bullet_allocation"
        ))
        or {}
    )
    payload_mode = str(
        allocation.get("allocation_mode") or ""
    ).strip().lower()
    if payload_mode in {
        "adaptive",
        "prefer_available_evidence",
        "all_canonical_before_fitting",
    }:
        return payload_mode
    return str(fallback_mode or "adaptive").strip().lower() or "adaptive"


def resolve_effective_fitting_bullet_ceiling(
    tailored_projects: dict[str, Any] | None,
    *,
    configured_max_bullets_per_project: int,
) -> int:
    """Resolve the pre-render ceiling from immutable allocation provenance.

    Adaptive and Prefer-available payloads retain the configured UI ceiling.
    An all-canonical payload is different: its persisted deterministic
    allocation is the authoritative pre-fit content, so every allocated bullet
    must reach the first render even if mutable/restored widget state says 4.
    """
    try:
        configured = max(1, int(configured_max_bullets_per_project))
    except (TypeError, ValueError):
        configured = 1

    projects_state = tailored_projects or {}
    if (
        resolve_fitting_bullet_allocation_mode(
            projects_state,
            fallback_mode="adaptive",
        )
        != "all_canonical_before_fitting"
    ):
        return configured

    allocated_counts = [
        len(project.get("draft_bullets", []) or [])
        for project in (
            projects_state.get("recommended_projects", []) or []
        )
        if isinstance(project, dict)
    ]
    return max([configured, *allocated_counts])

@dataclass
class PreparedFittingInput:
    """One immutable-in-practice pre-render input shared with Phase 9F."""

    fitting_input_snapshot: dict[str, Any]
    working_projects: dict[str, Any] | None
    working_skills: dict[str, Any] | None
    source_signature: str
    density_mode: str
    density_max_fill: float | None
    requested_project_header_layout: str
    active_project_header_layout: str
    active_project_metadata_style: str
    lock_policy: dict[str, Any]
    payload_max_bullets: int
    effective_max_projects: int
    effective_max_bullets: int


def prepare_fitting_input_snapshot(
    *,
    saved_resume_docx_path: str | Path,
    tailored_projects: dict[str, Any] | None,
    tailored_skills: dict[str, Any] | None,
    max_projects: int,
    max_bullets_per_project: int,
    spacing_mode: str,
    project_spacing_pt: int,
    after_projects_spacing_pt: int,
    blank_lines_between_projects: int,
    blank_lines_after_projects: int,
    add_spacing_before_first_project: bool,
    use_compact_before_delete: bool,
    prefer_balanced_bullets: bool,
    allow_skills_compaction: bool,
    lock_projects: bool,
    lock_skills: bool,
    minimum_total_skills: int,
    page_density_mode: str,
    allow_margin_compaction: bool,
    project_header_layout: str,
    project_metadata_style: str,
    source_artifact_identity: dict[str, Any] | None,
) -> PreparedFittingInput:
    """Build caller/prepared fitting provenance before the first render.

    The existing source-display fallback and deterministic ceiling preparation
    remain exactly where they were.  This helper records both sides of that
    boundary once so Phase 9F can persist and pass the same pre-render state to
    the expensive fitter.
    """
    density_mode = _normalise_page_density_mode(page_density_mode)
    density_max_fill = _PAGE_DENSITY_MAX_FILL[density_mode]
    requested_project_header_layout = normalise_project_header_layout(
        project_header_layout
    )
    active_project_header_layout = (
        "stacked"
        if requested_project_header_layout == "auto"
        else requested_project_header_layout
    )
    active_project_metadata_style = normalise_project_metadata_style(
        project_metadata_style
    )
    lock_policy = build_fitting_lock_policy(
        lock_projects=lock_projects,
        lock_skills=lock_skills,
    )
    payload_max_bullets = resolve_effective_fitting_bullet_ceiling(
        tailored_projects,
        configured_max_bullets_per_project=max_bullets_per_project,
    )
    effective_max_projects = (
        999999 if lock_policy["lock_projects"] else max_projects
    )
    effective_max_bullets = (
        999999
        if lock_policy["lock_projects"]
        else payload_max_bullets
    )

    # Preserve caller input separately.  The fallback returns a deep copy, so
    # neither this preparation nor later candidate search mutates it.
    working_projects = (
        apply_source_project_display_fallbacks(
            saved_resume_docx_path,
            tailored_projects,
        )
        if tailored_projects
        else None
    )
    working_skills = deepcopy(tailored_skills) if tailored_skills else None

    if working_projects and not lock_policy["lock_projects"]:
        visible_projects = (
            working_projects.get("recommended_projects", []) or []
        )[:max_projects]
        for project in visible_projects:
            project["draft_bullets"] = (
                project.get("draft_bullets", []) or []
            )[:payload_max_bullets]
            project["compact_bullets"] = (
                project.get("compact_bullets", []) or []
            )[:payload_max_bullets]
            sync_project_bullet_metadata(
                project,
                bullet_texts=project["draft_bullets"],
            )
        working_projects["recommended_projects"] = visible_projects

    source_signature = source_docx_signature(saved_resume_docx_path)
    source_docx_bytes = Path(saved_resume_docx_path).read_bytes()
    fitting_input_snapshot = build_fitting_input_snapshot(
        source_docx_sha256=hashlib.sha256(source_docx_bytes).hexdigest(),
        source_docx_byte_size=len(source_docx_bytes),
        source_artifact_identity=source_artifact_identity,
        caller_projects=tailored_projects,
        caller_skills=tailored_skills,
        prepared_projects=working_projects,
        prepared_skills=working_skills,
        fitter_invocation={
            # These two fields describe the limits the renderer/search actually
            # receives.  Preserve the incoming values separately because the
            # lock policy and canonical allocation can resolve them differently.
            "max_projects": effective_max_projects,
            "max_bullets_per_project": effective_max_bullets,
            "requested_max_projects": max_projects,
            "requested_max_bullets_per_project": max_bullets_per_project,
            "prepared_payload_max_bullets_per_project": payload_max_bullets,
            "spacing_mode": spacing_mode,
            "project_spacing_pt": project_spacing_pt,
            "after_projects_spacing_pt": after_projects_spacing_pt,
            "blank_lines_between_projects": blank_lines_between_projects,
            "blank_lines_after_projects": blank_lines_after_projects,
            "add_spacing_before_first_project": add_spacing_before_first_project,
            "use_compact_before_delete": use_compact_before_delete,
            "prefer_balanced_bullets": prefer_balanced_bullets,
            "allow_skills_compaction": allow_skills_compaction,
            "lock_projects": lock_projects,
            "lock_skills": lock_skills,
            "minimum_total_skills": minimum_total_skills,
            "page_density_mode": density_mode,
            "page_density_max_fill": density_max_fill,
            "allow_margin_compaction": allow_margin_compaction,
            "project_header_layout": active_project_header_layout,
            "requested_project_header_layout": requested_project_header_layout,
            "project_metadata_style": active_project_metadata_style,
        },
        fitting_policy_version=PHASE6C_FITTING_VERSION,
    )
    return PreparedFittingInput(
        fitting_input_snapshot=fitting_input_snapshot,
        working_projects=working_projects,
        working_skills=working_skills,
        source_signature=source_signature,
        density_mode=density_mode,
        density_max_fill=density_max_fill,
        requested_project_header_layout=requested_project_header_layout,
        active_project_header_layout=active_project_header_layout,
        active_project_metadata_style=active_project_metadata_style,
        lock_policy=lock_policy,
        payload_max_bullets=payload_max_bullets,
        effective_max_projects=effective_max_projects,
        effective_max_bullets=effective_max_bullets,
    )

def generate_tailored_resume_copy_fit_one_page(
    *,
    saved_resume_docx_path: str | Path,
    tailored_projects: dict[str, Any] | None = None,
    tailored_skills: dict[str, Any] | None = None,
    application_id: int | None = None,
    max_projects: int = 3,
    max_bullets_per_project: int = 3,
    spacing_mode: str = "paragraph_spacing",
    project_spacing_pt: int = 10,
    after_projects_spacing_pt: int = 10,
    blank_lines_between_projects: int = 1,
    blank_lines_after_projects: int = 1,
    add_spacing_before_first_project: bool = False,
    use_compact_before_delete: bool = False,
    prefer_balanced_bullets: bool = False,
    allow_skills_compaction: bool = False,
    lock_projects: bool = False,
    lock_skills: bool = False,
    minimum_total_skills: int = DEFAULT_MINIMUM_TOTAL_SKILLS,
    page_density_mode: str = "balanced",
    allow_margin_compaction: bool = False,
    project_header_layout: str = "auto",
    project_metadata_style: str = "pipes",
    generation_id: str | None = None,
    source_artifact_identity: dict[str, Any] | None = None,
    prepared_fitting_input: PreparedFittingInput | None = None,
) -> dict[str, Any]:
    """
    Fit the tailored resume to one page using layout-aware candidate probes.

    Phase 5.1 does not rerun résumé analysis, project selection, or the LLM.
    It compares the already-generated Projects and Skills using deterministic
    evidence-loss estimates and actual rendered PDF space saved.
    """
    if not tailored_projects and not tailored_skills:
        raise ValueError(
            "Generate a tailored Projects section or Skills section first."
        )

    fitting_started_at = time.perf_counter()

    def _fit_log(message: str) -> None:
        print(f"[FIT] {message}", flush=True)

    preparation = prepared_fitting_input or prepare_fitting_input_snapshot(
        saved_resume_docx_path=saved_resume_docx_path,
        tailored_projects=tailored_projects,
        tailored_skills=tailored_skills,
        max_projects=max_projects,
        max_bullets_per_project=max_bullets_per_project,
        spacing_mode=spacing_mode,
        project_spacing_pt=project_spacing_pt,
        after_projects_spacing_pt=after_projects_spacing_pt,
        blank_lines_between_projects=blank_lines_between_projects,
        blank_lines_after_projects=blank_lines_after_projects,
        add_spacing_before_first_project=add_spacing_before_first_project,
        use_compact_before_delete=use_compact_before_delete,
        prefer_balanced_bullets=prefer_balanced_bullets,
        allow_skills_compaction=allow_skills_compaction,
        lock_projects=lock_projects,
        lock_skills=lock_skills,
        minimum_total_skills=minimum_total_skills,
        page_density_mode=page_density_mode,
        allow_margin_compaction=allow_margin_compaction,
        project_header_layout=project_header_layout,
        project_metadata_style=project_metadata_style,
        source_artifact_identity=source_artifact_identity,
    )
    if not isinstance(preparation, PreparedFittingInput):
        raise TypeError("prepared_fitting_input must be a PreparedFittingInput")

    density_mode = preparation.density_mode
    density_max_fill = preparation.density_max_fill
    requested_project_header_layout = preparation.requested_project_header_layout
    active_project_header_layout = preparation.active_project_header_layout
    active_project_metadata_style = preparation.active_project_metadata_style
    project_header_compaction_used = False
    lock_policy = preparation.lock_policy
    effective_max_projects = preparation.effective_max_projects
    effective_max_bullets = preparation.effective_max_bullets

    # Preserve completed historical outputs. Temporary candidates are still
    # deleted explicitly by the fitting loop.
    attempt_logs: list[dict[str, Any]] = []
    working_projects = deepcopy(preparation.working_projects)
    working_skills = deepcopy(preparation.working_skills)

    active_changes: list[dict[str, Any]] = []

    source_signature = preparation.source_signature
    fitting_input_snapshot = preparation.fitting_input_snapshot
    render_layout_options = {
        "max_projects": effective_max_projects,
        "max_bullets_per_project": effective_max_bullets,
        "spacing_mode": spacing_mode,
        "project_spacing_pt": project_spacing_pt,
        "after_projects_spacing_pt": after_projects_spacing_pt,
        "blank_lines_between_projects": blank_lines_between_projects,
        "blank_lines_after_projects": blank_lines_after_projects,
        "add_spacing_before_first_project": (
            add_spacing_before_first_project
        ),
        "project_metadata_style": active_project_metadata_style,
    }
    render_cache: dict[str, dict[str, Any]] = {}
    active_margin_profile = "source"
    margin_compaction_used = False
    optimization_stats: dict[str, int] = {
        "candidate_state_requests": 0,
        "render_cache_hits": 0,
        "render_cache_misses": 0,
        "libreoffice_batch_processes": 0,
        "libreoffice_fallback_processes": 0,
        "reduction_tier_batches": 0,
        "reduction_candidates_generated": 0,
        "reduction_candidates_rendered": 0,
        "reduction_candidates_skipped_by_tier_stop": 0,
        "restoration_batch_count": 0,
        "restoration_candidates_rendered": 0,
        "coarse_render_count": 0,
        "exact_render_count": 0,
        "local_refinement_render_count": 0,
        "render_budget": FIT_RENDER_BUDGET,
        "render_budget_used": 0,
        "render_budget_exhausted": 0,
        "candidate_states_rendered": 0,
        "render_verification_failed": 0,
        "libreoffice_timeout_count": 0,
    }
    timing_stats: dict[str, float] = {
        "render_elapsed_seconds": 0.0,
        "libreoffice_elapsed_seconds": 0.0,
        "candidate_generation_elapsed_seconds": 0.0,
    }

    def optimization_summary() -> dict[str, Any]:
        return {
            "fitting_optimization_version": (
                PHASE6C1_OPTIMIZATION_VERSION
            ),
            "fitting_search_algorithm_version": (
                PHASE6C_SEARCH_ALGORITHM_VERSION
            ),
            "render_state_fingerprint_version": (
                PHASE6C1_OPTIMIZATION_VERSION
            ),
            "fitting_input_snapshot": deepcopy(fitting_input_snapshot),
            "fitting_input_fingerprint": fitting_input_snapshot[
                "fitting_input_fingerprint"
            ],
            "section_locks": dict(lock_policy),
            "render_cache_entry_count": len(render_cache),
            "libreoffice_process_count": (
                optimization_stats[
                    "libreoffice_batch_processes"
                ]
                + optimization_stats[
                    "libreoffice_fallback_processes"
                ]
            ),
            "fitting_elapsed_seconds": round(
                time.perf_counter() - fitting_started_at,
                3,
            ),
            **{
                key: round(value, 3)
                for key, value in timing_stats.items()
            },
            **optimization_stats,
        }

    def _empty_fill_metrics() -> dict[str, Any]:
        return {
            "page_fill_ratio": None,
            "estimated_unused_page_ratio": None,
            "last_page_fill_ratio": None,
            "overflow_ratio": None,
            "occupied_page_units": None,
            "measurement_method": "unavailable",
        }

    def _candidate_counts(
        projects_state: dict[str, Any] | None,
        skills_state: dict[str, Any] | None,
    ) -> dict[str, int]:
        projects = (
            (projects_state or {}).get(
                "recommended_projects",
                [],
            )
            or []
        )
        return {
            "project_count": len(projects),
            "bullet_count": sum(
                len(project.get("draft_bullets", []) or [])
                for project in projects
            ),
            "skill_line_count": len(
                (skills_state or {}).get(
                    "skill_lines",
                    [],
                )
                or []
            ),
            "skill_item_count": count_skill_items(
                skills_state
            ),
        }

    initial_counts = _candidate_counts(working_projects, working_skills)
    _fit_log(
        "start "
        f"projects={initial_counts['project_count']} "
        f"bullets={initial_counts['bullet_count']} "
        f"skills={initial_counts['skill_item_count']}"
    )

    def _prepare_candidate(
        specification: dict[str, Any],
    ) -> dict[str, Any]:
        projects_state = specification.get("projects")
        skills_state = specification.get("skills")
        margin_profile = _normalise_margin_profile(
            specification.get("margin_profile") or active_margin_profile
        )
        candidate_project_header_layout = normalise_project_header_layout(
            specification.get("project_header_layout")
            or active_project_header_layout
        )
        if candidate_project_header_layout == "auto":
            candidate_project_header_layout = "stacked"
        candidate_layout_options = {
            **render_layout_options,
            "margin_profile": margin_profile,
            "project_header_layout": candidate_project_header_layout,
        }
        fingerprint = build_render_state_fingerprint(
            source_signature=source_signature,
            projects_state=projects_state,
            skills_state=skills_state,
            layout_options=candidate_layout_options,
        )
        docx_path = generate_tailored_resume_copy(
            saved_resume_docx_path=saved_resume_docx_path,
            tailored_projects=projects_state,
            tailored_skills=skills_state,
            application_id=application_id,
            max_projects=effective_max_projects,
            max_bullets_per_project=effective_max_bullets,
            spacing_mode=spacing_mode,
            project_spacing_pt=project_spacing_pt,
            after_projects_spacing_pt=after_projects_spacing_pt,
            blank_lines_between_projects=(
                blank_lines_between_projects
            ),
            blank_lines_after_projects=(
                blank_lines_after_projects
            ),
            add_spacing_before_first_project=(
                add_spacing_before_first_project
            ),
            margin_profile=margin_profile,
            project_header_layout=candidate_project_header_layout,
            project_metadata_style=active_project_metadata_style,
        )
        return {
            **specification,
            "margin_profile": margin_profile,
            "project_header_layout": candidate_project_header_layout,
            "docx_path": Path(docx_path),
            "render_state_fingerprint": fingerprint,
        }

    def _complete_prepared_candidate(
        prepared: dict[str, Any],
        *,
        pdf_path: Path | None,
        batch_label: str,
        batch_size: int,
        cache_hit: bool,
    ) -> dict[str, Any]:
        docx_path = Path(prepared["docx_path"])
        fingerprint = str(
            prepared["render_state_fingerprint"]
        )
        cached = render_cache.get(fingerprint)

        if cache_hit and cached is not None:
            expected_pdf_path = (
                PREVIEW_DIR.resolve()
                / f"{docx_path.stem}.pdf"
            )
            expected_pdf_path.parent.mkdir(
                parents=True,
                exist_ok=True,
            )
            expected_pdf_path.write_bytes(
                cached["pdf_bytes"]
            )
            pdf_path = expected_pdf_path
            page_count = int(cached["page_count"])
            fill_metrics = deepcopy(
                cached["fill_metrics"]
            )
        elif pdf_path is not None:
            page_count = count_pdf_pages(pdf_path)
            if page_count is None:
                # A PDF without a page count is not a verified render. Do not
                # infer its fit from a converter exit code or its file size.
                fill_metrics = _empty_fill_metrics()
                pdf_path = None
            else:
                fill_metrics = measure_pdf_page_fill(
                    pdf_path
                )
                try:
                    pdf_bytes = Path(pdf_path).read_bytes()
                except OSError:
                    pdf_bytes = b""

                if pdf_bytes:
                    render_cache[fingerprint] = {
                        "pdf_bytes": pdf_bytes,
                        "page_count": page_count,
                        "fill_metrics": deepcopy(
                            fill_metrics
                        ),
                    }
        else:
            page_count = None
            fill_metrics = _empty_fill_metrics()

        entry: dict[str, Any] = {
            "attempt": len(attempt_logs) + 1,
            "attempt_type": prepared["attempt_type"],
            "docx_path": str(docx_path),
            "pdf_path": (
                str(pdf_path)
                if pdf_path is not None
                else None
            ),
            "page_count": page_count,
            "render_cache_hit": cache_hit,
            "render_state_fingerprint": fingerprint,
            "libreoffice_batch_label": batch_label,
            "libreoffice_batch_size": batch_size,
        }

        if page_count is not None:
            entry.update(
                _candidate_counts(
                    prepared.get("projects"),
                    prepared.get("skills"),
                )
            )
            entry.update(fill_metrics)

        change_applied = prepared.get(
            "change_applied"
        )
        if change_applied is not None:
            entry["change_applied"] = change_applied
        if prepared.get("restoration_candidate"):
            entry["restoration_candidate"] = True
        if (
            prepared.get(
                "restoration_quality_gain"
            )
            is not None
        ):
            entry["restoration_quality_gain"] = (
                prepared[
                    "restoration_quality_gain"
                ]
            )
        if prepared.get("probe_candidate"):
            entry["probe_candidate"] = True
        if prepared.get("quality_loss") is not None:
            entry["quality_loss"] = prepared[
                "quality_loss"
            ]

        attempt_logs.append(entry)
        return {
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "page_count": page_count,
            "fill_metrics": fill_metrics,
            "attempt_entry": entry,
        }

    def render_candidates_batch(
        specifications: list[dict[str, Any]],
        *,
        batch_label: str,
    ) -> list[dict[str, Any]]:
        if not specifications:
            return []

        remaining_budget = max(
            0,
            FIT_RENDER_BUDGET
            - optimization_stats["render_budget_used"],
        )
        if remaining_budget <= 0:
            optimization_stats["render_budget_exhausted"] = 1
            _fit_log(
                f"render budget exhausted label={batch_label} "
                f"budget={FIT_RENDER_BUDGET}"
            )
            return []
        if len(specifications) > remaining_budget:
            optimization_stats["render_budget_exhausted"] = 1
            specifications = specifications[:remaining_budget]
            _fit_log(
                f"render budget limited label={batch_label} "
                f"candidates={len(specifications)} remaining={remaining_budget}"
            )

        batch_started_at = time.perf_counter()

        optimization_stats[
            "candidate_state_requests"
        ] += len(specifications)

        preparation_started_at = time.perf_counter()
        prepared_candidates = [
            _prepare_candidate(specification)
            for specification in specifications
        ]
        timing_stats["candidate_generation_elapsed_seconds"] += (
            time.perf_counter() - preparation_started_at
        )
        uncached: list[dict[str, Any]] = []

        for prepared in prepared_candidates:
            fingerprint = str(
                prepared[
                    "render_state_fingerprint"
                ]
            )
            if fingerprint in render_cache:
                optimization_stats[
                    "render_cache_hits"
                ] += 1
            else:
                optimization_stats[
                    "render_cache_misses"
                ] += 1
                uncached.append(prepared)

        converted: dict[str, Path | None] = {}
        if uncached:
            conversion_started_at = time.perf_counter()
            converted, diagnostics = (
                convert_docx_batch_to_pdf_if_possible(
                    [
                        prepared["docx_path"]
                        for prepared in uncached
                    ]
                )
            )
            conversion_elapsed = time.perf_counter() - conversion_started_at
            timing_stats["libreoffice_elapsed_seconds"] += conversion_elapsed
            _fit_log(
                "libreoffice batch "
                f"label={batch_label} candidates={len(uncached)} "
                f"elapsed={conversion_elapsed:.3f}s"
            )
            optimization_stats[
                "libreoffice_batch_processes"
            ] += int(
                diagnostics.get(
                    "batch_process_count",
                    0,
                )
                or 0
            )
            if diagnostics.get("timed_out"):
                optimization_stats["libreoffice_timeout_count"] += 1
            optimization_stats[
                "libreoffice_fallback_processes"
            ] += int(
                diagnostics.get(
                    "fallback_process_count",
                    0,
                )
                or 0
            )

        results: list[dict[str, Any]] = []
        batch_size = len(prepared_candidates)

        for prepared in prepared_candidates:
            fingerprint = str(
                prepared[
                    "render_state_fingerprint"
                ]
            )
            cache_hit = fingerprint in render_cache
            pdf_path = None

            if not cache_hit:
                pdf_path = converted.get(
                    str(
                        Path(
                            prepared["docx_path"]
                        ).resolve()
                    )
                )

            completed = _complete_prepared_candidate(
                prepared,
                pdf_path=pdf_path,
                batch_label=batch_label,
                batch_size=batch_size,
                cache_hit=cache_hit,
            )
            if not cache_hit:
                optimization_stats["render_budget_used"] += 1
                optimization_stats["candidate_states_rendered"] += 1
            if (
                completed.get("pdf_path") is None
                or completed.get("page_count") is None
            ):
                optimization_stats["render_verification_failed"] = 1
            attempt_entry = completed["attempt_entry"]
            counts = _candidate_counts(
                prepared.get("projects"),
                prepared.get("skills"),
            )
            _fit_log(
                "render "
                f"{attempt_entry['attempt']} label={batch_label} "
                f"pages={completed.get('page_count')} "
                f"bullets={counts['bullet_count']} "
                f"elapsed={time.perf_counter() - batch_started_at:.3f}s"
            )
            results.append(completed)

        timing_stats["render_elapsed_seconds"] += (
            time.perf_counter() - batch_started_at
        )

        return results

    def render_candidate(
        projects_state: dict[str, Any] | None,
        skills_state: dict[str, Any] | None,
        *,
        attempt_type: str,
        change_applied: dict[str, Any] | None = None,
        restoration_candidate: bool = False,
        restoration_quality_gain: int | None = None,
        probe_candidate: bool = False,
        quality_loss: int | None = None,
        margin_profile: str | None = None,
        project_header_layout_override: str | None = None,
    ) -> dict[str, Any]:
        rendered = render_candidates_batch(
            [
                {
                    "projects": projects_state,
                    "skills": skills_state,
                    "attempt_type": attempt_type,
                    "change_applied": change_applied,
                    "restoration_candidate": (
                        restoration_candidate
                    ),
                    "restoration_quality_gain": (
                        restoration_quality_gain
                    ),
                    "probe_candidate": probe_candidate,
                    "quality_loss": quality_loss,
                    "margin_profile": (
                        margin_profile or active_margin_profile
                    ),
                    "project_header_layout": (
                        project_header_layout_override
                        or active_project_header_layout
                    ),
                }
            ],
            batch_label=attempt_type,
        )
        if rendered:
            return rendered[0]
        return {
            "docx_path": None,
            "pdf_path": None,
            "page_count": None,
            "fill_metrics": _empty_fill_metrics(),
            "attempt_entry": {
                "attempt": len(attempt_logs) + 1,
                "attempt_type": attempt_type,
                "page_count": None,
                "render_budget_exhausted": True,
            },
        }

    def build_reduction_candidates(
        projects_state: dict[str, Any] | None,
        skills_state: dict[str, Any] | None,
    ) -> list[dict[str, Any]]:
        """Build the existing evidence-safe one-step candidates without rendering."""
        candidates: list[dict[str, Any]] = []

        if (
            use_compact_before_delete
            and projects_state
            and not lock_policy["lock_projects"]
        ):
            compact_projects, changed, change = apply_compact_bullets_once(
                projects_state
            )
            if changed:
                change = deepcopy(change)
                change["section"] = "projects"
                candidates.append(
                    {
                        "projects": compact_projects,
                        "skills": skills_state,
                        "change": change,
                        "quality_loss": _project_reduction_quality_loss(change),
                        "candidate_order": 1,
                    }
                )

        if (
            allow_skills_compaction
            and skills_state
            and not lock_policy["lock_skills"]
        ):
            compact_skills, changed, change = compact_skills_one_step(
                skills_state,
                minimum_total_items=minimum_total_skills,
            )
            if changed:
                candidates.append(
                    {
                        "projects": projects_state,
                        "skills": compact_skills,
                        "change": change,
                        "quality_loss": _skill_reduction_quality_loss(change),
                        "candidate_order": 0,
                    }
                )

        if projects_state and not lock_policy["lock_projects"]:
            project_reductions = build_evidence_aware_project_reductions(
                projects_state,
                prefer_balanced_bullets=prefer_balanced_bullets,
            )
            for reduction_index, (reduced_projects, raw_change) in enumerate(
                project_reductions,
                start=2,
            ):
                change = deepcopy(raw_change)
                change["section"] = "projects"
                candidates.append(
                    {
                        "projects": reduced_projects,
                        "skills": skills_state,
                        "change": change,
                        "quality_loss": _project_reduction_quality_loss(change),
                        "candidate_order": reduction_index,
                    }
                )
        return candidates

    def build_safe_reduction_path(
        projects_state: dict[str, Any] | None,
        skills_state: dict[str, Any] | None,
    ) -> list[dict[str, Any]]:
        """Create cumulative low-risk states using the current Phase 6C order."""
        planning_started_at = time.perf_counter()
        path: list[dict[str, Any]] = []
        current_projects = deepcopy(projects_state)
        current_skills = deepcopy(skills_state)
        cumulative_changes: list[dict[str, Any]] = []

        while True:
            candidates = build_reduction_candidates(
                current_projects,
                current_skills,
            )
            if not candidates:
                break
            optimization_stats["reduction_candidates_generated"] += len(candidates)
            tier_groups = group_candidates_by_protection_tier(candidates)
            protection_tier, lowest_tier = tier_groups[0]
            next_candidate = lowest_tier[0]
            change = deepcopy(next_candidate["change"])
            cumulative_changes.append(change)
            current_projects = deepcopy(next_candidate["projects"])
            current_skills = deepcopy(next_candidate["skills"])
            path.append(
                {
                    "projects": current_projects,
                    "skills": current_skills,
                    "changes": deepcopy(cumulative_changes),
                    "last_change": change,
                    "protection_tier": protection_tier,
                }
            )

        timing_stats["candidate_generation_elapsed_seconds"] += (
            time.perf_counter() - planning_started_at
        )
        return path

    current_render = render_candidate(
        working_projects,
        working_skills,
        attempt_type="full",
    )

    if current_render["pdf_path"] is None:
        _fit_log(
            "complete status=verification_failed "
            f"renders={optimization_stats['candidate_states_rendered']} "
            f"elapsed={time.perf_counter() - fitting_started_at:.3f}s"
        )
        return {
            "generation_id": generation_id,
            "fitting_version": PHASE6C_FITTING_VERSION,
            "docx_path": current_render["docx_path"],
            "pdf_path": None,
            "page_count": None,
            "fit_one_page": False,
            "fit_status": "verification_failed",
            "attempts": attempt_logs,
            "tailored_projects_used": working_projects,
            "tailored_skills_used": working_skills,
            "page_fill_ratio": None,
            "estimated_unused_page_ratio": None,
            "page_density_mode": density_mode,
            "density_target_max": density_max_fill,
            "margin_profile": active_margin_profile,
            "margin_compaction_used": margin_compaction_used,
            "project_header_layout_requested": requested_project_header_layout,
            "project_header_layout_used": active_project_header_layout,
            "project_metadata_style": active_project_metadata_style,
            "project_header_compaction_used": project_header_compaction_used,
            **optimization_summary(),
            "note": (
                "Could not verify page count because LibreOffice is unavailable, "
                "timed out, or DOCX-to-PDF conversion failed. DOCX generation "
                "still worked, but this result is not verified as one page."
            ),
        }

    _fit_log(
        "initial render "
        f"pages={current_render['page_count']} "
        f"overflow={_rendered_overflow_value(current_render):.3f} "
        f"elapsed={time.perf_counter() - fitting_started_at:.3f}s"
    )

    # Auto preserves the clearer stacked display until physical rendering
    # proves that inline metadata saves meaningful space. This is a format
    # choice with zero evidence loss, and is intentionally evaluated before
    # margin compaction or evidence-removing candidates.
    if (
        requested_project_header_layout == "auto"
        and working_projects
        and int(current_render["page_count"]) > 1
    ):
        inline_render = render_candidate(
            working_projects,
            working_skills,
            attempt_type="project_header_inline_compaction",
            change_applied={
                "section": "formatting",
                "change_type": "project_header_layout",
                "from": "stacked",
                "to": "inline",
                "evidence_loss_score": 0,
            },
            probe_candidate=True,
            quality_loss=0,
            project_header_layout_override="inline",
        )
        if inline_render["pdf_path"] is not None:
            baseline_overflow = _rendered_overflow_value(current_render)
            inline_overflow = _rendered_overflow_value(inline_render)
            inline_improved = (
                int(inline_render.get("page_count") or 99) <= 1
                or inline_overflow
                <= baseline_overflow - _LAYOUT_EFFECT_THRESHOLD
            )
            inline_render["attempt_entry"]["project_header_auto_probe"] = True
            inline_render["attempt_entry"]["project_header_auto_accepted"] = (
                inline_improved
            )
            if inline_improved:
                _delete_generated_output(
                    current_render.get("docx_path"),
                    current_render.get("pdf_path"),
                )
                current_render["attempt_entry"]["superseded_output_deleted"] = True
                current_render = inline_render
                active_project_header_layout = "inline"
                project_header_compaction_used = True
            else:
                _delete_generated_output(
                    inline_render.get("docx_path"),
                    inline_render.get("pdf_path"),
                )
                inline_render["attempt_entry"]["temporary_output_deleted"] = True

    fitting_render: dict[str, Any] | None = None

    if allow_margin_compaction and int(current_render["page_count"]) > 1:
        margin_profiles = _available_margin_compaction_profiles(
            saved_resume_docx_path
        )
        margin_candidates = [
            render_candidate(
                working_projects,
                working_skills,
                attempt_type=f"margin_{profile}",
                probe_candidate=True,
                quality_loss=0,
                margin_profile=profile,
            )
            for profile in margin_profiles
        ]
        usable_margin_candidates = [
            candidate for candidate in margin_candidates
            if candidate.get("pdf_path") is not None
        ]
        if usable_margin_candidates:
            baseline_overflow = _rendered_overflow_value(current_render)
            reaching_one_page = [
                candidate for candidate in usable_margin_candidates
                if int(candidate.get("page_count") or 99) <= 1
            ]
            chosen_margin = (
                reaching_one_page[0]
                if reaching_one_page
                else min(usable_margin_candidates, key=_rendered_overflow_value)
            )
            chosen_overflow = _rendered_overflow_value(chosen_margin)
            margin_improved = (
                int(chosen_margin.get("page_count") or 99) <= 1
                or chosen_overflow <= baseline_overflow - _LAYOUT_EFFECT_THRESHOLD
            )
            for candidate in usable_margin_candidates:
                if candidate is not chosen_margin or not margin_improved:
                    _delete_generated_output(
                        candidate.get("docx_path"),
                        candidate.get("pdf_path"),
                    )
            if margin_improved:
                _delete_generated_output(
                    current_render.get("docx_path"),
                    current_render.get("pdf_path"),
                )
                current_render = chosen_margin
                active_margin_profile = str(
                    chosen_margin.get("margin_profile") or "source"
                )
                margin_compaction_used = active_margin_profile != "source"

    if int(current_render["page_count"]) <= 1:
        fitting_render = current_render

    if fitting_render is None:
        reduction_path = build_safe_reduction_path(
            working_projects,
            working_skills,
        )
        _fit_log(
            "reduction stage=coarse "
            f"path_states={len(reduction_path)} "
            f"budget={FIT_RENDER_BUDGET}"
        )
        if reduction_path:
            coarse_counts: list[int] = []
            next_count = FIT_COARSE_INITIAL_REMOVALS
            while next_count < len(reduction_path):
                coarse_counts.append(next_count)
                next_count *= 2
            coarse_counts.append(len(reduction_path))
            coarse_counts = list(dict.fromkeys(coarse_counts))

            coarse_specs = [
                {
                    "projects": reduction_path[count - 1]["projects"],
                    "skills": reduction_path[count - 1]["skills"],
                    "attempt_type": "coarse_reduction",
                    "change_applied": {
                        "change_type": "cumulative_reduction",
                        "reduction_count": count,
                        "last_change": reduction_path[count - 1]["last_change"],
                    },
                    "probe_candidate": True,
                    "quality_loss": sum(
                        _skill_reduction_quality_loss(change)
                        if str(change.get("section", "projects")) == "skills"
                        else _project_reduction_quality_loss(change)
                        for change in reduction_path[count - 1]["changes"]
                    ),
                    "margin_profile": active_margin_profile,
                }
                for count in coarse_counts
            ]
            rendered_coarse = render_candidates_batch(
                coarse_specs,
                batch_label="coarse_reduction",
            )
            optimization_stats["coarse_render_count"] += len(rendered_coarse)

            coarse_results: dict[int, dict[str, Any]] = {}
            for count, rendered in zip(coarse_counts, rendered_coarse):
                if rendered.get("pdf_path") is None:
                    break
                coarse_results[count] = rendered

            if optimization_stats["render_verification_failed"]:
                fitting_render = None
            else:
                fit_count = next(
                    (
                        count
                        for count in coarse_counts
                        if (
                            count in coarse_results
                            and int(coarse_results[count]["page_count"]) <= 1
                        )
                    ),
                    None,
                )
                if fit_count is not None:
                    lower_count = max(
                        [
                            count
                            for count in [0, *coarse_counts]
                            if count < fit_count
                            and (
                                count == 0
                                or int(coarse_results[count]["page_count"]) > 1
                            )
                        ],
                        default=0,
                    )
                    exact_fit_count = fit_count
                    _fit_log(
                        "reduction stage=exact "
                        f"bracket={lower_count}:{fit_count}"
                    )
                    while (
                        exact_fit_count - lower_count > 1
                        and not optimization_stats["render_budget_exhausted"]
                    ):
                        midpoint = (lower_count + exact_fit_count) // 2
                        midpoint_state = reduction_path[midpoint - 1]
                        rendered_midpoint = render_candidate(
                            midpoint_state["projects"],
                            midpoint_state["skills"],
                            attempt_type="exact_boundary",
                            change_applied={
                                "change_type": "cumulative_reduction",
                                "reduction_count": midpoint,
                                "last_change": midpoint_state["last_change"],
                            },
                            probe_candidate=True,
                            margin_profile=active_margin_profile,
                        )
                        optimization_stats["exact_render_count"] += 1
                        if rendered_midpoint.get("pdf_path") is None:
                            break
                        if int(rendered_midpoint["page_count"]) <= 1:
                            exact_fit_count = midpoint
                        else:
                            lower_count = midpoint
                        _delete_generated_output(
                            rendered_midpoint.get("docx_path"),
                            rendered_midpoint.get("pdf_path"),
                        )

                    if not optimization_stats["render_verification_failed"]:
                        boundary_before_projects = (
                            deepcopy(working_projects)
                            if exact_fit_count == 1
                            else deepcopy(
                                reduction_path[exact_fit_count - 2]["projects"]
                            )
                        )
                        boundary_before_skills = (
                            deepcopy(working_skills)
                            if exact_fit_count == 1
                            else deepcopy(
                                reduction_path[exact_fit_count - 2]["skills"]
                            )
                        )
                        boundary_before_changes = (
                            []
                            if exact_fit_count == 1
                            else deepcopy(
                                reduction_path[exact_fit_count - 2]["changes"]
                            )
                        )
                        local_candidates = build_reduction_candidates(
                            boundary_before_projects,
                            boundary_before_skills,
                        )
                        local_tiers = group_candidates_by_protection_tier(
                            local_candidates
                        )
                        local_candidates = (
                            local_tiers[0][1][:FIT_LOCAL_REFINEMENT_LIMIT]
                            if local_tiers
                            else []
                        )
                        _fit_log(
                            "reduction stage=local_refinement "
                            f"candidates={len(local_candidates)}"
                        )
                        local_rendered = render_candidates_batch(
                            [
                                {
                                    "projects": candidate["projects"],
                                    "skills": candidate["skills"],
                                    "attempt_type": "local_refinement",
                                    "change_applied": candidate["change"],
                                    "probe_candidate": True,
                                    "quality_loss": candidate["quality_loss"],
                                    "margin_profile": active_margin_profile,
                                }
                                for candidate in local_candidates
                            ],
                            batch_label="local_refinement",
                        )
                        optimization_stats["local_refinement_render_count"] += len(
                            local_rendered
                        )
                        baseline_overflow = _rendered_overflow_value(
                            current_render
                        )
                        local_results: list[dict[str, Any]] = []
                        for candidate, rendered in zip(
                            local_candidates,
                            local_rendered,
                        ):
                            if rendered.get("pdf_path") is None:
                                continue
                            candidate_overflow = _rendered_overflow_value(rendered)
                            space_saved = max(
                                0.0,
                                baseline_overflow - candidate_overflow,
                            )
                            local_results.append(
                                {
                                    **candidate,
                                    "rendered": rendered,
                                    "space_saved_ratio": space_saved,
                                    "reaches_one_page": int(
                                        rendered["page_count"]
                                    ) <= 1,
                                }
                            )

                        selected = (
                            _choose_layout_aware_reduction(local_results)
                            if local_results
                            else None
                        )
                        if selected is not None and selected.get(
                            "reaches_one_page"
                        ):
                            working_projects = deepcopy(selected["projects"])
                            working_skills = deepcopy(selected["skills"])
                            active_changes = [
                                *boundary_before_changes,
                                deepcopy(selected["change"]),
                            ]
                        else:
                            selected_state = reduction_path[exact_fit_count - 1]
                            working_projects = deepcopy(selected_state["projects"])
                            working_skills = deepcopy(selected_state["skills"])
                            active_changes = deepcopy(selected_state["changes"])

                        for candidate in local_results:
                            _delete_generated_output(
                                candidate["rendered"].get("docx_path"),
                                candidate["rendered"].get("pdf_path"),
                            )
                        for rendered in rendered_coarse:
                            _delete_generated_output(
                                rendered.get("docx_path"),
                                rendered.get("pdf_path"),
                            )
                        final_render = render_candidate(
                            working_projects,
                            working_skills,
                            attempt_type="final_verification",
                            change_applied={
                                "change_type": "bounded_search_selected",
                                "reduction_count": len(active_changes),
                            },
                            margin_profile=active_margin_profile,
                        )
                        optimization_stats["exact_render_count"] += 1
                        if (
                            final_render.get("pdf_path") is not None
                            and int(final_render["page_count"]) <= 1
                        ):
                            _delete_generated_output(
                                current_render.get("docx_path"),
                                current_render.get("pdf_path"),
                            )
                            fitting_render = final_render
                            _fit_log(
                                "accepted reduction "
                                f"count={len(active_changes)} "
                                f"pages={final_render['page_count']}"
                            )

    if fitting_render is None:
        fit_status = (
            "verification_failed"
            if optimization_stats["render_verification_failed"]
            else (
                "search_exhausted"
                if optimization_stats["render_budget_exhausted"]
                else "unable_to_fit"
            )
        )
        _fit_log(
            "complete "
            f"status={fit_status} pages={current_render.get('page_count')} "
            f"renders={optimization_stats['candidate_states_rendered']} "
            f"elapsed={time.perf_counter() - fitting_started_at:.3f}s"
        )
        return {
            "generation_id": generation_id,
            "fitting_version": PHASE6C_FITTING_VERSION,
            "docx_path": current_render["docx_path"],
            "pdf_path": current_render["pdf_path"],
            "page_count": current_render["page_count"],
            "fit_one_page": False,
            "fit_status": fit_status,
            "attempts": attempt_logs,
            "tailored_projects_used": working_projects,
            "tailored_skills_used": working_skills,
            "page_fill_ratio": None,
            "estimated_unused_page_ratio": None,
            "page_density_mode": density_mode,
            "density_target_max": density_max_fill,
            "margin_profile": active_margin_profile,
            "margin_compaction_used": margin_compaction_used,
            "project_header_layout_requested": requested_project_header_layout,
            "project_header_layout_used": active_project_header_layout,
            "project_metadata_style": active_project_metadata_style,
            "project_header_compaction_used": project_header_compaction_used,
            **optimization_summary(),
            "note": (
                (
                    "Resume still exceeds one page while preserving the locked "
                    "Projects or Skills sections. Unlock a section, reduce "
                    "spacing, or allow more than one page."
                )
                if lock_policy["lock_projects"] or lock_policy["lock_skills"]
                else (
                    "Fit search reached its deterministic render budget before "
                    "one-page output was verified."
                    if optimization_stats["render_budget_exhausted"]
                    else (
                    "Resume still exceeds one page after all allowed whole-resume "
                    "reductions. Reduce spacing, lower the minimum Skills count, "
                    "or allow more than one page."
                    )
                )
            ),
        }

    best_projects = deepcopy(working_projects)
    best_skills = deepcopy(working_skills)
    best_render = fitting_render
    restored_change_count = 0
    permanently_rejected_restorations: set[str] = set()

    while (
        active_changes
        and density_max_fill is not None
        and optimization_stats["restoration_candidates_rendered"]
        < FIT_RESTORATION_RENDER_BUDGET
    ):
        candidate_results: list[
            dict[str, Any]
        ] = []
        best_fill = float(
            (
                best_render.get(
                    "fill_metrics",
                    {},
                )
                or {}
            ).get("page_fill_ratio")
            or 0.0
        )
        restoration_prepared: list[
            dict[str, Any]
        ] = []

        for change_index in (
            _whole_resume_restorable_change_indices(
                active_changes
            )
        ):
            source_change = active_changes[
                change_index
            ]
            change_key = _change_identity(
                source_change
            )

            if (
                change_key
                in permanently_rejected_restorations
            ):
                continue

            (
                restored_projects,
                restored_skills,
                restored,
                restore_info,
            ) = _restore_whole_resume_change(
                best_projects,
                best_skills,
                source_change,
            )

            if not restored:
                continue

            quality_gain = (
                _whole_resume_restoration_quality_gain(
                    source_change
                )
            )
            restoration_prepared.append(
                {
                    "change_index": change_index,
                    "change_key": change_key,
                    "projects": restored_projects,
                    "skills": restored_skills,
                    "quality_gain": quality_gain,
                    "attempt_type": restore_info.get(
                        "change_type",
                        "restore_content",
                    ),
                    "change_applied": restore_info,
                    "restoration_candidate": True,
                    "restoration_quality_gain": (
                        quality_gain
                    ),
                    "probe_candidate": True,
                }
            )

        if not restoration_prepared:
            break

        restoration_remaining = max(
            0,
            FIT_RESTORATION_RENDER_BUDGET
            - optimization_stats["restoration_candidates_rendered"],
        )
        restoration_prepared = sorted(
            restoration_prepared,
            key=lambda candidate: (
                -int(candidate["quality_gain"]),
                int(candidate["change_index"]),
            ),
        )[:restoration_remaining]
        if not restoration_prepared:
            break

        _fit_log(
            "restoration "
            f"candidates={len(restoration_prepared)} "
            f"remaining_budget={restoration_remaining}"
        )

        optimization_stats[
            "restoration_batch_count"
        ] += 1
        rendered_restorations = (
            render_candidates_batch(
                restoration_prepared,
                batch_label="restoration_pass",
            )
        )

        for prepared, rendered in zip(
            restoration_prepared,
            rendered_restorations,
        ):
            optimization_stats[
                "restoration_candidates_rendered"
            ] += 1

            if rendered["pdf_path"] is None:
                _delete_generated_output(
                    rendered["docx_path"],
                    None,
                )
                continue

            entry = rendered["attempt_entry"]
            change_key = str(
                prepared["change_key"]
            )

            if int(rendered["page_count"]) > 1:
                entry[
                    "restoration_accepted"
                ] = False
                entry["rejection_reason"] = (
                    "Restoring this content caused "
                    "the resume to exceed one page."
                )
                permanently_rejected_restorations.add(
                    change_key
                )
                _delete_generated_output(
                    rendered["docx_path"],
                    rendered["pdf_path"],
                )
                entry[
                    "temporary_output_deleted"
                ] = True
                continue

            candidate_fill = float(
                (
                    rendered.get(
                        "fill_metrics",
                        {},
                    )
                    or {}
                ).get("page_fill_ratio")
                or 0.0
            )

            if candidate_fill > density_max_fill:
                entry[
                    "restoration_accepted"
                ] = False
                entry["rejection_reason"] = (
                    "Restoration exceeded the "
                    "selected page-density limit."
                )
                entry[
                    "density_target_max"
                ] = density_max_fill
                permanently_rejected_restorations.add(
                    change_key
                )
                _delete_generated_output(
                    rendered["docx_path"],
                    rendered["pdf_path"],
                )
                entry[
                    "temporary_output_deleted"
                ] = True
                continue

            space_consumed = max(
                0.0,
                candidate_fill - best_fill,
            )
            restoration_efficiency = (
                float(prepared["quality_gain"])
                / max(
                    space_consumed,
                    _LAYOUT_EFFECT_THRESHOLD,
                )
            )
            entry[
                "space_consumed_ratio"
            ] = round(space_consumed, 3)
            entry[
                "restoration_efficiency_score"
            ] = round(
                restoration_efficiency,
                2,
            )

            candidate_results.append(
                {
                    "change_index": prepared[
                        "change_index"
                    ],
                    "change_key": change_key,
                    "projects": prepared[
                        "projects"
                    ],
                    "skills": prepared["skills"],
                    "rendered": rendered,
                    "quality_gain": prepared[
                        "quality_gain"
                    ],
                    "space_consumed_ratio": (
                        space_consumed
                    ),
                    "restoration_efficiency_score": (
                        restoration_efficiency
                    ),
                }
            )

        if not candidate_results:
            break

        chosen = max(
            candidate_results,
            key=lambda candidate: (
                float(
                    candidate[
                        "restoration_efficiency_score"
                    ]
                ),
                int(candidate["quality_gain"]),
            ),
        )

        for candidate in candidate_results:
            entry = candidate[
                "rendered"
            ]["attempt_entry"]
            entry["probe_selected"] = (
                candidate is chosen
            )

            if candidate is chosen:
                entry[
                    "restoration_accepted"
                ] = True
            else:
                entry[
                    "restoration_accepted"
                ] = False
                entry["rejection_reason"] = (
                    "Another one-page restoration "
                    "recovered more value per unit "
                    "of rendered space."
                )
                _delete_generated_output(
                    candidate[
                        "rendered"
                    ]["docx_path"],
                    candidate[
                        "rendered"
                    ]["pdf_path"],
                )
                entry[
                    "temporary_output_deleted"
                ] = True

        _delete_generated_output(
            best_render["docx_path"],
            best_render["pdf_path"],
        )
        best_render[
            "attempt_entry"
        ][
            "superseded_output_deleted"
        ] = True

        best_projects = deepcopy(
            chosen["projects"]
        )
        best_skills = deepcopy(
            chosen["skills"]
        )
        best_render = chosen["rendered"]
        active_changes.pop(
            int(chosen["change_index"])
        )
        restored_change_count += 1

    fill_metrics = best_render.get("fill_metrics", {}) or {}
    first_attempt_was_full_fit = (
        len(attempt_logs) == 1
        and not active_changes
        and restored_change_count == 0
    )

    if first_attempt_was_full_fit:
        note = "Generated resume fits within one page using the full tailored sections."
    elif restored_change_count > 0:
        note = (
            "Generated resume fits within one page after layout-aware fitting "
            "and a restoration pass recovered the strongest content that fit "
            "within the selected density target."
        )
    elif project_header_compaction_used and not active_changes:
        note = (
            "Generated resume fits within one page after Auto switched project "
            "headers from stacked to inline; no résumé evidence was removed."
        )
    elif margin_compaction_used and not active_changes:
        note = (
            "Generated resume fits within one page after safe margin compaction; "
            "no résumé evidence was removed to achieve the one-page fit."
        )
    elif density_mode == "none" and active_changes:
        note = (
            "Generated resume fits within one page. Fit only was selected, so "
            "the fitter kept the first evidence-safe one-page "
            "result instead of restoring content to fill spare space."
        )
    else:
        note = (
            "Generated resume fits within one page after comparing actual "
            "rendered space saved against deterministic evidence loss."
        )

    remaining_quality_loss = sum(
        _skill_reduction_quality_loss(change)
        if str(change.get("section", "projects")) == "skills"
        else _project_reduction_quality_loss(change)
        for change in active_changes
    )

    _fit_log(
        "complete "
        f"pages={best_render['page_count']} "
        f"bullets={_candidate_counts(best_projects, best_skills)['bullet_count']} "
        f"renders={optimization_stats['candidate_states_rendered']} "
        f"elapsed={time.perf_counter() - fitting_started_at:.3f}s"
    )

    return {
        "generation_id": generation_id,
            "fitting_version": PHASE6C_FITTING_VERSION,
        "docx_path": best_render["docx_path"],
        "pdf_path": best_render["pdf_path"],
        "page_count": best_render["page_count"],
        "fit_one_page": True,
        "fit_status": "verified_one_page",
        "attempts": attempt_logs,
        "tailored_projects_used": best_projects,
        "tailored_skills_used": best_skills,
        "page_fill_ratio": fill_metrics.get("page_fill_ratio"),
        "estimated_unused_page_ratio": fill_metrics.get(
            "estimated_unused_page_ratio"
        ),
        "page_fill_measurement_method": fill_metrics.get("measurement_method"),
        "page_density_mode": density_mode,
        "density_target_max": density_max_fill,
        "margin_profile": active_margin_profile,
        "margin_compaction_used": margin_compaction_used,
        "project_header_layout_requested": requested_project_header_layout,
        "project_header_layout_used": active_project_header_layout,
        "project_metadata_style": active_project_metadata_style,
        "project_header_compaction_used": project_header_compaction_used,
        **optimization_summary(),
        "fitting_objective": (
            "Protect unique requirement evidence, then minimise deterministic "
            "evidence loss per unit of actual rendered space saved; "
            "do not rerun analysis or project selection."
        ),
        "remaining_quality_loss": remaining_quality_loss,
        "restored_change_count": restored_change_count,
        "remaining_active_reductions": active_changes,
        "permanently_rejected_restoration_count": len(
            permanently_rejected_restorations
        ),
        "note": note,
    }



# def generate_tailored_resume_copy_fit_one_page(
#     *,
#     saved_resume_docx_path: str | Path,
#     tailored_projects: dict[str, Any] | None = None,
#     tailored_skills: dict[str, Any] | None = None,
#     application_id: int | None = None,
#     max_projects: int = 3,
#     max_bullets_per_project: int = 3,
#     max_attempts: int = 5,
#     spacing_mode: str = "paragraph_spacing",
#     project_spacing_pt: int = 10,
#     after_projects_spacing_pt: int = 10,
#     blank_lines_between_projects: int = 1,
#     blank_lines_after_projects: int = 1,
#     add_spacing_before_first_project: bool = False,
# ) -> dict[str, Any]:
#     """
#     Generate a tailored DOCX and keep it to one page if possible.

#     Flow:
#     1. Generate full truthful tailored version first.
#     2. Convert to PDF and count pages.
#     3. If it fits, keep it.
#     4. If it exceeds one page, compact and retry.
#     """
#     if not tailored_projects and not tailored_skills:
#         raise ValueError("Generate a tailored Projects section or Skills section first.")

#     cleanup_old_tailored_outputs_for_application(application_id)

#     attempt_logs = []
#     last_docx_path = None
#     last_pdf_path = None
#     last_page_count = None
#     last_projects_used = deepcopy(tailored_projects) if tailored_projects else None

#     attempt_limit = max_attempts if tailored_projects else 1

#     for attempt_index in range(attempt_limit):
#         if attempt_index == 0:
#             working_projects = deepcopy(tailored_projects) if tailored_projects else None
#             attempt_type = "full"
#         else:
#             working_projects = compact_tailored_projects_for_space(
#                 tailored_projects,
#                 attempt=attempt_index,
#             )
#             attempt_type = f"compact_{attempt_index}"

#         docx_path = generate_tailored_resume_copy(
#             saved_resume_docx_path=saved_resume_docx_path,
#             tailored_projects=working_projects,
#             tailored_skills=tailored_skills,
#             application_id=application_id,
#             max_projects=max_projects,
#             max_bullets_per_project=max_bullets_per_project,
#             spacing_mode=spacing_mode,
#             project_spacing_pt=project_spacing_pt,
#             after_projects_spacing_pt=after_projects_spacing_pt,
#             blank_lines_between_projects=blank_lines_between_projects,
#             blank_lines_after_projects=blank_lines_after_projects,
#             add_spacing_before_first_project=add_spacing_before_first_project,
#         )

#         last_docx_path = docx_path
#         last_projects_used = working_projects

#         pdf_path = convert_docx_to_pdf_if_possible(docx_path)

#         if pdf_path is None:
#             return {
#                 "docx_path": docx_path,
#                 "pdf_path": None,
#                 "page_count": None,
#                 "fit_one_page": None,
#                 "attempts": attempt_logs,
#                 "tailored_projects_used": working_projects,
#                 "note": (
#                     "Could not check page count because LibreOffice is unavailable "
#                     "or DOCX-to-PDF conversion failed. DOCX generation still worked."
#                 ),
#             }

#         page_count = count_pdf_pages(pdf_path)

#         last_pdf_path = pdf_path
#         last_page_count = page_count

#         attempt_logs.append(
#             {
#                 "attempt": attempt_index + 1,
#                 "attempt_type": attempt_type,
#                 "docx_path": str(docx_path),
#                 "pdf_path": str(pdf_path),
#                 "page_count": page_count,
#                 "project_count": len(working_projects.get("recommended_projects", [])) if working_projects else 0,
#                 "bullet_count": sum(
#                     len(project.get("draft_bullets", []) or [])
#                     for project in working_projects.get("recommended_projects", [])
#                 ) if working_projects else 0,
#             }
#         )

#         if page_count <= 1:
#             return {
#                 "docx_path": docx_path,
#                 "pdf_path": pdf_path,
#                 "page_count": page_count,
#                 "fit_one_page": True,
#                 "attempts": attempt_logs,
#                 "tailored_projects_used": working_projects,
#                 "note": (
#                     "Generated resume fits within one page."
#                     if attempt_index == 0
#                     else "Generated resume fits within one page after compacting."
#                 ),
#             }

#     return {
#         "docx_path": last_docx_path,
#         "pdf_path": last_pdf_path,
#         "page_count": last_page_count,
#         "fit_one_page": False,
#         "attempts": attempt_logs,
#         "tailored_projects_used": last_projects_used,
#         "note": (
#             "Resume still appears to exceed one page after compacting. "
#             "Try fewer projects, fewer bullets, or shorter skills."
#         ),
#     }





# ---------------------------------------------------------------------------
# Preview helpers
# ---------------------------------------------------------------------------

def extract_docx_preview_text(docx_path: str | Path) -> str:
    """
    Return a simple text preview of the generated DOCX.

    This is not a visual Word preview, but it works without LibreOffice.
    """
    document = Document(docx_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    return "\n".join(lines)


# def _find_libreoffice_executable() -> str | None:
#     """Find LibreOffice executable on Windows/Linux/Mac."""
#     soffice = shutil.which("soffice") or shutil.which("libreoffice")

#     if soffice:
#         return soffice

#     common_windows_paths = [
#         r"C:\Program Files\LibreOffice\program\soffice.exe",
#         r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
#     ]

#     for path in common_windows_paths:
#         if Path(path).exists():
#             return path

#     return None


# def convert_docx_to_pdf_if_possible(docx_path: str | Path) -> Path | None:
#     """
#     Convert DOCX to PDF using LibreOffice if it is installed.

#     Returns:
#         PDF path, or None if LibreOffice is unavailable or conversion fails.
#     """
#     docx_path = Path(docx_path)
#     PREVIEW_DIR.mkdir(parents=True, exist_ok=True)

#     # soffice = shutil.which("soffice") or shutil.which("libreoffice")
#     soffice = _find_libreoffice_executable()
#     if not soffice:
#         return None

#     try:
#         subprocess.run(
#             [
#                 soffice,
#                 "--headless",
#                 "--convert-to",
#                 "pdf",
#                 "--outdir",
#                 str(PREVIEW_DIR),
#                 str(docx_path),
#             ],
#             check=True,
#             stdout=subprocess.PIPE,
#             stderr=subprocess.PIPE,
#             timeout=60,
#             text=True,
#         )
#     except Exception as exc:
#         print(f"[LibreOffice conversion failed] {exc}")
#         return None

#     pdf_path = PREVIEW_DIR / f"{docx_path.stem}.pdf"

#     if pdf_path.exists():
#         return pdf_path

#     return None


def _find_libreoffice_executable() -> str | None:
    """Find LibreOffice executable on Windows/Linux/Mac."""

    # Prefer the normal Windows install path first.
    # This avoids accidentally picking soffice.COM from PATH.
    common_windows_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
    ]

    for path in common_windows_paths:
        if Path(path).exists():
            return path

    return shutil.which("soffice") or shutil.which("libreoffice")


def convert_docx_to_pdf_if_possible(docx_path: str | Path) -> Path | None:
    """
    Convert DOCX to PDF using LibreOffice if it is installed.

    Returns:
        PDF path, or None if LibreOffice is unavailable or conversion fails.
    """
    docx_path = Path(docx_path).resolve()
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    preview_dir = PREVIEW_DIR.resolve()

    soffice = _find_libreoffice_executable()
    if not soffice:
        print("[LibreOffice conversion failed] LibreOffice executable not found.")
        return None

    expected_pdf_path = preview_dir / f"{docx_path.stem}.pdf"

    # Remove old preview with the same name so we know whether this conversion produced a new file.
    if expected_pdf_path.exists():
        try:
            expected_pdf_path.unlink()
        except OSError:
            pass

    # Give each headless conversion its own LibreOffice profile.
    # This reduces collisions with normal LibreOffice sessions and
    # with the fitter's rapid sequence of conversion attempts.
    lo_profile_dir = (
        preview_dir
        / "lo_profiles"
        / uuid.uuid4().hex
    ).resolve()
    lo_profile_dir.mkdir(parents=True, exist_ok=True)
    lo_profile_uri = lo_profile_dir.as_uri()

    command = [
        soffice,
        "--headless",
        "--nologo",
        "--nodefault",
        "--nofirststartwizard",
        "--nolockcheck",
        f"-env:UserInstallation={lo_profile_uri}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(preview_dir),
        str(docx_path),
    ]

    # Do not pass Python installation overrides into LibreOffice.
    # They can produce warnings such as:
    # "Could not find platform independent libraries <prefix>".
    conversion_env = os.environ.copy()
    conversion_env.pop("PYTHONHOME", None)
    conversion_env.pop("PYTHONPATH", None)

    try:
        result = subprocess.run(
            command,
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=LIBREOFFICE_SINGLE_CONVERSION_TIMEOUT_SECONDS,
            text=True,
            env=conversion_env,
        )

        if result.returncode != 0:
            print("[LibreOffice conversion failed]")
            print("Command:", command)
            print("Return code:", result.returncode)
            print("STDOUT:", result.stdout)
            print("STDERR:", result.stderr)
            return None

        # On Windows, LibreOffice can return just before the PDF is
        # visible to Python. Poll briefly before declaring failure.
        for _ in range(20):
            try:
                if (
                    expected_pdf_path.exists()
                    and expected_pdf_path.stat().st_size > 0
                ):
                    return expected_pdf_path
            except OSError:
                pass

            time.sleep(0.1)

        print(
            "[LibreOffice conversion failed] "
            "Command succeeded but PDF was not created "
            "after waiting 2 seconds."
        )
        print("Expected PDF:", expected_pdf_path)
        print("STDOUT:", result.stdout)
        print("STDERR:", result.stderr)
        return None

    except subprocess.TimeoutExpired:
        print(
            "[FIT] LibreOffice conversion timed out "
            f"after {LIBREOFFICE_SINGLE_CONVERSION_TIMEOUT_SECONDS}s.",
            flush=True,
        )
        return None
    except Exception as exc:
        print(f"[LibreOffice conversion crashed] {exc}")
        return None

    finally:
        try:
            shutil.rmtree(
                lo_profile_dir,
                ignore_errors=True,
            )
        except OSError:
            pass


def convert_docx_batch_to_pdf_if_possible(
    docx_paths: list[str | Path],
) -> tuple[dict[str, Path | None], dict[str, Any]]:
    """
    Convert several DOCX files with one LibreOffice process.

    Missing outputs are retried individually through the existing converter.
    The returned dictionary is keyed by each resolved DOCX path.
    """
    resolved_paths: list[Path] = []
    seen: set[str] = set()

    for raw_path in docx_paths:
        path = Path(raw_path).resolve()
        key = str(path)
        if key in seen:
            continue
        seen.add(key)
        resolved_paths.append(path)

    diagnostics: dict[str, Any] = {
        "requested_count": len(resolved_paths),
        "batch_process_count": 0,
        "fallback_process_count": 0,
        "missing_source_count": 0,
        "batch_return_code": None,
        "timed_out": False,
        "timeout_seconds": None,
    }
    results: dict[str, Path | None] = {
        str(path): None
        for path in resolved_paths
    }

    valid_paths: list[Path] = []
    for path in resolved_paths:
        try:
            if path.exists() and path.stat().st_size > 0:
                valid_paths.append(path)
            else:
                diagnostics["missing_source_count"] += 1
        except OSError:
            diagnostics["missing_source_count"] += 1

    if not valid_paths:
        return results, diagnostics

    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    preview_dir = PREVIEW_DIR.resolve()
    soffice = _find_libreoffice_executable()

    if not soffice:
        print(
            "[LibreOffice batch conversion failed] "
            "LibreOffice executable not found."
        )
        return results, diagnostics

    expected_paths = {
        str(path): preview_dir / f"{path.stem}.pdf"
        for path in valid_paths
    }

    for expected_path in expected_paths.values():
        try:
            expected_path.unlink()
        except FileNotFoundError:
            pass
        except OSError:
            pass

    lo_profile_dir = (
        preview_dir
        / "lo_profiles"
        / uuid.uuid4().hex
    ).resolve()
    lo_profile_dir.mkdir(parents=True, exist_ok=True)

    command = [
        soffice,
        "--headless",
        "--nologo",
        "--nodefault",
        "--nofirststartwizard",
        "--nolockcheck",
        f"-env:UserInstallation={lo_profile_dir.as_uri()}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(preview_dir),
        *[str(path) for path in valid_paths],
    ]

    conversion_env = os.environ.copy()
    conversion_env.pop("PYTHONHOME", None)
    conversion_env.pop("PYTHONPATH", None)

    try:
        diagnostics["batch_process_count"] = 1
        timeout_seconds = max(
            LIBREOFFICE_BATCH_BASE_TIMEOUT_SECONDS,
            min(
                LIBREOFFICE_BATCH_MAX_TIMEOUT_SECONDS,
                LIBREOFFICE_BATCH_BASE_TIMEOUT_SECONDS
                + len(valid_paths) * LIBREOFFICE_BATCH_TIMEOUT_PER_CANDIDATE_SECONDS,
            ),
        )
        diagnostics["timeout_seconds"] = timeout_seconds
        result = subprocess.run(
            command,
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout_seconds,
            text=True,
            env=conversion_env,
        )
        diagnostics["batch_return_code"] = result.returncode

        if result.returncode != 0:
            print("[LibreOffice batch conversion failed]")
            print("Command:", command)
            print("Return code:", result.returncode)
            print("STDOUT:", result.stdout)
            print("STDERR:", result.stderr)

        for _ in range(40):
            pending = []
            for source_path in valid_paths:
                expected_path = expected_paths[str(source_path)]
                try:
                    ready = (
                        expected_path.exists()
                        and expected_path.stat().st_size > 0
                    )
                except OSError:
                    ready = False

                if ready:
                    results[str(source_path)] = expected_path
                else:
                    pending.append(source_path)

            if not pending:
                break
            time.sleep(0.1)

    except subprocess.TimeoutExpired:
        diagnostics["timed_out"] = True
        print(
            "[FIT] LibreOffice batch conversion timed out "
            f"after {diagnostics['timeout_seconds']}s.",
            flush=True,
        )
    except Exception as exc:
        print(f"[LibreOffice batch conversion crashed] {exc}")

    finally:
        try:
            shutil.rmtree(
                lo_profile_dir,
                ignore_errors=True,
            )
        except OSError:
            pass

    if diagnostics["timed_out"]:
        # Do not convert a timed-out batch one file at a time.  The batch did
        # not provide rendered verification, so callers must fail closed.
        return results, diagnostics

    # Retry only files that the batch did not produce. This keeps the existing
    # robust individual conversion path as a fallback.
    for source_path in valid_paths:
        if results[str(source_path)] is not None:
            continue

        diagnostics["fallback_process_count"] += 1
        results[str(source_path)] = (
            convert_docx_to_pdf_if_possible(source_path)
        )

    return results, diagnostics


def pdf_to_preview_pngs(
    pdf_path: str | Path,
    *,
    zoom: float = 1.35,
    max_pages: int = 5,
) -> list[bytes]:
    """Render PDF pages to opaque PNG bytes for reliable Streamlit preview.

    This avoids browser-native PDF viewer differences (including dark/black
    canvases in some browser/theme combinations). PDF generation and page-count
    logic remain unchanged.
    """
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError(
            "PyMuPDF is required for rendered PDF previews."
        ) from exc

    path = Path(pdf_path)
    if not path.is_file():
        raise FileNotFoundError(
            f"PDF preview file does not exist: {path}"
        )

    safe_zoom = min(4.0, max(0.5, float(zoom)))
    safe_max_pages = min(20, max(1, int(max_pages)))
    rendered: list[bytes] = []

    with fitz.open(str(path)) as document:
        for page_index in range(
            min(len(document), safe_max_pages)
        ):
            page = document[page_index]
            pixmap = page.get_pixmap(
                matrix=fitz.Matrix(safe_zoom, safe_zoom),
                alpha=False,
            )
            rendered.append(pixmap.tobytes("png"))

    return rendered


def pdf_to_preview_html(
    pdf_path: str | Path,
    *,
    max_width: int = 820,
    max_pages: int = 5,
    zoom: float = 1.35,
    include_download: bool = True,
) -> str:
    """Return a centered rasterized PDF preview and optional PDF download."""
    import html

    path = Path(pdf_path)
    if not path.is_file():
        raise FileNotFoundError(
            f"PDF preview file does not exist: {path}"
        )

    safe_width = min(1400, max(360, int(max_width)))
    pages = pdf_to_preview_pngs(
        path,
        zoom=zoom,
        max_pages=max_pages,
    )
    if not pages:
        raise RuntimeError("PDF preview renderer returned no pages.")

    parts = [
        '<div style="display:flex;flex-direction:column;'
        'align-items:center;gap:14px;width:100%;">'
    ]
    for page_index, page_png in enumerate(pages, start=1):
        encoded_png = base64.b64encode(page_png).decode("ascii")
        parts.append(
            '<div style="width:100%;display:flex;'
            'justify-content:center;">'
            f'<img src="data:image/png;base64,{encoded_png}" '
            f'alt="PDF preview page {page_index}" '
            'style="display:block;width:100%;'
            f'max-width:{safe_width}px;height:auto;background:white;'
            'border-radius:8px;" />'
            '</div>'
        )

    if include_download:
        pdf_b64 = base64.b64encode(path.read_bytes()).decode("ascii")
        safe_name = html.escape(path.name, quote=True)
        parts.append(
            f'<a href="data:application/pdf;base64,{pdf_b64}" '
            f'download="{safe_name}" '
            'style="display:inline-block;padding:8px 14px;'
            'border:1px solid rgba(128,128,128,.45);'
            'border-radius:8px;text-decoration:none;font-weight:600;">'
            'Download PDF</a>'
        )

    parts.append("</div>")
    return "".join(parts)


def pdf_to_iframe_html(
    pdf_path: str | Path,
    *,
    height: int = 800,
) -> str:
    """Compatibility wrapper returning the safe rasterized PDF preview."""
    _ = height
    return pdf_to_preview_html(pdf_path)





def cleanup_stale_libreoffice_profiles(
    *,
    max_age_hours: int = 24,
) -> dict[str, Any]:
    """
    Remove abandoned LibreOffice profile folders only when they are old.

    Recent profiles are left alone because another conversion could still be
    using them. The current converter normally removes its unique profile in a
    ``finally`` block, so this is only a fallback for interrupted processes.
    """
    max_age_seconds = max(1, int(max_age_hours)) * 3600
    now = time.time()
    removed: list[str] = []
    failed: list[dict[str, str]] = []

    candidates: list[Path] = []

    legacy_root = PREVIEW_DIR / "lo_profile"
    if legacy_root.exists():
        candidates.append(legacy_root)

    profiles_root = PREVIEW_DIR / "lo_profiles"
    if profiles_root.exists():
        candidates.extend(
            child
            for child in profiles_root.iterdir()
            if child.is_dir()
        )

    for candidate in candidates:
        try:
            mtimes = [candidate.stat().st_mtime]
            mtimes.extend(
                path.stat().st_mtime
                for path in candidate.rglob("*")
                if path.exists()
            )
            newest_mtime = max(mtimes)
        except OSError as exc:
            failed.append(
                {
                    "path": str(candidate),
                    "error": str(exc),
                }
            )
            continue

        if now - newest_mtime < max_age_seconds:
            continue

        try:
            shutil.rmtree(candidate, ignore_errors=False)
            removed.append(str(candidate))
        except OSError as exc:
            failed.append(
                {
                    "path": str(candidate),
                    "error": str(exc),
                }
            )

    if profiles_root.exists():
        try:
            if not any(profiles_root.iterdir()):
                profiles_root.rmdir()
                removed.append(str(profiles_root))
        except OSError:
            pass

    return {
        "removed_profile_directories": removed,
        "failed_files": failed,
    }


def cleanup_application_resume_files(
    application_id: int | None,
    *,
    delete_saved_resume: bool = True,
    delete_generated_outputs: bool = True,
    delete_libreoffice_profiles: bool = True,
) -> dict[str, Any]:
    """
    Delete app-owned résumé files for one application session.

    Safety:
    - Only files beginning with ``app_{application_id}_`` are considered.
    - The function never scans outside the three known application folders.
    - Files already downloaded by the user elsewhere are unaffected.
    """
    if application_id is None:
        return {
            "application_id": None,
            "deleted_file_count": 0,
            "deleted_files": [],
            "failed_files": [],
        }

    try:
        app_id = int(application_id)
    except (TypeError, ValueError):
        raise ValueError("application_id must be an integer.")

    if app_id < 0:
        raise ValueError("application_id must be non-negative.")

    patterns: list[Path] = []

    if delete_saved_resume:
        patterns.append(
            SAVED_RESUME_DIR / f"app_{app_id}_*.docx"
        )

    if delete_generated_outputs:
        patterns.extend(
            [
                TAILORED_RESUME_DIR
                / f"app_{app_id}_tailored_resume_*.docx",
                PREVIEW_DIR
                / f"app_{app_id}_tailored_resume_*.pdf",
                PREVIEW_DIR
                / f"app_{app_id}_tailored_resume_*.png",
            ]
        )

    deleted_files: list[str] = []
    failed_files: list[dict[str, str]] = []

    for pattern in patterns:
        for path in pattern.parent.glob(pattern.name):
            try:
                if path.is_file():
                    path.unlink()
                    deleted_files.append(str(path))
            except OSError as exc:
                failed_files.append(
                    {
                        "path": str(path),
                        "error": str(exc),
                    }
                )

    removed_profile_directories: list[str] = []

    if delete_libreoffice_profiles:
        profile_cleanup = cleanup_stale_libreoffice_profiles(
            max_age_hours=24,
        )
        removed_profile_directories = list(
            profile_cleanup.get(
                "removed_profile_directories",
                [],
            )
        )
        failed_files.extend(
            profile_cleanup.get("failed_files", [])
        )

    return {
        "application_id": app_id,
        "deleted_file_count": len(deleted_files),
        "deleted_files": deleted_files,
        "removed_profile_directories": (
            removed_profile_directories
        ),
        "failed_files": failed_files,
    }


# ---------------------------------------------------------------------------
# Cleaner
# ---------------------------------------------------------------------------

def cleanup_old_tailored_outputs_for_application(application_id: int | None) -> None:
    """
    Remove old generated DOCX/PDF/PNG outputs for this application session.

    Does not remove data/saved_resumes, because that is the original saved DOCX copy.
    """
    if application_id is None:
        return

    patterns = [
        TAILORED_RESUME_DIR / f"app_{application_id}_tailored_resume_*.docx",
        PREVIEW_DIR / f"app_{application_id}_tailored_resume_*.pdf",
        PREVIEW_DIR / f"app_{application_id}_tailored_resume_*.png",
    ]

    for pattern in patterns:
        for path in pattern.parent.glob(pattern.name):
            try:
                path.unlink()
            except OSError:
                pass
