"""Deterministic DOCX materialisation from an immutable résumé text snapshot."""

from __future__ import annotations

from pathlib import Path
from datetime import datetime
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.shared import Inches, Pt


IMMUTABLE_SNAPSHOT_DOCX_VERSION = "immutable-snapshot-docx-v1"


def _normalise_docx_container(path: Path) -> None:
    """Make ZIP metadata deterministic without changing document content."""
    with ZipFile(path, "r") as source:
        entries = [
            (item.filename, source.read(item.filename), item.external_attr)
            for item in source.infolist()
        ]
    normalized = path.with_suffix(".normalized.docx")
    with ZipFile(normalized, "w", compression=ZIP_DEFLATED) as target:
        for name, data, external_attr in sorted(entries):
            info = ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.external_attr = external_attr
            info.create_system = 0
            target.writestr(info, data)
    normalized.replace(path)


def materialise_immutable_snapshot_docx(
    *,
    resume_text: str,
    output_path: str | Path,
) -> Path:
    """Render frozen text without claiming to reproduce original fitted bytes."""
    text = str(resume_text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not text.strip():
        raise ValueError("The immutable résumé text snapshot is empty.")

    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        if line.upper() == line and any(character.isalpha() for character in line):
            run.bold = True
            run.font.size = Pt(10)

    properties = document.core_properties
    properties.title = "Immutable Blueprint Snapshot"
    properties.subject = IMMUTABLE_SNAPSHOT_DOCX_VERSION
    properties.author = "Job AI Helper"
    properties.created = datetime(2000, 1, 1, 0, 0, 0)
    properties.modified = datetime(2000, 1, 1, 0, 0, 0)
    document.save(target)
    _normalise_docx_container(target)
    return target
