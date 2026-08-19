from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt

from database import user_profile_manager as profile_manager
from resume_builder.project_header_format import (
    build_project_title,
    format_project_metadata,
    inline_project_header,
    normalise_project_header_layout,
    normalise_project_metadata_style,
    split_legacy_project_title,
)
from resume_builder.docx_projects_skills_replacer import (
    apply_source_project_display_fallbacks,
    replace_projects_section,
)
from tailoring.project_section_tailor import (
    _evidence_item_to_candidate,
    build_project_candidate_pool,
)


class ResumeFormatStylePolicyTests(unittest.TestCase):
    def test_pipe_groups_are_semantic_not_baked_into_title(self):
        project = {
            "title": "CyberSphere",
            "resume_header_tools": ["Unity Engine", "C#"],
            "resume_header_context": ["Team of 2", "Published on Google Play"],
        }
        self.assertEqual(build_project_title(project), "CyberSphere")
        self.assertEqual(
            format_project_metadata(project, style="pipes"),
            "Unity Engine, C# | Team of 2 | Published on Google Play",
        )
        self.assertEqual(
            inline_project_header(project, style="pipes"),
            "CyberSphere | Unity Engine, C# | Team of 2 | Published on Google Play",
        )

    def test_parentheses_remain_legacy_option(self):
        project = {
            "title": "CyberSphere",
            "resume_header_tools": ["Unity Engine", "C#"],
            "resume_header_context": ["Team of 2"],
        }
        self.assertEqual(
            inline_project_header(project, style="parentheses"),
            "CyberSphere (Unity Engine, C#, Team of 2)",
        )

    def test_subtitle_is_separate_from_metadata(self):
        project = {
            "title": "RequestFlow",
            "subtitle": "IT Service Request Tracker",
            "resume_header_tools": ["React", "TypeScript", "FastAPI"],
            "resume_header_context": ["Individual Project"],
        }
        self.assertEqual(
            build_project_title(project),
            "RequestFlow — IT Service Request Tracker",
        )
        self.assertEqual(
            format_project_metadata(project, style="pipes"),
            "React, TypeScript, FastAPI | Individual Project",
        )

    def test_legacy_title_parser(self):
        parsed = split_legacy_project_title(
            "CyberSphere (Unity Engine, Team of 2, Published on Google Play)"
        )
        self.assertEqual(parsed["title"], "CyberSphere")
        self.assertEqual(parsed["resume_header_tools"], ["Unity Engine"])
        self.assertEqual(
            parsed["resume_header_context"],
            ["Team of 2", "Published on Google Play"],
        )
        self.assertTrue(parsed["legacy_metadata_found"])

    def test_defaults_fail_safe(self):
        self.assertEqual(normalise_project_header_layout("nonsense"), "auto")
        self.assertEqual(normalise_project_metadata_style("nonsense"), "pipes")

    def test_evidence_candidate_propagates_metadata(self):
        candidate = _evidence_item_to_candidate({
            "category": "Project",
            "title": "CyberSphere",
            "subtitle": "",
            "description": "Built gameplay systems.",
            "period": "Jan 2018 - Feb 2018",
            "skills": ["Gameplay scripting"],
            "tools": ["Unity Engine", "C#"],
            "resume_header_tools": ["Unity Engine", "C#"],
            "resume_header_context": ["Team of 2", "Published on Google Play"],
            "impact": "Published mobile game",
        })
        self.assertIsNotNone(candidate)
        assert candidate is not None
        self.assertEqual(candidate["title"], "CyberSphere")
        self.assertEqual(candidate["resume_header_tools"], ["Unity Engine", "C#"])

    def test_resume_subtitle_survives_empty_evidence_display_fields(self):
        pool = build_project_candidate_pool(
            resume_profile={
                "projects": [
                    {
                        "title": (
                            "Job AI Helper — AI-Powered Resume Tailoring System"
                        ),
                        "date": "May 2026 - Aug 2026",
                        "bullets": ["Built the application."],
                    }
                ]
            },
            evidence_items=[
                {
                    "category": "Project",
                    "title": "Job AI Helper",
                    "description": "Built the application.",
                    "period": "May 2026 - Aug 2026",
                    "tools": ["Python", "Streamlit", "OpenAI API"],
                    "resume_header_tools": [],
                    "resume_header_context": [],
                }
            ],
        )
        self.assertEqual(len(pool), 1)
        project = pool[0]
        self.assertEqual(project["title"], "Job AI Helper")
        self.assertEqual(
            project["subtitle"],
            "AI-Powered Resume Tailoring System",
        )
        self.assertEqual(
            project["display_title"],
            "Job AI Helper — AI-Powered Resume Tailoring System",
        )
        self.assertEqual(
            project["canonical_tools"],
            ["Python", "Streamlit", "OpenAI API"],
        )

    def test_source_docx_fallback_preserves_stacked_groups_and_metadata_style(self):
        with tempfile.TemporaryDirectory() as temporary:
            source_path = Path(temporary) / "base.docx"
            document = Document()

            document.add_paragraph("WORK EXPERIENCE")
            experience_title = document.add_paragraph()
            experience_title.add_run("Example Company — Software Engineer")
            experience_title.add_run("\tMay 2025 - Apr 2026")
            document.add_paragraph(
                "Project Context | Python, Docker | Team of 2"
            )
            document.add_paragraph("• Preserved work-experience bullet.")

            document.add_paragraph("PROJECTS")
            title_paragraph = document.add_paragraph()
            title_run = title_paragraph.add_run(
                "Job AI Helper — AI-Powered Resume Tailoring System"
            )
            title_run.bold = True
            title_paragraph.add_run("\tMay 2026 - Aug 2026")

            metadata_paragraph = document.add_paragraph()
            metadata_paragraph.paragraph_format.space_before = Pt(2)
            metadata_paragraph.paragraph_format.space_after = Pt(1)
            metadata_run = metadata_paragraph.add_run(
                "Python, Streamlit, OpenAI, SQLite, ChromaDB"
                " | GitHub Actions | Individual Project"
            )
            metadata_run.italic = False

            document.add_paragraph("• Original project bullet.")
            document.add_paragraph("SKILLS")
            document.add_paragraph("Programming: Python")
            document.save(source_path)

            tailored = {
                "recommended_projects": [
                    {
                        "title": "Job AI Helper",
                        "subtitle": "",
                        "display_title": "Job AI Helper",
                        "resume_header_tools": [],
                        "resume_header_context": [],
                        "canonical_tools": [
                            "Python",
                            "Streamlit",
                            "OpenAI API",
                            "SQLite",
                            "ChromaDB",
                            "LiteLLM",
                            "pypdf",
                            "python-docx",
                            "GitHub",
                            "GitHub Actions (CI)",
                        ],
                        "period": "May 2026 - Aug 2026",
                        "draft_bullets": ["Tailored project bullet."],
                    }
                ]
            }

            enriched = apply_source_project_display_fallbacks(
                source_path,
                tailored,
            )
            project = enriched["recommended_projects"][0]
            self.assertEqual(
                project["subtitle"],
                "AI-Powered Resume Tailoring System",
            )
            self.assertEqual(
                project["resume_header_tools"],
                ["Python", "Streamlit", "OpenAI", "SQLite", "ChromaDB"],
            )
            self.assertEqual(
                project["resume_header_context"],
                ["GitHub Actions", "Individual Project"],
            )

            rendered = Document(source_path)
            before_texts = [paragraph.text for paragraph in rendered.paragraphs]
            work_heading_index = before_texts.index("WORK EXPERIENCE")
            projects_heading_index = before_texts.index("PROJECTS")
            work_before = before_texts[
                work_heading_index + 1 : projects_heading_index
            ]

            replace_projects_section(
                rendered,
                enriched,
                max_projects=1,
                max_bullets_per_project=1,
                project_header_layout="stacked",
                project_metadata_style="pipes",
            )

            texts = [paragraph.text for paragraph in rendered.paragraphs]
            self.assertIn(
                "Job AI Helper — AI-Powered Resume Tailoring System"
                "\tMay 2026 - Aug 2026",
                texts,
            )
            expected_metadata = (
                "Python, Streamlit, OpenAI, SQLite, ChromaDB"
                " | GitHub Actions | Individual Project"
            )
            self.assertIn(expected_metadata, texts)
            self.assertNotIn(
                "Python, Streamlit, OpenAI API, SQLite, ChromaDB, "
                "LiteLLM, pypdf, python-docx, GitHub, GitHub Actions (CI)",
                texts,
            )

            metadata = next(
                paragraph
                for paragraph in rendered.paragraphs
                if paragraph.text == expected_metadata
            )
            self.assertEqual(metadata.paragraph_format.space_before, Pt(2))
            self.assertEqual(metadata.paragraph_format.space_after, Pt(1))
            self.assertFalse(bool(metadata.runs[0].italic))

            work_heading_index = texts.index("WORK EXPERIENCE")
            projects_heading_index = texts.index("PROJECTS")
            work_after = texts[
                work_heading_index + 1 : projects_heading_index
            ]
            self.assertEqual(work_after, work_before)

    def test_explicit_structured_metadata_wins_over_source_fallback(self):
        with tempfile.TemporaryDirectory() as temporary:
            source_path = Path(temporary) / "base.docx"
            document = Document()
            document.add_paragraph("PROJECTS")
            title = document.add_paragraph()
            title.add_run("CyberSphere — Shooter Game")
            title.add_run("\tJan 2018 - Feb 2018")
            document.add_paragraph(
                "C#, Unity Engine | Published on Google Play | Team of 2"
            )
            document.add_paragraph("• Original bullet.")
            document.add_paragraph("SKILLS")
            document.add_paragraph("Game & Engine: C#")
            document.save(source_path)

            tailored = {
                "recommended_projects": [
                    {
                        "title": "CyberSphere",
                        "subtitle": "Updated Subtitle",
                        "resume_header_tools": ["Unity Engine"],
                        "resume_header_context": ["Individual Project"],
                        "canonical_tools": ["Unity Engine", "C#"],
                        "period": "Jan 2018 - Feb 2018",
                        "draft_bullets": ["Tailored bullet."],
                    }
                ]
            }
            enriched = apply_source_project_display_fallbacks(
                source_path,
                tailored,
            )
            project = enriched["recommended_projects"][0]
            self.assertEqual(project["subtitle"], "Updated Subtitle")
            self.assertEqual(
                project["resume_header_tools"],
                ["Unity Engine"],
            )
            self.assertEqual(
                project["resume_header_context"],
                ["Individual Project"],
            )

    def test_database_migration_preserves_id_and_canonical_tools(self):
        original_db = profile_manager.DB_PATH
        with tempfile.TemporaryDirectory() as temporary:
            profile_manager.DB_PATH = Path(temporary) / "applications.db"
            try:
                profile_manager.init_user_profile_library()
                item_id = profile_manager.create_evidence_item(
                    category="Project",
                    title=(
                        "CyberSphere (Unity Engine, Team of 2, "
                        "Published on Google Play)"
                    ),
                    description="Built gameplay systems.",
                    tools=["Unity Engine", "C#"],
                )
                self.assertEqual(
                    profile_manager.migrate_legacy_project_titles_to_structured_metadata(),
                    1,
                )
                item = profile_manager.get_evidence_item_by_id(item_id)
                self.assertIsNotNone(item)
                assert item is not None
                self.assertEqual(item["id"], item_id)
                self.assertEqual(item["title"], "CyberSphere")
                self.assertEqual(item["tools"], ["Unity Engine", "C#"])
                self.assertEqual(item["resume_header_tools"], ["Unity Engine"])
                self.assertEqual(
                    item["resume_header_context"],
                    ["Team of 2", "Published on Google Play"],
                )
            finally:
                profile_manager.DB_PATH = original_db

    def test_profile_preferences(self):
        original_db = profile_manager.DB_PATH
        with tempfile.TemporaryDirectory() as temporary:
            profile_manager.DB_PATH = Path(temporary) / "applications.db"
            try:
                profile_manager.init_user_profile_library()
                self.assertEqual(
                    profile_manager.get_resume_format_preferences(),
                    {"project_header_layout": "auto", "project_metadata_style": "pipes"},
                )
                profile_manager.update_resume_format_preferences(
                    project_header_layout="stacked",
                    project_metadata_style="parentheses",
                )
                self.assertEqual(
                    profile_manager.get_resume_format_preferences(),
                    {"project_header_layout": "stacked", "project_metadata_style": "parentheses"},
                )
            finally:
                profile_manager.DB_PATH = original_db


if __name__ == "__main__":
    unittest.main()
