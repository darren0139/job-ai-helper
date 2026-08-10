from __future__ import annotations

import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from resume_builder.docx_projects_skills_replacer import (
    _add_project_bullet_after,
)


def _attach_native_numbering(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")

    num_pr.append(ilvl)
    num_pr.append(num_id)
    ppr.append(num_pr)


class DocxProjectBulletPreservationTests(unittest.TestCase):
    def test_literal_bullet_template_preserves_one_text_bullet(self) -> None:
        doc = Document()
        template = doc.add_paragraph("• Original project bullet")
        anchor = doc.add_paragraph("Project title")

        created = _add_project_bullet_after(
            anchor,
            bullet="Tailored project evidence.",
            template=template,
        )

        self.assertEqual(created.text, "• Tailored project evidence.")
        self.assertTrue(
            created._p.pPr is None or created._p.pPr.numPr is None
        )

    def test_literal_bullet_is_not_duplicated_if_text_already_has_one(self) -> None:
        doc = Document()
        template = doc.add_paragraph("• Original project bullet")
        anchor = doc.add_paragraph("Project title")

        created = _add_project_bullet_after(
            anchor,
            bullet="• Tailored project evidence.",
            template=template,
        )

        self.assertEqual(created.text, "• Tailored project evidence.")
        self.assertNotIn("• •", created.text)

    def test_native_word_numbering_stays_native_without_text_bullet(self) -> None:
        doc = Document()
        template = doc.add_paragraph("Original project bullet")
        _attach_native_numbering(template)
        anchor = doc.add_paragraph("Project title")

        created = _add_project_bullet_after(
            anchor,
            bullet="Tailored project evidence.",
            template=template,
        )

        self.assertEqual(created.text, "Tailored project evidence.")
        self.assertIsNotNone(created._p.pPr)
        self.assertIsNotNone(created._p.pPr.numPr)
        self.assertFalse(created.text.startswith("•"))


if __name__ == "__main__":
    unittest.main()
