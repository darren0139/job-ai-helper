from __future__ import annotations

import copy
import hashlib
import json
import sqlite3
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from docx import Document

from analysis_stability.stable_evidence_scoring import SCORING_VERSION
from database import tailoring_version_manager as base_manager
from database.global_blueprint_manager import (
    init_global_blueprint_registry,
    list_active_global_blueprints_read_only,
)
from rag.jd_identity import source_version_id
from tailoring.capability_taxonomy import get_default_taxonomy
from tailoring.phase9c_blueprint_evaluation import _target_analysis
from tailoring.phase9f_jd_intake import build_transient_exact_jd_snapshot
from tailoring.phase9f_starting_source_artifacts import (
    Phase9FBArtifactError,
    resolve_starting_source_artifacts,
)
from tailoring.phase9f_starting_source_ranking import (
    BASE_RESUME_CONTENT_POLICY_VERSION,
    BASE_RESUME_FORMAT_VERSION,
    BASE_RESUME_VERSION_POLICY_VERSION,
    BLUEPRINT_FORMAT_VERSION,
    BLUEPRINT_IDENTITY_POLICY_VERSION,
    PHASE9F_B_EVIDENCE_POLICY_VERSION,
    PHASE9F_B_RANKING_POLICY_VERSION,
    ROLE_FAMILY_NEAR_TIE_TOLERANCES,
    canonical_json,
    fingerprint_value,
    normalize_active_blueprint_source,
    order_scored_candidates,
    prepare_ranking_context,
    rank_prepared_context,
    rank_starting_resume_sources,
    score_normalized_source,
    validate_ranked_candidate_analysis_snapshot,
)


REPO_ROOT = Path(__file__).resolve().parents[1]
GOLDEN = REPO_ROOT / "ci_fixtures" / (
    "phase9f_starting_source_ranking_golden.json"
)
PINNED_IDENTITY = REPO_ROOT / "ci_fixtures" / (
    "phase9f_b_ranking_identity_pinned_v1.json"
)


JD_TEXT = """AI Full-Stack Software Engineer

Responsibilities
- Build full-stack user-facing software applications with Python and React.
- Design secure backend APIs and relational database workflows.

Requirements
- Python backend API development.
- React frontend application development.
- PostgreSQL relational database design.
- Authentication workflows and secure database access.

Preferred Qualifications
- Row-Level Security policies and PostgREST integration.
"""

JD_PROFILE = {
    "job_title": "AI Full-Stack Software Engineer",
    "company": "Synthetic Product Lab",
    "location": "Singapore",
    "experience_level": "Junior",
    "responsibilities": [
        "Build full-stack user-facing software applications with Python and React.",
        "Design secure backend APIs and relational database workflows.",
    ],
    "required_skills": [
        "Python backend API development",
        "React frontend application development",
        "PostgreSQL relational database design",
        "Authentication workflows and secure database access",
    ],
    "preferred_skills": [
        "Row-Level Security policies and PostgREST integration"
    ],
    "tools_technologies": [
        "Python",
        "React",
        "PostgreSQL",
        "PostgREST",
    ],
    "soft_skills": [],
    "buzzwords": [],
    "deal_breakers": [],
}


def make_exact_jd(
    *,
    source_type: str = "pasted",
    library_jd_id: int = 0,
) -> dict:
    return build_transient_exact_jd_snapshot(
        raw_text=JD_TEXT,
        jd_profile=copy.deepcopy(JD_PROFILE),
        source_type=source_type,
        source_url=(
            "https://example.test/saved" if source_type == "saved" else ""
        ),
        library_jd_id=library_jd_id or None,
        saved_source_version_id=(
            source_version_id(JD_TEXT) if source_type == "saved" else ""
        ),
        model_calls=[],
    )


def resume_profile(*, strong: bool, marker: str = "") -> dict:
    bullet = (
        "Built full-stack user-facing React applications with Python backend "
        "APIs, PostgreSQL database design, authentication workflows, "
        "PostgREST, Row-Level Security policies, and secure database access."
        if strong
        else "Documented internal operational workflows and status updates."
    )
    skills = (
        {
            "languages": ["Python", "JavaScript"],
            "frameworks": ["React"],
            "tools": ["PostgreSQL", "PostgREST"],
            "concepts": ["Authentication", "Row-Level Security"],
            "platforms": [],
        }
        if strong
        else {
            "languages": ["C"],
            "frameworks": [],
            "tools": [],
            "concepts": ["Documentation"],
            "platforms": [],
        }
    )
    return {
        "name": f"Synthetic Candidate {marker}",
        "contact": {
            "email": "candidate@example.test",
            "phone": "",
            "linkedin": "",
            "github": "",
            "portfolio": "",
        },
        "summary": "Builds reliable software products.",
        "education": [
            {
                "school": "Example University",
                "degree": "BSc Computer Science",
                "graduation_date": "2026",
                "courses": [],
            }
        ],
        "experience": [
            {
                "title": "Software Engineer",
                "company": "Example Co",
                "date": "2025",
                "bullets": [bullet],
            }
        ],
        "projects": [
            {
                "title": "QueryAI" if strong else "Operations Notes",
                "date": "2026",
                "bullets": [bullet],
            }
        ],
        "skills": skills,
    }


def resume_text(profile: dict) -> str:
    lines = [str(profile.get("name") or "")]
    for section in ("experience", "projects"):
        for row in profile.get(section, []) or []:
            lines.append(str(row.get("title") or ""))
            lines.extend(str(value) for value in row.get("bullets", []) or [])
    for values in (profile.get("skills") or {}).values():
        lines.extend(str(value) for value in values or [])
    return "\n".join(lines)


def write_resume_artifacts(root: Path, text: str) -> tuple[Path, Path]:
    docx_path = root / "approved.docx"
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    document.save(docx_path)

    pdf_path = root / "approved.pdf"
    pdf = fitz.open()
    try:
        page = pdf.new_page(width=612, height=792)
        inserted = page.insert_textbox(
            fitz.Rect(40, 40, 572, 752),
            text,
            fontsize=7,
        )
        if inserted < 0:
            raise AssertionError("Synthetic test resume did not fit the PDF page.")
        pdf.save(pdf_path)
    finally:
        pdf.close()
    return docx_path, pdf_path


def resolved_blueprint_provenance(
    *,
    generation_id: str,
    artifact_hash_records: list[dict] | None = None,
) -> dict:
    return {
        "chain_status": "resolved",
        "source_resume_result_or_generation": {
            "source_generation": {
                "resolved": True,
                "approval_resolved": True,
                "fit_identity_match": True,
                "generation_id": generation_id,
            },
            "immutable_artifact_hash_records": artifact_hash_records or [],
        },
        "phase8_verification": {
            "resolved": True,
            "blueprint_ready": True,
        },
    }


def make_base(*, strong: bool, display_name: str = "Base Resume") -> tuple[dict, dict]:
    profile = resume_profile(strong=strong, marker="Base")
    text = resume_text(profile)
    artifact_bytes = b"synthetic-base-resume-docx"
    artifact_sha = hashlib.sha256(artifact_bytes).hexdigest()
    semantic = {
        "format_version": BASE_RESUME_FORMAT_VERSION,
        "content_policy_version": BASE_RESUME_CONTENT_POLICY_VERSION,
        "artifact": {
            "artifact_sha256": artifact_sha,
            "artifact_type": "docx",
            "artifact_size_bytes": len(artifact_bytes),
        },
        "resume_text": {
            "resume_text_sha256": hashlib.sha256(text.encode()).hexdigest(),
            "resume_text_char_count": len(text),
        },
        "structured_profile_fingerprint": fingerprint_value(profile),
        "profile_contract": {
            "extraction_policy_version": "synthetic",
            "resume_profile_prompt_sha256": "synthetic",
            "resume_profile_stability_version": "synthetic",
        },
    }
    content_fingerprint = fingerprint_value(semantic)
    version_identity = {
        "format_version": BASE_RESUME_FORMAT_VERSION,
        "version_policy_version": BASE_RESUME_VERSION_POLICY_VERSION,
        "master_content_fingerprint": content_fingerprint,
        "version_number": 1,
        "predecessor": {
            "master_version_id": "",
            "master_version_fingerprint": "",
        },
    }
    version_fingerprint = fingerprint_value(version_identity)
    version_id = version_fingerprint[:32]
    snapshot = {
        "format_version": BASE_RESUME_FORMAT_VERSION,
        "master_version_id": version_id,
        "master_version_fingerprint": version_fingerprint,
        "master_content_fingerprint": content_fingerprint,
        "version_number": 1,
        "predecessor": copy.deepcopy(version_identity["predecessor"]),
        "artifact": {
            "artifact_sha256": artifact_sha,
            "artifact_type": "docx",
            "artifact_size_bytes": len(artifact_bytes),
            "original_filename": "base.docx",
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        "resume_text": text,
        "resume_text_sha256": hashlib.sha256(text.encode()).hexdigest(),
        "resume_text_char_count": len(text),
        "structured_profile": copy.deepcopy(profile),
        "structured_profile_fingerprint": fingerprint_value(profile),
        "semantic_identity": copy.deepcopy(semantic),
        "version_identity": copy.deepcopy(version_identity),
        "extraction_provenance": {"call_count": 0},
    }
    row = {
        "master_version_id": version_id,
        "master_version_fingerprint": version_fingerprint,
        "master_content_fingerprint": content_fingerprint,
        "format_version": BASE_RESUME_FORMAT_VERSION,
        "content_policy_version": BASE_RESUME_CONTENT_POLICY_VERSION,
        "version_policy_version": BASE_RESUME_VERSION_POLICY_VERSION,
        "version_number": 1,
        "artifact_sha256": artifact_sha,
        "artifact_type": "docx",
        "artifact_size_bytes": len(artifact_bytes),
        "original_filename": "base.docx",
        "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "resume_text_sha256": hashlib.sha256(text.encode()).hexdigest(),
        "resume_text_char_count": len(text),
        "resume_text": text,
        "structured_profile_fingerprint": fingerprint_value(profile),
        "structured_profile": profile,
        "semantic_identity": semantic,
        "version_identity": version_identity,
        "master_snapshot": snapshot,
        "display_name": display_name,
        "created_at": "2026-08-15T00:00:00",
    }
    artifact = {
        "artifact_id": "base-artifact",
        "master_version_id": version_id,
        "artifact_kind": "original",
        "media_type": row["media_type"],
        "filename": "base.docx",
        "sha256": artifact_sha,
        "byte_size": len(artifact_bytes),
        "authoritative": True,
        "artifact_bytes": artifact_bytes,
        "created_at": "2026-08-15T00:00:00",
    }
    return row, artifact


def make_blueprint(
    *,
    strong: bool,
    role_family_id: str,
    role_family_label: str,
    marker: str,
    display_name: str | None = None,
    historical_score: int = 1,
    status: str = "active",
) -> dict:
    profile = resume_profile(strong=strong, marker=marker)
    text = resume_text(profile)
    frozen = {
        "resume_profile_snapshot": profile,
        "resume_text_snapshot": text,
    }
    semantic = {
        "phase9d_version": BLUEPRINT_FORMAT_VERSION,
        "fingerprint_policy_version": BLUEPRINT_IDENTITY_POLICY_VERSION,
        "role_family": {
            "role_family_id": role_family_id,
            "role_family_label": role_family_label,
        },
        "candidate": {
            "candidate_id": f"candidate-{marker}",
            "candidate_fingerprint": f"candidate-fingerprint-{marker}",
            "phase9b_version": "phase9b-blueprint-candidate-v3",
        },
        "evaluation": {
            "evaluation_id": f"evaluation-{marker}",
            "evaluation_fingerprint": f"evaluation-fingerprint-{marker}",
            "phase9c_version": "phase9c-cross-jd-evaluation-v1",
            "policy_version": "phase9c-same-family-explicit-scope-v3",
            "evidence_link_version": "phase9c-full-snapshot-evidence-v2",
            "scoring_version": SCORING_VERSION,
            "taxonomy_version": get_default_taxonomy().version,
        },
        "resume_snapshot": {
            "complete_snapshot_fingerprint": fingerprint_value(frozen),
            "resume_profile_snapshot_fingerprint": fingerprint_value(profile),
            "resume_text_snapshot_sha256": hashlib.sha256(text.encode()).hexdigest(),
        },
        "stable_input_provenance": [],
    }
    blueprint_fingerprint = fingerprint_value(semantic)
    blueprint_id = blueprint_fingerprint[:32]
    snapshot = {
        "phase9d_version": BLUEPRINT_FORMAT_VERSION,
        "fingerprint_policy_version": BLUEPRINT_IDENTITY_POLICY_VERSION,
        "blueprint_id": blueprint_id,
        "blueprint_fingerprint": blueprint_fingerprint,
        "role_family_id": role_family_id,
        "role_family_label": role_family_label,
        "semantic_identity": semantic,
        "frozen_resume_snapshot": frozen,
        "phase9b_candidate_semantic_snapshot": {},
        "phase9c_evaluation_snapshot": {
            "evaluation_id": f"evaluation-{marker}",
            "aggregate_result": {
                "provisional": False,
                "mean_score": historical_score,
            },
        },
        "phase9c_semantic_identity": {},
        "provenance": {},
    }
    return {
        "blueprint_id": blueprint_id,
        "blueprint_fingerprint": blueprint_fingerprint,
        "phase9d_version": BLUEPRINT_FORMAT_VERSION,
        "fingerprint_policy_version": BLUEPRINT_IDENTITY_POLICY_VERSION,
        "role_family_id": role_family_id,
        "role_family_label": role_family_label,
        "version_number": 1,
        "status": status,
        "candidate_id": f"candidate-{marker}",
        "candidate_fingerprint": f"candidate-fingerprint-{marker}",
        "evaluation_id": f"evaluation-{marker}",
        "evaluation_fingerprint": f"evaluation-fingerprint-{marker}",
        "semantic_identity": semantic,
        "blueprint_snapshot": snapshot,
        "display_name": display_name or f"{marker} Blueprint",
        "notes": "",
        "created_at": "2026-08-15T00:00:00",
    }


def metric_candidate(name: str, relationship: str, values: dict, confidence: str) -> dict:
    return {
        "source_type": "global_blueprint",
        "source_id": name,
        "source_version": 1,
        "source_fingerprint": f"source-{name}",
        "source_content_fingerprint": f"content-{name}",
        "normalized_source_fingerprint": hashlib.sha256(name.encode()).hexdigest(),
        "source_display_name": name,
        "source_role_family_id": name,
        "source_role_family_label": name,
        "role_family_relationship": relationship,
        "role_family_prior_eligible": bool(
            relationship == "same_family" and confidence in {"medium", "high"}
        ),
        "role_family_prior_applied": False,
        "ranking_reason": "canonical_metrics_order",
        "deterministic_alignment_score": values["overall"],
        "required_core_coverage_score": values["required_core"],
        "preferred_coverage_score": values["preferred"],
        "evidence_strength_score": values["evidence"],
        "important_gap_count": values["important_gaps"],
        "deal_breaker_gap_count": values["deal_breaker_gaps"],
        "stable_input_fingerprint": f"input-{name}",
        "comparison_result_fingerprint": f"result-{name}",
    }


class Phase9FStartingSourceRankingTests(unittest.TestCase):
    def setUp(self) -> None:
        self.jd = make_exact_jd()
        self.base, self.base_artifact = make_base(strong=False)
        self.same_family = make_blueprint(
            strong=True,
            role_family_id="ai_fullstack_software_engineering",
            role_family_label="AI & Full-Stack Software Engineering",
            marker="same",
        )

    def test_pinned_ranking_identity_contract(self):
        pinned = json.loads(PINNED_IDENTITY.read_text(encoding="utf-8"))
        base, artifact = make_base(strong=False)
        blueprint = make_blueprint(
            strong=True,
            role_family_id="ai_fullstack_software_engineering",
            role_family_label="AI & Full-Stack Software Engineering",
            marker="identity-pin",
        )
        result = rank_starting_resume_sources(
            exact_jd=make_exact_jd(),
            current_base_resume=base,
            current_base_artifact=artifact,
            global_blueprints=[blueprint],
        )
        actual = {
            "ranking_input_fingerprint": result[
                "ranking_input_fingerprint"
            ],
            "ranking_fingerprint": result["ranking_fingerprint"],
            "winner": result["recommended_source"][
                "normalized_source_fingerprint"
            ],
            "candidate_order": [
                row["normalized_source_fingerprint"]
                for row in result["ranked_candidates"]
            ],
            "candidate_metrics": [
                {
                    key: row[key]
                    for key in (
                        "normalized_source_fingerprint",
                        "deterministic_alignment_score",
                        "required_core_coverage_score",
                        "preferred_coverage_score",
                        "evidence_strength_score",
                        "important_gap_count",
                        "deal_breaker_gap_count",
                    )
                }
                for row in result["ranked_candidates"]
            ],
            "comparison_result_fingerprints": [
                row["comparison_result_fingerprint"]
                for row in result["ranked_candidates"]
            ],
        }
        self.assertEqual(actual, {
            key: value
            for key, value in pinned.items()
            if key != "fixture_version"
        })

    def test_complete_candidate_analysis_snapshot_is_retained_and_valid(self):
        result = self.rank()
        for candidate in result["ranked_candidates"]:
            validated = validate_ranked_candidate_analysis_snapshot(candidate)
            self.assertEqual(
                validated["candidate_analysis_snapshot_fingerprint"],
                candidate["candidate_analysis_snapshot_fingerprint"],
            )
            snapshot = validated["candidate_analysis_snapshot"]
            self.assertEqual(
                snapshot["stable_analysis_snapshot"][
                    "deterministic_alignment_score"
                ],
                candidate["deterministic_alignment_score"],
            )

    def test_candidate_analysis_snapshot_tampering_fails_closed(self):
        candidate = copy.deepcopy(self.rank()["recommended_source"])
        candidate["candidate_analysis_snapshot"]["resume_text_snapshot"] += (
            " tampered"
        )
        with self.assertRaisesRegex(ValueError, "fingerprint"):
            validate_ranked_candidate_analysis_snapshot(candidate)

    def rank(self, *, base=True, blueprints=None, jd=None):
        return rank_starting_resume_sources(
            exact_jd=jd or self.jd,
            current_base_resume=self.base if base else None,
            current_base_artifact=self.base_artifact if base else None,
            global_blueprints=(
                [self.same_family] if blueprints is None else blueprints
            ),
        )

    def test_semantically_identical_jd_provenance_does_not_change_ranking(self):
        pasted = make_exact_jd(source_type="pasted")
        saved = make_exact_jd(source_type="saved", library_jd_id=42)
        first = self.rank(jd=pasted)
        second = self.rank(jd=saved)
        self.assertEqual(first["status"], "ranked")
        self.assertEqual(first["ranking_input_fingerprint"], second["ranking_input_fingerprint"])
        self.assertEqual(first["ranking_fingerprint"], second["ranking_fingerprint"])
        self.assertNotEqual(first["jd_provenance"], second["jd_provenance"])

    def test_corrupt_current_base_fails_complete_ranking_closed(self):
        corrupt = copy.deepcopy(self.base)
        corrupt["resume_text"] += " corrupt"
        result = rank_starting_resume_sources(
            exact_jd=self.jd,
            current_base_resume=corrupt,
            current_base_artifact=self.base_artifact,
            global_blueprints=[self.same_family],
        )
        self.assertEqual(result["status"], "integrity_failed")
        self.assertIsNone(result["recommended_source"])
        self.assertEqual(result["ranked_candidates"], [])
        self.assertTrue(any("base_resume" == row["source_type"] for row in result["integrity_issues"]))

    def test_corrupt_active_blueprint_fails_complete_ranking_closed(self):
        corrupt = copy.deepcopy(self.same_family)
        corrupt["blueprint_snapshot"]["frozen_resume_snapshot"]["resume_text_snapshot"] += " corrupt"
        result = self.rank(blueprints=[corrupt])
        self.assertEqual(result["status"], "integrity_failed")
        self.assertIsNone(result["recommended_source"])
        self.assertTrue(any("global_blueprint" == row["source_type"] for row in result["integrity_issues"]))

    def test_missing_base_is_valid_absence(self):
        result = self.rank(base=False)
        self.assertEqual(result["status"], "ranked")
        self.assertEqual(result["recommended_source"]["source_type"], "global_blueprint")

    def test_blueprint_can_beat_a_weaker_base(self):
        result = self.rank()
        self.assertEqual(result["status"], "ranked")
        self.assertEqual(
            result["recommended_source"]["source_id"],
            self.same_family["blueprint_id"],
        )

    def test_base_can_beat_a_weaker_blueprint(self):
        strong_base, strong_artifact = make_base(strong=True)
        weak_blueprint = make_blueprint(
            strong=False,
            role_family_id="ai_fullstack_software_engineering",
            role_family_label="AI & Full-Stack Software Engineering",
            marker="weak-same-family",
        )
        result = rank_starting_resume_sources(
            exact_jd=self.jd,
            current_base_resume=strong_base,
            current_base_artifact=strong_artifact,
            global_blueprints=[weak_blueprint],
        )
        self.assertEqual(result["status"], "ranked")
        self.assertEqual(
            result["recommended_source"]["source_type"],
            "base_resume",
        )

    def test_no_eligible_sources_is_explicit_and_has_no_winner(self):
        result = self.rank(base=False, blueprints=[])
        self.assertEqual(result["status"], "no_eligible_sources")
        self.assertIsNone(result["recommended_source"])
        self.assertEqual(result["ranked_candidates"], [])

    def test_superseded_blueprint_is_intentionally_excluded(self):
        superseded = copy.deepcopy(self.same_family)
        superseded["status"] = "superseded"
        result = self.rank(blueprints=[superseded])
        self.assertEqual(result["status"], "ranked")
        self.assertEqual(result["recommended_source"]["source_type"], "base_resume")
        self.assertEqual(result["excluded_sources"][0]["reason"], "superseded_or_inactive_blueprint")

    def test_historical_approval_scores_do_not_influence_current_jd_winner(self):
        weak_high_history = make_blueprint(
            strong=False,
            role_family_id="frontend_software_engineering",
            role_family_label="Frontend Software Engineering",
            marker="weak-history",
            historical_score=100,
        )
        strong_low_history = make_blueprint(
            strong=True,
            role_family_id="backend_cloud_software_engineering",
            role_family_label="Backend & Cloud Software Engineering",
            marker="strong-history",
            historical_score=1,
        )
        result = self.rank(base=False, blueprints=[weak_high_history, strong_low_history])
        self.assertEqual(result["recommended_source"]["source_id"], strong_low_history["blueprint_id"])

    def test_fresh_blueprint_score_matches_phase9c_target_semantics(self):
        normalized = normalize_active_blueprint_source(self.same_family)
        context = prepare_ranking_context(
            exact_jd=self.jd,
            current_base_resume=None,
            current_base_artifact=None,
            global_blueprints=[self.same_family],
        )
        scored = score_normalized_source(normalized, context["_exact_jd"])
        candidate = {
            "resume_profile_snapshot": copy.deepcopy(
                normalized["resume_profile_snapshot"]
            ),
            "resume_text_snapshot": normalized["resume_text_snapshot"],
        }
        phase9c = _target_analysis(
            candidate,
            self.jd,
            self.jd["canonicalisation"],
        )
        for field in (
            "deterministic_alignment_score",
            "required_core_coverage_score",
            "preferred_coverage_score",
            "evidence_strength_score",
        ):
            self.assertEqual(scored[field], int(phase9c[field]))
        self.assertEqual(
            scored["canonical_requirement_scope_fingerprint"],
            self.jd["canonical_requirement_fingerprint"],
        )

    def test_distinct_visible_evidence_changes_canonical_outcomes(self):
        weak_blueprint = make_blueprint(
            strong=False,
            role_family_id="backend_cloud_software_engineering",
            role_family_label="Backend & Cloud Software Engineering",
            marker="distinct-weak",
        )
        context = prepare_ranking_context(
            exact_jd=self.jd,
            current_base_resume=None,
            current_base_artifact=None,
            global_blueprints=[self.same_family, weak_blueprint],
        )
        scored = [
            score_normalized_source(source, context["_exact_jd"])
            for source in context["_normalized_sources"]
        ]
        by_id = {row["source_id"]: row for row in scored}
        strong = by_id[self.same_family["blueprint_id"]]
        weak = by_id[weak_blueprint["blueprint_id"]]
        self.assertNotEqual(
            strong["canonical_requirement_results"],
            weak["canonical_requirement_results"],
        )
        self.assertNotEqual(
            strong["comparison_result_fingerprint"],
            weak["comparison_result_fingerprint"],
        )

    def test_golden_calibration_cases_and_selected_tolerances(self):
        fixture = json.loads(GOLDEN.read_text(encoding="utf-8"))
        self.assertEqual(fixture["selected_tolerances"], ROLE_FAMILY_NEAR_TIE_TOLERANCES)
        for case in fixture["cases"]:
            with self.subTest(case=case["name"]):
                rows = [
                    metric_candidate(
                        "cross_family",
                        "cross_family",
                        case["cross_family"],
                        case["jd_confidence"],
                    ),
                    metric_candidate(
                        "same_family",
                        "same_family",
                        case["same_family"],
                        case["jd_confidence"],
                    ),
                ]
                ordered = order_scored_candidates(rows)
                self.assertEqual(ordered[0]["source_id"], case["expected_winner"])
                self.assertEqual(ordered[0]["role_family_prior_applied"], case["prior_applied"])

    def test_materially_stronger_cross_family_blueprint_wins(self):
        cross = make_blueprint(
            strong=True,
            role_family_id="backend_cloud_software_engineering",
            role_family_label="Backend & Cloud Software Engineering",
            marker="cross-strong",
        )
        same_weak = make_blueprint(
            strong=False,
            role_family_id="ai_fullstack_software_engineering",
            role_family_label="AI & Full-Stack Software Engineering",
            marker="same-weak",
        )
        result = self.rank(base=False, blueprints=[same_weak, cross])
        self.assertEqual(result["recommended_source"]["source_id"], cross["blueprint_id"])
        self.assertFalse(result["recommended_source"]["role_family_prior_applied"])

    def test_reverse_enumeration_is_order_independent(self):
        cross = make_blueprint(
            strong=False,
            role_family_id="backend_cloud_software_engineering",
            role_family_label="Backend & Cloud Software Engineering",
            marker="cross",
        )
        first = self.rank(blueprints=[self.same_family, cross])
        second = self.rank(blueprints=[cross, self.same_family])
        self.assertEqual(first["ranking_input_fingerprint"], second["ranking_input_fingerprint"])
        self.assertEqual(first["ranking_fingerprint"], second["ranking_fingerprint"])
        self.assertEqual(
            [row["source_id"] for row in first["ranked_candidates"]],
            [row["source_id"] for row in second["ranked_candidates"]],
        )

    def test_display_name_does_not_change_semantic_identity(self):
        renamed_base = copy.deepcopy(self.base)
        renamed_base["display_name"] = "Renamed Base"
        renamed_blueprint = copy.deepcopy(self.same_family)
        renamed_blueprint["display_name"] = "Renamed Blueprint"
        first = self.rank()
        second = rank_starting_resume_sources(
            exact_jd=self.jd,
            current_base_resume=renamed_base,
            current_base_artifact=self.base_artifact,
            global_blueprints=[renamed_blueprint],
        )
        self.assertEqual(first["ranking_input_fingerprint"], second["ranking_input_fingerprint"])
        self.assertEqual(first["ranking_fingerprint"], second["ranking_fingerprint"])

    def test_changed_semantic_inputs_make_cached_result_stale(self):
        first_context = prepare_ranking_context(
            exact_jd=self.jd,
            current_base_resume=self.base,
            current_base_artifact=self.base_artifact,
            global_blueprints=[self.same_family],
        )
        first = rank_prepared_context(first_context)
        changed = make_blueprint(
            strong=False,
            role_family_id="backend_cloud_software_engineering",
            role_family_label="Backend & Cloud Software Engineering",
            marker="new-active",
        )
        second_context = prepare_ranking_context(
            exact_jd=self.jd,
            current_base_resume=self.base,
            current_base_artifact=self.base_artifact,
            global_blueprints=[self.same_family, changed],
        )
        self.assertNotEqual(
            first["ranking_input_fingerprint"],
            second_context["ranking_input_fingerprint"],
        )

    def test_zero_cost_and_no_input_mutation(self):
        original_jd = canonical_json(self.jd)
        original_base = canonical_json(self.base)
        original_blueprint = canonical_json(self.same_family)
        module_source = (REPO_ROOT / "tailoring" / "phase9f_starting_source_ranking.py").read_text(encoding="utf-8")
        self.assertNotIn("import llm", module_source)
        self.assertNotIn("jd_chroma_rag", module_source)
        self.assertNotIn("global_blueprint_manager", module_source)
        self.assertNotIn("global_master_resume_manager", module_source)
        with patch("database.tailoring_version_manager._connect", side_effect=AssertionError("persistence access")):
            result = self.rank()
        self.assertEqual(
            result["zero_cost_diagnostics"],
            {
                "model_call_count": 0,
                "embedding_call_count": 0,
                "chroma_read_count": 0,
                "chroma_write_count": 0,
                "persistence_write_count": 0,
            },
        )
        self.assertEqual(canonical_json(self.jd), original_jd)
        self.assertEqual(canonical_json(self.base), original_base)
        self.assertEqual(canonical_json(self.same_family), original_blueprint)

    def test_policy_and_evidence_versions_are_in_semantic_identity(self):
        result = self.rank()
        semantic = result["semantic_identity"]
        self.assertEqual(
            semantic["ranking_policy"]["policy_version"],
            PHASE9F_B_RANKING_POLICY_VERSION,
        )
        self.assertEqual(
            semantic["scoring"]["evidence_policy_version"],
            PHASE9F_B_EVIDENCE_POLICY_VERSION,
        )
        self.assertEqual(
            semantic["scoring"]["scoring_version"],
            SCORING_VERSION,
        )

    def test_active_blueprint_loader_succeeds_with_sqlite_query_only(self):
        with tempfile.TemporaryDirectory() as temporary:
            old_path = base_manager.DB_PATH
            base_manager.DB_PATH = Path(temporary) / "phase9f_b.sqlite"
            try:
                init_global_blueprint_registry()
                connection = sqlite3.connect(base_manager.DB_PATH)
                connection.row_factory = sqlite3.Row
                connection.execute("PRAGMA query_only = ON")
                with patch(
                    "database.global_blueprint_manager._connect",
                    return_value=connection,
                ):
                    self.assertEqual(
                        list_active_global_blueprints_read_only(),
                        [],
                    )
            finally:
                base_manager.DB_PATH = old_path

    def test_blueprint_artifacts_resolve_from_immutable_fit_provenance(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            blueprint = copy.deepcopy(self.same_family)
            frozen = blueprint["blueprint_snapshot"]["frozen_resume_snapshot"]
            docx_path, pdf_path = write_resume_artifacts(
                root,
                frozen["resume_text_snapshot"],
            )
            candidate = blueprint["blueprint_snapshot"][
                "phase9b_candidate_semantic_snapshot"
            ]
            candidate.update(
                {
                    "source_generation_id": "approved-generation",
                    "source_verification_fingerprint": "phase8-fingerprint",
                    "fit_result": {
                        "generation_id": "bound-fit-result-generation",
                        "fit_one_page": True,
                        "page_count": 1,
                        "docx_path": str(docx_path),
                        "pdf_path": str(pdf_path),
                    },
                }
            )
            context = prepare_ranking_context(
                exact_jd=self.jd,
                current_base_resume=None,
                current_base_artifact=None,
                global_blueprints=[blueprint],
            )
            result = rank_prepared_context(context)
            resolved = resolve_starting_source_artifacts(
                ranked_candidate=result["recommended_source"],
                normalized_source=context["_normalized_sources"][0],
                current_base_artifact=None,
                current_base_preview_artifact=None,
                global_blueprints=[blueprint],
                blueprint_provenance=resolved_blueprint_provenance(
                    generation_id="approved-generation"
                ),
            )
            self.assertEqual(
                [row["artifact_type"] for row in resolved["artifacts"]],
                ["docx", "pdf"],
            )
            self.assertEqual(
                resolved["preview_pdf"]["artifact_bytes"],
                pdf_path.read_bytes(),
            )
            self.assertTrue(
                all(
                    row["verification_method"]
                    == "frozen_snapshot_token_multiset_v1"
                    for row in resolved["artifacts"]
                )
            )

    def test_changed_blueprint_artifact_fails_frozen_snapshot_verification(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            blueprint = copy.deepcopy(self.same_family)
            frozen = blueprint["blueprint_snapshot"]["frozen_resume_snapshot"]
            docx_path, pdf_path = write_resume_artifacts(
                root,
                frozen["resume_text_snapshot"],
            )
            document = Document()
            document.add_paragraph(
                "Unrelated replacement resume content " * 20
            )
            document.save(docx_path)
            candidate = blueprint["blueprint_snapshot"][
                "phase9b_candidate_semantic_snapshot"
            ]
            candidate.update(
                {
                    "source_generation_id": "approved-generation",
                    "source_verification_fingerprint": "phase8-fingerprint",
                    "fit_result": {
                        "generation_id": "fit-generation",
                        "fit_one_page": True,
                        "page_count": 1,
                        "docx_path": str(docx_path),
                        "pdf_path": str(pdf_path),
                    },
                }
            )
            context = prepare_ranking_context(
                exact_jd=self.jd,
                current_base_resume=None,
                current_base_artifact=None,
                global_blueprints=[blueprint],
            )
            result = rank_prepared_context(context)
            with self.assertRaisesRegex(
                Phase9FBArtifactError,
                "could not be safely verified",
            ):
                resolve_starting_source_artifacts(
                    ranked_candidate=result["recommended_source"],
                    normalized_source=context["_normalized_sources"][0],
                    current_base_artifact=None,
                    current_base_preview_artifact=None,
                    global_blueprints=[blueprint],
                    blueprint_provenance=resolved_blueprint_provenance(
                        generation_id="approved-generation"
                    ),
                )

    def test_authoritative_blueprint_hash_rejects_byte_change(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            blueprint = copy.deepcopy(self.same_family)
            frozen = blueprint["blueprint_snapshot"]["frozen_resume_snapshot"]
            docx_path, pdf_path = write_resume_artifacts(
                root,
                frozen["resume_text_snapshot"],
            )
            candidate = blueprint["blueprint_snapshot"][
                "phase9b_candidate_semantic_snapshot"
            ]
            candidate.update(
                {
                    "source_generation_id": "approved-generation",
                    "source_verification_fingerprint": "phase8-fingerprint",
                    "fit_result": {
                        "generation_id": "fit-generation",
                        "fit_one_page": True,
                        "page_count": 1,
                        "docx_path": str(docx_path),
                        "pdf_path": str(pdf_path),
                    },
                }
            )
            records = [
                {
                    "artifact_kind": kind,
                    "artifact_sha256": hashlib.sha256(
                        path.read_bytes()
                    ).hexdigest(),
                    "artifact_size": path.stat().st_size,
                }
                for kind, path in (("docx", docx_path), ("pdf", pdf_path))
            ]
            context = prepare_ranking_context(
                exact_jd=self.jd,
                current_base_resume=None,
                current_base_artifact=None,
                global_blueprints=[blueprint],
            )
            result = rank_prepared_context(context)
            provenance = resolved_blueprint_provenance(
                generation_id="approved-generation",
                artifact_hash_records=records,
            )
            resolved = resolve_starting_source_artifacts(
                ranked_candidate=result["recommended_source"],
                normalized_source=context["_normalized_sources"][0],
                current_base_artifact=None,
                current_base_preview_artifact=None,
                global_blueprints=[blueprint],
                blueprint_provenance=provenance,
            )
            self.assertTrue(
                all(
                    row["verification_method"]
                    == "authoritative_immutable_application_result_sha256"
                    for row in resolved["artifacts"]
                )
            )

            docx_path.write_bytes(docx_path.read_bytes() + b"changed")
            with self.assertRaisesRegex(
                Phase9FBArtifactError,
                "could not be safely verified",
            ):
                resolve_starting_source_artifacts(
                    ranked_candidate=result["recommended_source"],
                    normalized_source=context["_normalized_sources"][0],
                    current_base_artifact=None,
                    current_base_preview_artifact=None,
                    global_blueprints=[blueprint],
                    blueprint_provenance=provenance,
                )

    def test_missing_blueprint_artifact_fails_read_only(self):
        blueprint = copy.deepcopy(self.same_family)
        candidate = blueprint["blueprint_snapshot"][
            "phase9b_candidate_semantic_snapshot"
        ]
        candidate.update(
            {
                "source_generation_id": "approved-generation",
                "source_verification_fingerprint": "phase8-fingerprint",
                "fit_result": {
                    "generation_id": "fit-generation",
                    "fit_one_page": True,
                    "page_count": 1,
                    "docx_path": "missing-approved.docx",
                    "pdf_path": "missing-approved.pdf",
                },
            }
        )
        context = prepare_ranking_context(
            exact_jd=self.jd,
            current_base_resume=None,
            current_base_artifact=None,
            global_blueprints=[blueprint],
        )
        result = rank_prepared_context(context)
        with self.assertRaisesRegex(
            Phase9FBArtifactError,
            "No approved fitted Blueprint artifact",
        ):
            resolve_starting_source_artifacts(
                ranked_candidate=result["recommended_source"],
                normalized_source=context["_normalized_sources"][0],
                current_base_artifact=None,
                current_base_preview_artifact=None,
                global_blueprints=[blueprint],
                blueprint_provenance=resolved_blueprint_provenance(
                    generation_id="approved-generation"
                ),
            )

    def test_base_resume_artifact_still_uses_authoritative_stored_hash(self):
        context = prepare_ranking_context(
            exact_jd=self.jd,
            current_base_resume=self.base,
            current_base_artifact=self.base_artifact,
            global_blueprints=[],
        )
        result = rank_prepared_context(context)
        resolved = resolve_starting_source_artifacts(
            ranked_candidate=result["recommended_source"],
            normalized_source=context["_normalized_sources"][0],
            current_base_artifact=self.base_artifact,
            current_base_preview_artifact=None,
            global_blueprints=[],
        )
        self.assertEqual(
            resolved["artifacts"][0]["sha256"],
            self.base_artifact["sha256"],
        )

        changed = copy.deepcopy(self.base_artifact)
        changed["artifact_bytes"] += b"changed"
        with self.assertRaisesRegex(
            Phase9FBArtifactError,
            "failed hash or size validation",
        ):
            resolve_starting_source_artifacts(
                ranked_candidate=result["recommended_source"],
                normalized_source=context["_normalized_sources"][0],
                current_base_artifact=changed,
                current_base_preview_artifact=None,
                global_blueprints=[],
            )


if __name__ == "__main__":
    unittest.main()
